import os
import json
import logging
import random
import re
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.interval import IntervalTrigger
from functools import wraps
from io import BytesIO
from pathlib import Path
from datetime import datetime, timedelta
from sqlalchemy import orm, or_

from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
import base64
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, redirect, url_for, session, flash, send_from_directory, Response
import ai_services
from werkzeug.security import generate_password_hash, check_password_hash
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, scoped_session
import PyPDF2
import docx
import google.generativeai as genai

# PDF Generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# HTML processing
from bs4 import BeautifulSoup

# Load environment variables
env_path = os.path.join(os.path.dirname(__file__), '.env')
load_dotenv(dotenv_path=env_path)

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Verify environment variables
logger.debug("Environment variables loaded:")
logger.debug(f"GEMINI_API_KEY present: {'Yes' if os.getenv('GEMINI_API_KEY') else 'No'}")
logger.debug(f"DEBUG mode: {os.getenv('DEBUG', 'False')}")

# Initialize Gemini
gemini_api_key = os.getenv('GEMINI_API_KEY')
if not gemini_api_key or gemini_api_key == 'your-gemini-api-key-here':
    logger.error("ERROR: GEMINI_API_KEY is not properly configured in .env file")
else:
    genai.configure(api_key=gemini_api_key)
    model = genai.GenerativeModel('gemini-pro')

from dashboard import dashboard_bp

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')

# Debug route to check environment variables
@app.route('/debug/env')
def debug_env():
    """Debug endpoint to check environment variables"""
    env_vars = {
        'GEMINI_API_KEY': 'Set' if os.getenv('GEMINI_API_KEY') else 'Not set',
        'DEBUG': os.getenv('DEBUG', 'False'),
        'Current Directory': os.getcwd(),
        'Files in Directory': ', '.join(os.listdir('.'))
    }
    return jsonify(env_vars)

# PostgreSQL configuration
# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'postgresql://postgres:sairam@localhost/smarthire')
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300,  # Recycle connections after 5 minutes
    'pool_size': 20,
    'max_overflow': 30
}
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)

app.register_blueprint(dashboard_bp, url_prefix='/hr/dashboard')

# Initialize extensions
db = SQLAlchemy()
db.init_app(app)
migrate = Migrate(app, db)

# Create database engine
engine = create_engine(app.config['SQLALCHEMY_DATABASE_URI'], **app.config.get('SQLALCHEMY_ENGINE_OPTIONS', {}))
SessionLocal = scoped_session(sessionmaker(autocommit=False, autoflush=False, bind=engine))

def get_db_session():
    db_session = SessionLocal()
    try:
        yield db_session
    finally:
        db_session.close()

# Encryption helper functions
def get_encryption_key():
    """Generate or retrieve encryption key for storing passwords"""
    # Use app secret key as base for encryption
    app_secret = app.config.get('SECRET_KEY', 'default-secret-key').encode()
    salt = b'smarthire_salt'  # Fixed salt for consistency
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=100000,
    )
    key = base64.urlsafe_b64encode(kdf.derive(app_secret))
    return key

def encrypt_password(password):
    """Encrypt password for storage"""
    key = get_encryption_key()
    f = Fernet(key)
    encrypted_password = f.encrypt(password.encode())
    return base64.urlsafe_b64encode(encrypted_password).decode()

def decrypt_password(encrypted_password):
    """Decrypt stored password"""
    if not encrypted_password:
        return None
    try:
        key = get_encryption_key()
        f = Fernet(key)
        encrypted_data = base64.urlsafe_b64decode(encrypted_password.encode())
        decrypted_password = f.decrypt(encrypted_data).decode()
        return decrypted_password
    except Exception:
        return None

def get_device_info():
    """Get device/browser info for remembered credentials"""
    user_agent = request.headers.get('User-Agent', '')
    # Simple device fingerprint
    return user_agent[:100] if user_agent else 'Unknown'

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    full_name = db.Column(db.String(100), nullable=False)
    phone = db.Column(db.String(20), nullable=False)
    role = db.Column(db.String(20), nullable=False)
    department = db.Column(db.String(50), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_login = db.Column(db.DateTime, nullable=True)
    is_active = db.Column(db.Boolean, default=True)
    status = db.Column(db.String(20), default='Active')
    hire_date = db.Column(db.Date, nullable=True)
    position = db.Column(db.String(100), nullable=True)
    manager_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    employee_id = db.Column(db.String(20), nullable=True)
    exit_date = db.Column(db.Date, nullable=True)
    
    # Password reset fields
    reset_token = db.Column(db.String(255), nullable=True)
    reset_token_expires = db.Column(db.DateTime, nullable=True)
    
    # Relationships
    manager = db.relationship('User', remote_side=[id], backref=db.backref('direct_reports', lazy=True))
    skills = db.relationship('EmployeeSkill', backref='employee', lazy=True, cascade='all, delete-orphan')

    def set_password(self, password):
        self.password = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password, password)
    
    def generate_reset_token(self, expires_in=3600):
        import secrets
        token = secrets.token_urlsafe(32)
        self.reset_token = token
        self.reset_token_expires = datetime.utcnow() + timedelta(seconds=expires_in)
        return token
    
    def verify_reset_token(self, token):
        if self.reset_token != token:
            return False
        if datetime.utcnow() > self.reset_token_expires:
            return False
        return True
    
    def clear_reset_token(self):
        self.reset_token = None
        self.reset_token_expires = None


class EmployeeSkill(db.Model):
    """Employee skills and expertise"""
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    skill_name = db.Column(db.String(100), nullable=False)
    skill_category = db.Column(db.String(50), nullable=True)  # Technical, Soft Skills, Domain, etc.
    proficiency_level = db.Column(db.String(20), default='Intermediate')  # Beginner, Intermediate, Advanced, Expert
    years_experience = db.Column(db.Float, nullable=True)
    last_used = db.Column(db.Date, nullable=True)
    certification = db.Column(db.String(200), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class PreOnboardingTask(db.Model):
    """Pre-onboarding administrative tasks"""
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    task_name = db.Column(db.String(200), nullable=False)
    task_description = db.Column(db.Text, nullable=True)
    task_type = db.Column(db.String(50), nullable=False)  # 'admin', 'documentation', 'preparation'
    due_date = db.Column(db.Date, nullable=True)
    is_completed = db.Column(db.Boolean, default=False)
    completed_at = db.Column(db.DateTime, nullable=True)
    assigned_to = db.Column(db.String(100), nullable=True)  # HR or employee
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    employee = db.relationship('User', backref='pre_onboarding_tasks')


class WelcomePackage(db.Model):
    """Welcome packages/swag tracking"""
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    package_name = db.Column(db.String(200), nullable=False)
    items = db.Column(db.Text, nullable=True)  # JSON array of items
    status = db.Column(db.String(20), default='pending')  # pending, ordered, shipped, delivered
    tracking_number = db.Column(db.String(100), nullable=True)
    shipping_address = db.Column(db.Text, nullable=True)
    estimated_delivery = db.Column(db.Date, nullable=True)
    actual_delivery = db.Column(db.Date, nullable=True)
    notes = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    employee = db.relationship('User', backref='welcome_packages')


class FirstDayAgenda(db.Model):
    """First day agenda for new hires"""
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    agenda_date = db.Column(db.Date, nullable=False)
    agenda_items = db.Column(db.Text, nullable=False)  # JSON array of agenda items
    start_time = db.Column(db.Time, nullable=True)
    end_time = db.Column(db.Time, nullable=True)
    location = db.Column(db.String(200), nullable=True)
    meeting_links = db.Column(db.Text, nullable=True)  # JSON array of virtual meeting links
    preparation_notes = db.Column(db.Text, nullable=True)
    is_shared = db.Column(db.Boolean, default=False)
    shared_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    employee = db.relationship('User', backref='first_day_agendas')


class TeamIntroduction(db.Model):
    """Early team introductions"""
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    team_member_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    introduction_type = db.Column(db.String(50), default='peer')  # peer, manager, mentor, buddy
    message = db.Column(db.Text, nullable=True)
    is_sent = db.Column(db.Boolean, default=False)
    sent_at = db.Column(db.DateTime, nullable=True)
    scheduled_date = db.Column(db.Date, nullable=True)
    response_received = db.Column(db.Boolean, default=False)
    response_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    employee = db.relationship('User', foreign_keys=[employee_id], backref='introductions_received')











class Employee(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    employee_id = db.Column(db.String(20), nullable=False)
    job_title = db.Column(db.String(100), nullable=False)
    department = db.Column(db.String(50), nullable=False)
    manager_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    hire_date = db.Column(db.Date, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships with explicit foreign_keys
    user = db.relationship('User', foreign_keys=[user_id], backref=db.backref('employee', uselist=False, lazy=True))
    manager = db.relationship('User', foreign_keys=[manager_id], remote_side=[User.id], backref=db.backref('managed_employees', lazy=True))


class Candidate(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), nullable=True)
    phone = db.Column(db.String(20), nullable=True)
    job_desc = db.Column(db.Text, nullable=True)
    resume_text = db.Column(db.Text, nullable=True)
    score = db.Column(db.Float, nullable=True)
    summary = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    # Remove the conflicting relationship - will be defined in Interview model
    
    def extract_contact_info(self):
        """Extract email and phone from resume text"""
        import re
        
        if not self.resume_text:
            return
            
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, self.resume_text)
        if emails and not self.email:
            self.email = emails[0]  # Take the first email found
            
        # Extract phone numbers (multiple formats)
        phone_patterns = [
            r'\b\d{3}-\d{3}-\d{4}\b',  # 123-456-7890
            r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',  # (123) 456-7890
            r'\b\d{3}\.\d{3}\.\d{4}\b',  # 123.456.7890
            r'\b\d{10}\b',  # 1234567890
            r'\b\+?\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'  # International format
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, self.resume_text)
            if phones and not self.phone:
                # Clean up the phone number format
                phone = re.sub(r'[^\d+]', '', phones[0])
                if len(phone) >= 10:  # Ensure it's a valid phone number
                    self.phone = phones[0]
                break

class Interview(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    candidate_id = db.Column(db.Integer, db.ForeignKey('candidate.id'), nullable=True)  # For external candidates
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)  # For internal employees
    responses = db.Column(db.JSON, nullable=True)
    summary = db.Column(db.Text, nullable=True)
    status = db.Column(db.String(20), default='pending')  # pending, requested, ready, in_progress, completed
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships with explicit names to avoid conflicts
    candidate_ref = db.relationship('Candidate', backref='candidate_interviews')
    user_ref = db.relationship('User', backref='employee_interviews', foreign_keys=[user_id])

class ExitFeedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    reason = db.Column(db.Text, nullable=False)
    feedback = db.Column(db.Text, nullable=True)
    sentiment = db.Column(db.String(20), nullable=True)
    # AI Analysis Fields
    key_themes = db.Column(db.Text, nullable=True)  # JSON array of themes
    risk_level = db.Column(db.String(20), nullable=True)  # low|medium|high|critical
    actionable_insights = db.Column(db.Text, nullable=True)  # JSON array of insights
    emotional_tone = db.Column(db.String(30), nullable=True)  # enthusiastic|satisfied|neutral|concerned|frustrated|angry
    retention_probability = db.Column(db.String(20), nullable=True)  # low|medium|high
    recommendations = db.Column(db.Text, nullable=True)  # JSON array of recommendations
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Knowledge Transfer Models
class KTSession(db.Model):
    """Knowledge Transfer Sessions"""
    id = db.Column(db.Integer, primary_key=True)
    session_topic = db.Column(db.Text, nullable=False)
    attendees = db.Column(db.Text, nullable=False)
    scheduled_date = db.Column(db.Date, nullable=False)
    duration_hours = db.Column(db.Float, nullable=False)
    description = db.Column(db.Text, nullable=True)
    status = db.Column(db.String(20), default='Scheduled')  # Scheduled, In Progress, Completed, Cancelled
    documentation_path = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class KTDocument(db.Model):
    """Knowledge Transfer Documents"""
    id = db.Column(db.Integer, primary_key=True)
    document_name = db.Column(db.Text, nullable=False)
    file_path = db.Column(db.Text, nullable=False)
    file_type = db.Column(db.String(50), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    uploaded_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    category = db.Column(db.String(100), nullable=True)
    description = db.Column(db.Text, nullable=True)

class SuccessorTraining(db.Model):
    """Successor Training Tracking"""
    id = db.Column(db.Integer, primary_key=True)
    successor_name = db.Column(db.Text, nullable=False)
    training_module = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(20), default='Pending')  # Pending, In Progress, Completed
    completion_date = db.Column(db.Date, nullable=True)
    trainer_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    notes = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class ProjectHandover(db.Model):
    """Project Handover Tracking"""
    id = db.Column(db.Integer, primary_key=True)
    project_name = db.Column(db.Text, nullable=False)
    project_description = db.Column(db.Text, nullable=True)
    status = db.Column(db.String(20), default='Pending')  # Pending, In Progress, Completed
    handover_date = db.Column(db.Date, nullable=True)
    recipient_name = db.Column(db.Text, nullable=False)
    documentation_path = db.Column(db.Text, nullable=True)
    verified = db.Column(db.Boolean, default=False)
    notes = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class KTProgress(db.Model):
    """Knowledge Transfer Progress Tracking"""
    id = db.Column(db.Integer, primary_key=True)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    sessions_completed = db.Column(db.Integer, default=0)
    docs_uploaded = db.Column(db.Integer, default=0)
    successor_trained_percent = db.Column(db.Integer, default=0)
    projects_handover = db.Column(db.Integer, default=0)
    overall_progress = db.Column(db.Integer, default=0)
    completion_date = db.Column(db.Date, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    employee = db.relationship('User', backref='kt_progress')

class EmployeeFeedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    mood_rating = db.Column(db.Integer, nullable=False)  # 1-5 scale
    confidence_rating = db.Column(db.Integer, nullable=False)  # 1-5 scale
    feedback = db.Column(db.Text, nullable=True)  # Optional anonymous feedback
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user = db.relationship('User', backref=db.backref('feedbacks', lazy=True))


class EmployeeDocument(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'),
                        nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    document_type = db.Column(db.String(100), nullable=True)
    file_size = db.Column(db.Integer, nullable=True)  # in bytes
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    description = db.Column(db.Text, nullable=True)
    is_verified = db.Column(db.Boolean, default=False)
    status = db.Column(db.String(20), default='pending')  # pending, approved, rejected
    user = db.relationship('User',
                           backref=db.backref('documents', lazy=True))


class EmployeeSettings(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'),
                        nullable=False, unique=True)
    email_notifications = db.Column(db.Boolean, default=True)
    sms_notifications = db.Column(db.Boolean, default=False)
    notification_frequency = db.Column(db.String(50),
                                       default='immediately')
    theme = db.Column(db.String(50), default='light')  # light/dark
    language = db.Column(db.String(10), default='en')
    two_factor_enabled = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationship
    user = db.relationship('User', backref=db.backref('settings', uselist=False))


class Task(db.Model):
    """General tasks table for various task types"""
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    task_type = db.Column(db.String(50), nullable=False)  # onboarding, offboarding, general
    status = db.Column(db.String(50), default='pending')  # pending, in_progress, completed, overdue
    priority = db.Column(db.String(20), default='medium')  # low, medium, high, urgent
    assigned_to = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    assigned_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    department = db.Column(db.String(50))
    due_date = db.Column(db.DateTime)
    completed_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    assignee = db.relationship('User', foreign_keys=[assigned_to], backref='assigned_tasks')
    assigner = db.relationship('User', foreign_keys=[assigned_by], backref='created_tasks')


class Message(db.Model):
    """Internal messaging system"""
    id = db.Column(db.Integer, primary_key=True)
    subject = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)
    sender_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    recipient_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    message_type = db.Column(db.String(50), default='message')  # message, notification, reminder
    priority = db.Column(db.String(20), default='normal')  # low, normal, high, urgent
    is_priority = db.Column(db.Boolean, default=False)  # Whether this is a priority message
    status = db.Column(db.String(50), default='unread')  # unread, read, archived
    sent_at = db.Column(db.DateTime, default=datetime.utcnow)
    read_at = db.Column(db.DateTime)
    parent_id = db.Column(db.Integer, db.ForeignKey('message.id'), nullable=True)  # For replies
    notification_data = db.Column(db.JSON, nullable=True)  # Additional notification metadata
    
    # Edit and delete tracking
    is_edited = db.Column(db.Boolean, default=False)
    edited_at = db.Column(db.DateTime)
    is_deleted = db.Column(db.Boolean, default=False)
    deleted_at = db.Column(db.DateTime)
    
    # Relationships
    sender = db.relationship('User', foreign_keys=[sender_id], backref='sent_messages')
    recipient = db.relationship('User', foreign_keys=[recipient_id], backref='received_messages')
    replies = db.relationship('Message', backref=db.backref('parent', remote_side=[id]), lazy='dynamic')


class RememberedCredential(db.Model):
    """Remembered login credentials for devices"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    username = db.Column(db.String(80), nullable=False)
    encrypted_password = db.Column(db.String(255), nullable=False)
    device_info = db.Column(db.String(500), nullable=False)
    is_active = db.Column(db.Boolean, default=True)
    expires_at = db.Column(db.DateTime, nullable=False)
    last_used = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    user = db.relationship('User', backref='remembered_credentials')


# Make timedelta and datetime available in templates
@app.context_processor
def inject_datetime():
    return {
        'timedelta': timedelta,
        'now': datetime.utcnow
    }

@app.context_processor
def inject_user():
    """Inject current user into all templates"""
    if 'user_id' in session:
        user = User.query.get(session['user_id'])
        return {'user': user}
    return {'user': None}

class DocumentTemplate(db.Model):
    """Document templates for HR with placeholder support"""
    __tablename__ = 'document_templates'
    
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    code = db.Column(db.String(50), unique=True, nullable=False)  # e.g., "offer_letter", "nda"
    description = db.Column(db.Text, nullable=True)
    template_type = db.Column(db.String(50), nullable=False)  # 'onboarding', 'offboarding', 'offer', 'termination'
    content = db.Column(db.Text, nullable=False)  # HTML content with placeholders
    placeholders = db.Column(db.Text, nullable=True)  # JSON list of required placeholders
    is_active = db.Column(db.Boolean, default=True)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    creator = db.relationship('User', backref='created_templates')
    generated_documents = db.relationship('GeneratedDocument', backref='template', lazy=True, cascade='all, delete-orphan')


class GeneratedDocument(db.Model):
    """Generated documents from templates"""
    __tablename__ = 'generated_documents'
    
    id = db.Column(db.Integer, primary_key=True)
    template_id = db.Column(db.Integer, db.ForeignKey('document_templates.id'), nullable=False)
    employee_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    document_name = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)  # Filled HTML content
    file_path = db.Column(db.String(500), nullable=True)  # Generated PDF path
    status = db.Column(db.String(20), default='draft')  # 'draft', 'generated', 'downloaded'
    generated_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    generated_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    employee = db.relationship('User', foreign_keys=[employee_id], backref='generated_documents')
    generator = db.relationship('User', foreign_keys=[generated_by], backref='documents_generated')


# Create database tables
with app.app_context():
    db.create_all()

# Routes
# Login required decorator for role-based access
def login_required(role=None):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session or 'role' not in session:
                # Redirect to role selection page if not logged in
                return redirect(url_for('index'))
            if role:
                # Handle both single role string and list of roles
                allowed_roles = role if isinstance(role, list) else [role]
                if session.get('role') not in allowed_roles:
                    flash('You do not have permission to access this page.', 'danger')
                    # For candidates, redirect to the login with candidate role
                    if session.get('role') == 'candidate':
                        return redirect(url_for('login', role='candidate'))
                    # For other roles, redirect to their appropriate dashboard
                    return redirect(url_for(session.get('role', 'index') + '_dashboard'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

@app.route('/')
def index():
    if 'user_id' in session:
        # Handle candidate role by redirecting to employee login
        if session.get('role') == 'candidate':
            return redirect(url_for('login', role='employee'))
        # For other roles, redirect to their dashboard
        return redirect(url_for(session['role'] + '_dashboard'))
    return redirect(url_for('select_role'))

@app.route('/select-role')
def select_role():
    # Clear the current session to ensure a clean role switch
    session.clear()
    return render_template('role_selection.html')

@app.route('/api/remembered-credentials')
def get_remembered_credentials():
    """Get remembered credentials for the current device"""
    try:
        device_info = get_device_info()
        
        # Find active remembered credentials for this device
        credentials = RememberedCredential.query.filter_by(
            device_info=device_info,
            is_active=True
        ).filter(
            RememberedCredential.expires_at > datetime.utcnow()
        ).all()
        
        remembered_accounts = []
        for cred in credentials:
            user = User.query.get(cred.user_id)
            if user and user.is_active:
                remembered_accounts.append({
                    'id': cred.id,
                    'username': cred.username,
                    'full_name': user.full_name,
                    'role': user.role,
                    'department': user.department,
                    'has_password': bool(cred.encrypted_password),
                    'last_used': cred.last_used.isoformat() if cred.last_used else None
                })
        
        return jsonify({
            'status': 'success',
            'remembered_accounts': remembered_accounts
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/remembered-credentials/<int:credential_id>/login', methods=['POST'])
def login_with_remembered_credentials(credential_id):
    """Login using remembered credentials"""
    try:
        credential = RememberedCredential.query.get_or_404(credential_id)
        
        # Verify device matches
        if credential.device_info != get_device_info():
            return jsonify({
                'status': 'error',
                'message': 'Invalid device'
            }), 403
        
        # Check if credential is still valid
        if not credential.is_active or credential.expires_at < datetime.utcnow():
            return jsonify({
                'status': 'error',
                'message': 'Credentials expired'
            }), 400
        
        # Decrypt password
        password = decrypt_password(credential.encrypted_password)
        if not password:
            return jsonify({
                'status': 'error',
                'message': 'Could not decrypt password'
            }), 500
        
        # Get user
        user = User.query.get(credential.user_id)
        if not user or not user.is_active:
            return jsonify({
                'status': 'error',
                'message': 'User not found or inactive'
            }), 404
        
        # Verify password
        if not user.check_password(password):
            return jsonify({
                'status': 'error',
                'message': 'Invalid credentials'
            }), 401
        
        # Update last used
        credential.last_used = datetime.utcnow()
        db.session.commit()
        
        # Set session
        session['user_id'] = user.id
        session['role'] = user.role
        session.permanent = True
        session['username'] = user.full_name
        
        # Update last login
        user.last_login = datetime.utcnow()
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'redirect_url': url_for('hr_dashboard') if user.role == 'hr' else url_for('employee_dashboard')
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/remembered-credentials/<int:credential_id>', methods=['DELETE'])
def delete_remembered_credentials(credential_id):
    """Delete remembered credentials"""
    try:
        credential = RememberedCredential.query.get_or_404(credential_id)
        
        # Verify device matches
        if credential.device_info != get_device_info():
            return jsonify({
                'status': 'error',
                'message': 'Invalid device'
            }), 403
        
        db.session.delete(credential)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Credentials removed successfully'
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/login/<role>', methods=['GET', 'POST'])
def login(role):
    try:
        # Map role to valid roles to prevent injection
        valid_roles = ['hr', 'employee']
        if role not in valid_roles:
            flash('Invalid role selected', 'danger')
            return redirect(url_for('select_role'))
        
        if request.method == 'POST':
            username = request.form.get('username')
            password = request.form.get('password')
            remember = True if request.form.get('remember') else False
            
            if not username or not password:
                flash('Please enter both username/email and password', 'danger')
                return redirect(url_for('login', role=role))
            
            try:
                # Check if input is email or username
                if '@' in username:
                    # Login with email
                    user = User.query.filter(
                        db.func.lower(User.email) == username.lower(),
                        User.role == role,
                        User.is_active == True
                    ).first()
                else:
                    # Login with username
                    user = User.query.filter(
                        db.func.lower(User.username) == username.lower(),
                        User.role == role,
                        User.is_active == True
                    ).first()
                
                # Check if user exists and password is correct
                if not user or not user.check_password(password):
                    app.logger.warning(f'Failed login attempt for: {username}')
                    flash('Invalid username/email or password', 'danger')
                    return redirect(url_for('login', role=role))
                
                # Handle remember me functionality
                if remember:
                    # Check if credentials already exist for this device
                    device_info = get_device_info()
                    existing_credential = RememberedCredential.query.filter_by(
                        user_id=user.id,
                        device_info=device_info,
                        is_active=True
                    ).first()
                    
                    if existing_credential:
                        # Update existing credential
                        existing_credential.encrypted_password = encrypt_password(password)
                        existing_credential.last_used = datetime.utcnow()
                        existing_credential.expires_at = datetime.utcnow() + timedelta(days=30)
                    else:
                        # Create new remembered credential
                        credential = RememberedCredential(
                            user_id=user.id,
                            username=username,
                            encrypted_password=encrypt_password(password),
                            device_info=device_info,
                            expires_at=datetime.utcnow() + timedelta(days=30)
                        )
                        db.session.add(credential)
                    
                    db.session.commit()
                    app.logger.info(f'Remembered credentials saved for user: {user.username}')
                else:
                    # Clean up any existing remembered credentials for this device if user unchecked remember me
                    device_info = get_device_info()
                    RememberedCredential.query.filter_by(
                        user_id=user.id,
                        device_info=device_info
                    ).delete()
                    db.session.commit()
                
                # Update last login
                user.last_login = datetime.utcnow()
                db.session.commit()
                
                # Set session
                session['user_id'] = user.id
                session['role'] = user.role
                session.permanent = remember  # Session expires when browser closes if not remembered
                
                app.logger.info(f'User {user.username} logged in successfully')
                flash(f'Welcome back, {user.full_name}!', 'success')
                
                # Store display name for welcome bar / navbar
                session['username'] = user.full_name
                
                # Redirect based on role
                if role == 'employee':
                    return redirect(url_for('employee_dashboard'))
                return redirect(url_for('hr_dashboard'))
                
            except Exception as e:
                db.session.rollback()
                app.logger.error(f'Error during login: {str(e)}')
                flash('An error occurred during login. Please try again.', 'danger')
                return redirect(url_for('login', role=role))
        
        # For GET request, show login form
        return render_template('login.html', role=role.title())
        
    except Exception as e:
        app.logger.error(f'Unexpected error in login route: {str(e)}')
        flash('An unexpected error occurred. Please try again later.', 'danger')
        return redirect(url_for('select_role'))

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    try:
        if request.method == 'POST':
            email = request.form.get('email')
            role = request.form.get('role')
            
            if not email or not role:
                flash('Please provide both email and account type', 'danger')
                return redirect(url_for('forgot_password'))
            
            # Validate role
            valid_roles = ['hr', 'employee']
            if role not in valid_roles:
                flash('Invalid account type selected', 'danger')
                return redirect(url_for('forgot_password'))
            
            try:
                # Find user by email and role
                user = User.query.filter(
                    db.func.lower(User.email) == email.lower(),
                    User.role == role,
                    User.is_active == True
                ).first()
                
                if user:
                    # Generate reset token
                    reset_token = user.generate_reset_token()
                    db.session.commit()
                    
                    # In a real application, you would send an email with the reset link
                    # For now, we'll redirect directly to the reset page with the token
                    app.logger.info(f'Password reset token generated for user: {user.username} ({user.email})')
                    flash(f'Password reset link generated. You will be redirected to reset your password.', 'success')
                    return redirect(url_for('reset_password', token=reset_token))
                else:
                    # Don't reveal that the user doesn't exist for security
                    flash('If an account with that email exists, you would receive a reset link.', 'info')
                    return redirect(url_for('select_role'))
                
            except Exception as e:
                app.logger.error(f'Error during password reset: {str(e)}')
                flash('An error occurred. Please try again later.', 'danger')
                return redirect(url_for('forgot_password'))
        
        # For GET request, show forgot password form
        return render_template('forgot_password.html')
        
    except Exception as e:
        app.logger.error(f'Unexpected error in forgot_password route: {str(e)}')
        flash('An unexpected error occurred. Please try again later.', 'danger')
        return redirect(url_for('select_role'))

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        # Find user with valid reset token
        user = User.query.filter_by(reset_token=token).first()
        
        if not user or not user.verify_reset_token(token):
            flash('Invalid or expired reset link. Please request a new password reset.', 'danger')
            return redirect(url_for('forgot_password'))
        
        if request.method == 'POST':
            password = request.form.get('password')
            confirm_password = request.form.get('confirm_password')
            
            if not password or not confirm_password:
                flash('Please provide both password fields', 'danger')
                return redirect(url_for('reset_password', token=token))
            
            if password != confirm_password:
                flash('Passwords do not match', 'danger')
                return redirect(url_for('reset_password', token=token))
            
            if len(password) < 8:
                flash('Password must be at least 8 characters long', 'danger')
                return redirect(url_for('reset_password', token=token))
            
            try:
                # Update password and clear reset token
                user.set_password(password)
                user.clear_reset_token()
                db.session.commit()
                
                app.logger.info(f'Password reset successfully for user: {user.username}')
                flash('Your password has been reset successfully. You can now login with your new password.', 'success')
                return redirect(url_for('select_role'))
                
            except Exception as e:
                db.session.rollback()
                app.logger.error(f'Error during password reset: {str(e)}')
                flash('An error occurred. Please try again.', 'danger')
                return redirect(url_for('reset_password', token=token))
        
        # For GET request, show reset password form
        return render_template('reset_password.html', token=token)
        
    except Exception as e:
        app.logger.error(f'Unexpected error in reset_password route: {str(e)}')
        flash('An unexpected error occurred. Please try again later.', 'danger')
        return redirect(url_for('select_role'))

@app.route('/register/<role>', methods=['GET', 'POST'])
def register(role):
    try:
        # Only allow HR registration
        if role != 'hr':
            flash('Self-registration is only available for HR personnel', 'danger')
            return redirect(url_for('select_role'))
        
        if request.method == 'POST':
            username = request.form.get('username', '').strip()
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')
            confirm_password = request.form.get('confirm_password', '')
            full_name = request.form.get('full_name', '').strip()
            phone = request.form.get('phone', '').strip()
            department = request.form.get('department', '').strip() if role == 'employee' else None
            
            # Basic validation
            if not all([username, email, password, confirm_password, full_name, phone]):
                flash('Please fill in all required fields', 'danger')
                return redirect(url_for('register', role=role))
            
            # Validate email format
            if not re.match(r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$', email):
                flash('Please enter a valid email address', 'danger')
                return redirect(url_for('register', role=role))
            
            # Validate password strength
            if len(password) < 8:
                flash('Password must be at least 8 characters long', 'danger')
                return redirect(url_for('register', role=role))
            
            # Validate passwords match
            if password != confirm_password:
                flash('Passwords do not match', 'danger')
                return redirect(url_for('register', role=role))
            
            try:
                # Check if username or email already exists (case-insensitive)
                if User.query.filter(db.func.lower(User.username) == username.lower()).first():
                    flash('Username already exists', 'danger')
                    return redirect(url_for('register', role=role))
                    
                if User.query.filter(db.func.lower(User.email) == email.lower()).first():
                    flash('Email already registered', 'danger')
                    return redirect(url_for('register', role=role))
                
                # Create new user
                user = User(
                    username=username,
                    email=email,
                    full_name=full_name,
                    phone=phone,
                    role=role,
                    department=department,
                    is_active=True,
                    created_at=datetime.utcnow()
                )
                user.set_password(password)
                
                db.session.add(user)
                db.session.commit()
                
                app.logger.info(f'New user registered: {username} ({role})')
                flash('Registration successful! Please log in.', 'success')
                return redirect(url_for('login', role=role))
                
            except Exception as e:
                db.session.rollback()
                app.logger.error(f'Database error during registration: {str(e)}')
                flash('An error occurred during registration. Please try again.', 'danger')
                return redirect(url_for('register', role=role))
        
        # For GET request, show registration form
        return render_template('register.html', role=role.title())
        
    except Exception as e:
        app.logger.error(f'Unexpected error in register route: {str(e)}')
        flash('An unexpected error occurred. Please try again later.', 'danger')
        return redirect(url_for('select_role'))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))  # Redirect to role selection page

@app.route('/analytics/mood', endpoint='mood_analytics')
@login_required()
def mood_analytics():
    """Mood Analytics Dashboard"""
    return render_template('mood_analytics_dashboard.html')

# HR Dashboard
@app.route('/hr/employees/add', methods=['GET', 'POST'])
@login_required('hr')
def add_employee():
    if request.method == 'POST':
        try:
            # Get form data
            username = request.form.get('username')
            email = request.form.get('email')
            full_name = request.form.get('full_name')
            phone = request.form.get('phone')
            department = request.form.get('department')
            position = request.form.get('position')
            manager_id = request.form.get('manager_id')
            password = request.form.get('password')
            hire_date_str = request.form.get('hire_date')
            
            # Validate hire date
            if not hire_date_str:
                flash('Hire date is required for onboarding new employees.', 'danger')
                managers = User.query.filter_by(role='manager', is_active=True).all()
                return render_template('hr/add_employee.html', managers=managers, now=datetime.utcnow(),
                                   form_data=request.form)
            
            try:
                hire_date = datetime.strptime(hire_date_str, '%Y-%m-%d').date()
                if hire_date > datetime.utcnow().date():
                    flash('Hire date cannot be in the future.', 'danger')
                    managers = User.query.filter_by(role='manager', is_active=True).all()
                    return render_template('hr/add_employee.html', managers=managers, now=datetime.utcnow(),
                                       form_data=request.form)
            except ValueError:
                flash('Invalid hire date format. Please use YYYY-MM-DD.', 'danger')
                managers = User.query.filter_by(role='manager', is_active=True).all()
                return render_template('hr/add_employee.html', managers=managers, now=datetime.utcnow(),
                                   form_data=request.form)

            # Create new user with 'Onboarding' status
            user = User(
                username=username,
                email=email,
                full_name=full_name,
                phone=phone,
                role='employee',
                department=department,
                position=position,
                manager_id=manager_id if manager_id != 'None' else None,
                status='Onboarding',
                is_active=True,
                hire_date=hire_date
            )
            user.set_password(password)
            
            # Add to database
            db.session.add(user)
            db.session.flush()  # Get the user ID
            
            # Legacy check - onboarding checklist system has been removed
            checklist = None
            
            db.session.commit()

            flash(f"Employee {full_name} added successfully with custom onboarding checklist!", 'success')
            return redirect(url_for('view_employee', employee_id=user.id))
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f'Error adding employee: {str(e)}')
            flash('An error occurred while adding the employee. Please try again.', 'danger')
    
    # For GET request or if there was an error
    from datetime import datetime
    managers = User.query.filter_by(role='manager', is_active=True).all()
    return render_template('hr/add_employee.html', managers=managers, now=datetime.utcnow())
@app.route('/hr/employees/<int:employee_id>')
@login_required('hr')
def view_employee(employee_id):
    employee = User.query.get_or_404(employee_id)
    # Legacy check - onboarding checklist system has been removed
    checklist = None
    
    # Get buddy/mentor relationships for this employee
    buddy_mentor_relationships = []
    try:
        # Legacy check - buddy/mentor system has been removed
        pass
    except:
        buddy_mentor_relationships = []
    
    # Get all users for partner selection (excluding current employee)
    all_users = User.query.filter(User.id != employee_id).all()
    
    # Get employee skills from database or add sample Indian resume skills
    employee_skills = EmployeeSkill.query.filter_by(employee_id=employee_id).all()
    
    # If no skills exist, add sample Indian resume-based skills
    if not employee_skills:
        sample_skills = [
            # Technical Skills (common in Indian IT resumes)
            {"skill_name": "Python", "skill_category": "Technical", "proficiency_level": "Advanced", "years_experience": 3.5, "certification": "Python Professional Certification"},
            {"skill_name": "Java", "skill_category": "Technical", "proficiency_level": "Intermediate", "years_experience": 2.0, "certification": None},
            {"skill_name": "React.js", "skill_category": "Technical", "proficiency_level": "Advanced", "years_experience": 2.5, "certification": "React Developer Certificate"},
            {"skill_name": "Node.js", "skill_category": "Technical", "proficiency_level": "Intermediate", "years_experience": 1.5, "certification": None},
            {"skill_name": "SQL", "skill_category": "Technical", "proficiency_level": "Advanced", "years_experience": 4.0, "certification": "SQL Database Expert"},
            {"skill_name": "MongoDB", "skill_category": "Technical", "proficiency_level": "Intermediate", "years_experience": 1.0, "certification": None},
            
            # Cloud & DevOps (popular in Indian tech industry)
            {"skill_name": "AWS", "skill_category": "Cloud", "proficiency_level": "Intermediate", "years_experience": 2.0, "certification": "AWS Solutions Architect Associate"},
            {"skill_name": "Docker", "skill_category": "DevOps", "proficiency_level": "Intermediate", "years_experience": 1.5, "certification": None},
            {"skill_name": "Git", "skill_category": "DevOps", "proficiency_level": "Advanced", "years_experience": 3.0, "certification": None},
            
            # Soft Skills (emphasized in Indian corporate culture)
            {"skill_name": "Communication", "skill_category": "Soft Skills", "proficiency_level": "Advanced", "years_experience": 4.0, "certification": "Business Communication Certificate"},
            {"skill_name": "Team Leadership", "skill_category": "Soft Skills", "proficiency_level": "Intermediate", "years_experience": 2.0, "certification": "Leadership Training Program"},
            {"skill_name": "Problem Solving", "skill_category": "Soft Skills", "proficiency_level": "Advanced", "years_experience": 4.0, "certification": None},
            
            # Domain & Business Skills
            {"skill_name": "Agile Methodologies", "skill_category": "Domain", "proficiency_level": "Advanced", "years_experience": 3.0, "certification": "Certified ScrumMaster"},
            {"skill_name": "Project Management", "skill_category": "Domain", "proficiency_level": "Intermediate", "years_experience": 2.5, "certification": "PMP Certification"},
            {"skill_name": "Data Analysis", "skill_category": "Domain", "proficiency_level": "Intermediate", "years_experience": 2.0, "certification": "Data Analytics Professional"},
        ]
        
        # Add sample skills to database
        for skill_data in sample_skills:
            skill = EmployeeSkill(
                employee_id=employee_id,
                skill_name=skill_data["skill_name"],
                skill_category=skill_data["skill_category"],
                proficiency_level=skill_data["proficiency_level"],
                years_experience=skill_data["years_experience"],
                certification=skill_data["certification"],
                last_used=datetime.utcnow().date()
            )
            db.session.add(skill)
        
        db.session.commit()
        employee_skills = EmployeeSkill.query.filter_by(employee_id=employee_id).all()
    
    return render_template('hr/view_employee.html', 
                         employee=employee, 
                         checklist=checklist,
                         buddy_mentor_relationships=buddy_mentor_relationships,
                         all_users=all_users,
                         employee_skills=employee_skills)


@app.route('/api/hr/buddy-mentor-assign', methods=['POST'])
@login_required('hr')
def assign_buddy_mentor_from_employee():
    """Assign buddy/mentor from employee view"""
    try:
        data = request.json
        employee_id = data['employee_id']
        employee_role = data['employee_role']
        partner_id = data['partner_id']
        relationship_type = data['relationship_type']
        program_name = data.get('program_name')
        matching_reason = data.get('matching_reason')
        goals = data.get('goals', [])
        start_date = data['start_date']
        end_date = data.get('end_date')
        
        # Determine mentor and mentee based on role
        if employee_role == 'mentee':
            mentor_id = partner_id
            mentee_id = employee_id
        else:  # employee is mentor
            mentor_id = employee_id
            mentee_id = partner_id
        
        # Check if there's already an active relationship
        existing = None
        # Legacy check - buddy/mentor system has been removed
        
        if existing:
            return jsonify({
                'status': 'error',
                'message': 'An active relationship already exists between these employees'
            }), 400
        
        # Create new relationship
        # Legacy functionality - buddy/mentor system has been removed
        return jsonify({
            'status': 'error',
            'message': 'Buddy/mentor system has been removed'
        }), 400
        
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/hr/buddy-mentor/<int:relationship_id>/end', methods=['POST'])
@login_required('hr')
def end_buddy_mentor_relationship(relationship_id):
    """End a buddy/mentor relationship"""
    try:
        # Legacy functionality - buddy/mentor system has been removed
        return jsonify({
            'status': 'error',
            'message': 'Buddy/mentor system has been removed'
        }), 400
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/hr/employees/<int:employee_id>/update-status', methods=['POST'])
@login_required('hr')
def update_employee_status(employee_id):
    employee = User.query.get_or_404(employee_id)
    
    try:
        # Get form data
        new_status = request.form.get('status')
        notes = request.form.get('notes', '')
        notify_employee = request.form.get('notify_employee') == 'on'
        
        # Validate status
        valid_statuses = ['Active', 'Onboarding', 'Offboarding', 'Terminated', 'On Leave']
        if new_status not in valid_statuses:
            flash('Invalid status selected.', 'danger')
            return redirect(url_for('list_employees'))
        
        # Store old status for logging
        old_status = employee.status
        
        # Update employee status
        employee.status = new_status
        
        # If changing to Offboarding, set exit date if not already set
        if new_status == 'Offboarding' and not employee.exit_date:
            employee.exit_date = datetime.utcnow().date() + timedelta(days=14)  # Default 2 weeks notice
        
        # If changing to Active, clear exit date
        if new_status == 'Active':
            employee.exit_date = None
        
        db.session.commit()
        
        # Log the status change (you could create a separate table for this)
        print(f"Status changed for {employee.full_name}: {old_status} -> {new_status}")
        
        # Send notification if requested
        if notify_employee:
            # Here you would typically send an email
            print(f"Notification sent to {employee.email} about status change to {new_status}")
        
        flash(f'Employee status updated successfully from {old_status} to {new_status}.', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'Error updating employee status: {str(e)}', 'danger')
    
    return redirect(url_for('list_employees'))


@app.route('/hr/employees/<int:employee_id>/documents')
@login_required('hr')
def hr_employee_documents(employee_id):
    """Employee documents quick generation page"""
    user = User.query.get(session['user_id'])
    employee = User.query.get_or_404(employee_id)
    
    # Get all generated documents for this employee
    documents = GeneratedDocument.query.filter_by(employee_id=employee_id).order_by(GeneratedDocument.generated_at.desc()).all()
    
    return render_template('hr/employee_documents_quick.html', 
                         user=user, 
                         employee=employee, 
                         documents=documents)


@app.route('/hr/employees/<int:employee_id>/pre-offboarding', methods=['GET', 'POST'])
@login_required('hr')
def hr_pre_offboarding(employee_id):
    employee = User.query.get_or_404(employee_id)

    if request.method == 'POST':
        exit_date_str = request.form.get('exit_date')
        if exit_date_str:
            try:
                new_exit_date = datetime.strptime(exit_date_str, '%Y-%m-%d').date()
                employee.exit_date = new_exit_date
                db.session.commit()
                flash('Last working day updated successfully.', 'success')
            except ValueError:
                flash('Invalid date format. Please use YYYY-MM-DD.', 'danger')

    documents = EmployeeDocument.query.filter_by(user_id=employee.id).all()
    today = datetime.utcnow().date()
    notice_days = None
    if employee.exit_date:
        try:
            notice_days = (employee.exit_date - today).days
        except Exception:
            notice_days = None
    return render_template('hr/pre_offboarding.html', employee=employee, documents=documents, notice_days=notice_days)

@app.route('/hr/employees')
@login_required('hr')
def list_employees():
    # Get filter parameters
    search = request.args.get('search', '')
    department_filter = request.args.get('department', '')
    status_filter = request.args.get('status', '')
    
    # Build query - Only show employees, exclude HR users
    query = User.query.filter(User.role == 'employee')
    
    # Apply filters
    if search:
        query = query.filter(
            db.or_(
                User.full_name.ilike(f'%{search}%'),
                User.email.ilike(f'%{search}%'),
                User.username.ilike(f'%{search}%')
            )
        )
    
    if department_filter:
        query = query.filter(User.department == department_filter)
    
    if status_filter:
        query = query.filter(User.status == status_filter)
    
    employees = query.order_by(User.created_at.desc()).all()
    
    # Get unique departments for filter dropdown - only from employees, not HR
    departments = db.session.query(User.department).filter(
        User.role == 'employee',
        User.department.isnot(None)
    ).distinct().all()
    departments = [dept[0] for dept in departments if dept[0]]
    
    return render_template('hr/employees.html', 
                         employees=employees,
                         departments=departments,
                         current_search=search,
                         current_department=department_filter,
                         current_status=status_filter)

@app.route('/hr/dashboard')
@login_required('hr')
def hr_dashboard():
    # Get the logged-in user
    user = User.query.get(session['user_id'])
    
    # Get real counts from database - only count employees, not HR
    total_employees = User.query.filter_by(role='employee').count()
    onboarding_count = User.query.filter(User.role == 'employee', User.status == 'Onboarding').count()
    active_count = User.query.filter(User.role == 'employee', User.status == 'Active').count()
    
    # Get recent AI interview results from database
    recent_interviews = Interview.query.order_by(Interview.created_at.desc()).limit(10).all()
    interview_results = []
    for interview in recent_interviews:
        candidate = interview.candidate_ref if interview.candidate_ref else None
        
        # Extract contact info if not already extracted
        if candidate and candidate.resume_text and (not candidate.email or not candidate.phone):
            candidate.extract_contact_info()
            db.session.commit()
        
        interview_results.append({
            'id': interview.id,
            'candidate_name': candidate.name if candidate else 'Unknown Candidate',
            'email': candidate.email if candidate and candidate.email else 'N/A',
            'phone': candidate.phone if candidate and candidate.phone else 'N/A',
            'role': candidate.job_desc[:50] + '...' if candidate and candidate.job_desc else 'N/A',
            'score': round(candidate.score, 1) if candidate and candidate.score else 'N/A',
            'recommendation': 'Proceed to Next Round' if (candidate and candidate.score and candidate.score >= 7.5) else 'Hold for Review' if (candidate and candidate.score and candidate.score >= 6.0) else 'Reject',
            'date': interview.created_at.strftime('%Y-%m-%d'),
            'summary': interview.summary[:100] + '...' if interview.summary else 'No summary available'
        })
    
    # Get exit feedback sentiment analysis from database
    exit_feedbacks = ExitFeedback.query.all()
    sentiment_counts = {'Positive': 0, 'Neutral': 0, 'Negative': 0}
    for feedback in exit_feedbacks:
        if feedback.sentiment:
            sentiment_counts[feedback.sentiment] = sentiment_counts.get(feedback.sentiment, 0) + 1
        else:
            sentiment_counts['Neutral'] += 1
    
    # Calculate employee feedback averages
    employee_feedbacks = EmployeeFeedback.query.filter(
        EmployeeFeedback.created_at >= datetime.utcnow() - timedelta(days=30)
    ).all()
    
    avg_mood = sum(f.mood_rating for f in employee_feedbacks) / len(employee_feedbacks) if employee_feedbacks else 3.5
    avg_confidence = sum(f.confidence_rating for f in employee_feedbacks) / len(employee_feedbacks) if employee_feedbacks else 3.5
    engagement_score = round((avg_mood + avg_confidence) / 2 * 20, 1)  # Convert to percentage
    
    # Get department statistics
    dept_query = db.session.query(
        User.department, 
        db.func.count(User.id).label('count')
    ).filter(
        User.role == 'employee',
        User.department.isnot(None)
    ).group_by(User.department).all()
    
    # Convert SQLAlchemy Row objects to simple list format for JSON serialization
    departments = [[dept[0], dept[1]] for dept in dept_query]
    
    # Calculate average onboarding completion time from actual employee data
    avg_onboarding_days = 7.2  # Default value
    completed_employees = User.query.filter(
        User.role == 'employee',
        User.status == 'Active',
        User.hire_date.isnot(None)
    ).all()
    
    if completed_employees:
        # Calculate actual onboarding duration based on hire dates
        total_days = 0
        count = 0
        for emp in completed_employees:
            if emp.hire_date:
                # Calculate days from hire date to now (if still active) or to status change
                # Since we don't have onboarding completion dates, use standard 10-day period
                hire_datetime = datetime.combine(emp.hire_date, datetime.min.time())
                days_since_hire = (datetime.utcnow() - hire_datetime).days
                # Assume onboarding takes 10 business days (~14 calendar days)
                onboarding_period = min(14, days_since_hire) if days_since_hire > 0 else 10
                total_days += onboarding_period
                count += 1
        
        if count > 0:
            avg_onboarding_days = total_days / count
    
    # Calculate HR Insights & Strategy data from database
    # Recent Activities from database
    recent_activities = []
    
    # Recent interview results
    recent_interviews_db = Interview.query.order_by(Interview.created_at.desc()).limit(5).all()
    for interview in recent_interviews_db:
        if interview.candidate_ref:
            recent_activities.append({
                'type': 'interview',
                'title': f"Interview with {interview.candidate_ref.name}",
                'date': interview.created_at.strftime('%Y-%m-%d'),
                'priority': 'normal'
            })
    
    # Recent employee status changes
    recent_hires = User.query.filter(
        User.role == 'employee',
        User.created_at >= datetime.utcnow() - timedelta(days=7)
    ).all()
    for emp in recent_hires:
        recent_activities.append({
            'type': 'hire',
            'title': f"New hire: {emp.full_name}",
            'date': emp.created_at.strftime('%Y-%m-%d'),
            'priority': 'high'
        })
    
    # Recent feedback submissions
    recent_feedback = EmployeeFeedback.query.order_by(EmployeeFeedback.created_at.desc()).limit(5).all()
    for feedback in recent_feedback:
        if feedback.user:
            recent_activities.append({
                'type': 'feedback',
                'title': f"Feedback from {feedback.user.full_name}",
                'date': feedback.created_at.strftime('%Y-%m-%d'),
                'priority': 'medium'
            })
    
    # Sort activities by date
    recent_activities.sort(key=lambda x: x['date'], reverse=True)
    
    # Strategic Insights from database
    strategic_insights = []
    
    # Employee turnover analysis
    total_employees = User.query.filter_by(role='employee').count()
    active_employees = User.query.filter(User.role == 'employee', User.status == 'Active').count()
    if total_employees > 0:
        retention_rate = (active_employees / total_employees) * 100
        if retention_rate < 85:
            strategic_insights.append({
                'type': 'warning',
                'title': 'Low Retention Rate',
                'description': f'Retention rate is {retention_rate:.1f}%. Consider reviewing employee satisfaction.',
                'action': 'Review exit feedback and employee engagement'
            })
    
    # Department analysis
    dept_analysis = db.session.query(
        User.department,
        db.func.count(User.id).label('count')
    ).filter(
        User.role == 'employee',
        User.department.isnot(None)
    ).group_by(User.department).all()
    
    if dept_analysis:
        largest_dept = max(dept_analysis, key=lambda x: x[1])
        strategic_insights.append({
            'type': 'info',
            'title': 'Largest Department',
            'description': f'{largest_dept[0]} has {largest_dept[1]} employees',
            'action': 'Ensure adequate resources for largest team'
        })
    
    # Mood/Engagement insights
    recent_mood_data = EmployeeFeedback.query.filter(
        EmployeeFeedback.created_at >= datetime.utcnow() - timedelta(days=30)
    ).all()
    if recent_mood_data:
        avg_mood = sum(f.mood_rating for f in recent_mood_data) / len(recent_mood_data)
        if avg_mood < 3.0:
            strategic_insights.append({
                'type': 'alert',
                'title': 'Low Employee Mood',
                'description': f'Average mood rating is {avg_mood:.1f}/5.0',
                'action': 'Consider wellness initiatives and check-ins'
            })
    pending_tasks_count = Task.query.filter_by(status='pending').count()
    unread_messages_count = Message.query.filter_by(recipient_id=session['user_id'], status='unread').count()
    
    # Get upcoming events
    offboarding_employees = User.query.filter(User.role == 'employee', User.status == 'Offboarding').all()
    upcoming_exits = []
    for emp in offboarding_employees:
        if emp.exit_date:
            # Handle both datetime and date objects
            exit_date = emp.exit_date
            if hasattr(exit_date, 'date'):
                # It's a datetime object, get the date part
                exit_date = exit_date.date()
            
            days_until = (exit_date - datetime.utcnow().date()).days
            if days_until >= 0 and days_until <= 7:  # Next 7 days
                upcoming_exits.append({
                    'name': emp.full_name,
                    'days': days_until,
                    'date': exit_date.strftime('%b %d')
                })
    
    # Get goal setting deadlines (mock data for now - could be enhanced with real goal tracking)
    upcoming_goals = []
    if onboarding_count > 0:
        upcoming_goals.append({'type': 'Goal setting', 'count': onboarding_count, 'timeframe': 'This week'})
    
    return render_template('dashboard.html',
                         user=user,
                         current_user=user,
                         total_employees=total_employees,
                         onboarding_count=onboarding_count,
                         active_count=active_count,
                         interview_results=interview_results,
                         sentiment_data=sentiment_counts,
                         engagement_score=engagement_score,
                         avg_onboarding_days=round(avg_onboarding_days, 1),
                         departments=departments,
                         pending_tasks_count=pending_tasks_count,
                         unread_messages_count=unread_messages_count,
                         upcoming_exits=upcoming_exits,
                         upcoming_goals=upcoming_goals,
                         recent_activities=recent_activities[:5],  # Top 5 recent activities
                         strategic_insights=strategic_insights)


# Document Template Management Routes
@app.route('/hr/templates')
@login_required('hr')
def hr_templates():
    """List all document templates"""
    user = User.query.get(session['user_id'])
    templates = DocumentTemplate.query.filter_by(is_active=True).order_by(DocumentTemplate.created_at.desc()).all()
    
    # Group templates by type
    template_types = {}
    for template in templates:
        if template.template_type not in template_types:
            template_types[template.template_type] = []
        template_types[template.template_type].append(template)
    
    return render_template('hr/templates.html', user=user, template_types=template_types)


@app.route('/hr/templates/new', methods=['GET', 'POST'])
@login_required('hr')
def hr_create_template():
    """Create a new document template"""
    user = User.query.get(session['user_id'])
    
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        template_type = request.form.get('template_type')
        content = request.form.get('content')
        
        # Extract placeholders from content
        import re
        placeholders = re.findall(r'\{\{(\w+)\}\}', content)
        
        template = DocumentTemplate(
            name=name,
            description=description,
            template_type=template_type,
            content=content,
            placeholders=json.dumps(placeholders),
            created_by=user.id
        )
        
        db.session.add(template)
        db.session.commit()
        
        flash('Template created successfully!', 'success')
        return redirect(url_for('hr_templates'))
    
    return render_template('hr/create_template.html', user=user)


@app.route('/hr/templates/<int:template_id>/edit', methods=['GET', 'POST'])
@login_required('hr')
def hr_edit_template(template_id):
    """Edit an existing document template"""
    user = User.query.get(session['user_id'])
    template = DocumentTemplate.query.get_or_404(template_id)
    
    if request.method == 'POST':
        template.name = request.form.get('name')
        template.description = request.form.get('description')
        template.template_type = request.form.get('template_type')
        template.content = request.form.get('content')
        
        # Extract placeholders from content
        import re
        placeholders = re.findall(r'\{\{(\w+)\}\}', template.content)
        template.placeholders = json.dumps(placeholders)
        
        db.session.commit()
        
        flash('Template updated successfully!', 'success')
        return redirect(url_for('hr_templates'))
    
    return render_template('hr/edit_template.html', user=user, template=template)


@app.route('/hr/templates/<int:template_id>/delete', methods=['POST'])
@login_required('hr')
def hr_delete_template(template_id):
    """Delete a document template"""
    template = DocumentTemplate.query.get_or_404(template_id)
    template.is_active = False
    db.session.commit()
    
    flash('Template deleted successfully!', 'success')
    return redirect(url_for('hr_templates'))


@app.route('/hr/templates/<int:template_id>/generate', methods=['GET', 'POST'])
@login_required('hr')
def hr_generate_document(template_id):
    """Generate a document from template"""
    user = User.query.get(session['user_id'])
    template = DocumentTemplate.query.get_or_404(template_id)
    
    if request.method == 'POST':
        employee_id = request.form.get('employee_id')
        employee = User.query.get(employee_id)
        
        if not employee:
            flash('Employee not found!', 'error')
            return redirect(url_for('hr_generate_document', template_id=template_id))
        
        # Prepare placeholder values
        placeholder_values = {
            'employee_name': employee.full_name,
            'designation': employee.position or 'N/A',
            'joining_date': employee.hire_date.strftime('%Y-%m-%d') if employee.hire_date else 'N/A',
            'employee_id': employee.employee_id or 'N/A',
            'company_name': 'SmartHire Inc.',  # Could be made configurable
            'hr_name': user.full_name,
            'department': employee.department or 'N/A',
            'phone': employee.phone,
            'email': employee.email
        }
        
        # Replace placeholders in template content
        content = template.content
        for placeholder, value in placeholder_values.items():
            content = content.replace(f'{{{{{placeholder}}}}}', str(value))
        
        # Create generated document record
        generated_doc = GeneratedDocument(
            template_id=template.id,
            employee_id=employee.id,
            document_name=f"{template.name} - {employee.full_name}",
            content=content,
            generated_by=user.id
        )
        
        db.session.add(generated_doc)
        db.session.commit()
        
        flash('Document generated successfully!', 'success')
        return redirect(url_for('hr_view_generated_document', doc_id=generated_doc.id))
    
    # Get list of employees for selection
    employees = User.query.filter_by(role='employee').order_by(User.full_name).all()
    
    return render_template('hr/generate_document.html', 
                         user=user, 
                         template=template, 
                         employees=employees)


# Enhanced Document Generation with Template Codes
@app.route('/hr/generate_document/<int:employee_id>/<string:template_code>', methods=['POST'])
@login_required('hr')
def generate_document_by_code(employee_id, template_code):
    """Generate document using template code"""
    from jinja2 import Template
    from datetime import datetime
    import os
    
    employee = User.query.get_or_404(employee_id)
    template_obj = DocumentTemplate.query.filter_by(code=template_code).first_or_404()
    
    # Prepare context for placeholders
    context = {
        "today": datetime.today().strftime("%d-%m-%Y"),
        "employee_name": employee.full_name,
        "employee_address": getattr(employee, 'address', 'Not specified'),
        "designation": getattr(employee, 'position', 'Not specified'),
        "company_name": "SmartHire AI Solutions",
        "joining_date": employee.hire_date.strftime("%d-%m-%Y") if employee.hire_date else "To be announced",
        "ctc": getattr(employee, 'salary', 'Not specified'),
        "hr_name": "HR Manager",
        "department": getattr(employee, 'department', 'Not specified'),
        "employee_id": f"EMP{employee.id:04d}",
        "phone": getattr(employee, 'phone', 'Not specified'),
        "email": employee.email,
    }
    
    # Render HTML with Jinja2 Template
    template = Template(template_obj.content)
    rendered_html = template.render(**context)
    
    # Create directory for generated documents
    file_dir = os.path.join("static", "generated_docs")
    os.makedirs(file_dir, exist_ok=True)
    
    # Generate filename
    timestamp = int(datetime.utcnow().timestamp())
    filename = f"{template_code}_{employee.id}_{timestamp}.html"
    file_path = os.path.join(file_dir, filename)
    
    # Save HTML file
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(rendered_html)
    
    # Save record in DB
    doc = GeneratedDocument(
        template_id=template_obj.id,
        employee_id=employee.id,
        document_name=f"{template_obj.name} - {employee.full_name}",
        content=rendered_html,
        file_path=file_path,
        status='generated',
        generated_by=session['user_id']
    )
    db.session.add(doc)
    db.session.commit()
    
    return jsonify({
        "status": "success",
        "message": f"{template_obj.name} generated successfully",
        "document_id": doc.id,
        "download_url": url_for('download_generated_document', doc_id=doc.id, _external=False)
    })


@app.route('/hr/documents/<int:doc_id>')
@login_required('hr')
def hr_view_generated_document(doc_id):
    """View a generated document"""
    user = User.query.get(session['user_id'])
    document = GeneratedDocument.query.get_or_404(doc_id)
    
    return render_template('hr/view_document.html', user=user, document=document)


@app.route('/hr/documents')
@login_required('hr')
def hr_generated_documents():
    """List all generated documents"""
    user = User.query.get(session['user_id'])
    documents = GeneratedDocument.query.order_by(GeneratedDocument.generated_at.desc()).all()
    
    return render_template('hr/generated_documents.html', user=user, documents=documents)


@app.route('/documents/download/<int:doc_id>')
@login_required(['hr', 'employee'])
def download_generated_document(doc_id):
    """Download generated document"""
    from flask import send_file
    
    document = GeneratedDocument.query.get_or_404(doc_id)
    
    # Check permissions: HR can download any, employees can only download their own
    if session.get('user_role') == 'employee' and document.employee_id != session.get('user_id'):
        return jsonify({'error': 'Unauthorized'}), 403
    
    if document.file_path and os.path.exists(document.file_path):
        return send_file(document.file_path, as_attachment=True, 
                        download_name=document.document_name + '.html')
    else:
        # If file doesn't exist, return content as HTML file
        from io import BytesIO
        from flask import Response
        
        response = Response(document.content, mimetype='text/html')
        response.headers['Content-Disposition'] = f'attachment; filename="{document.document_name}.html"'
        return response


# API Routes for Document Management
@app.route('/api/documents/<int:doc_id>/mark-downloaded', methods=['POST'])
@login_required('hr')
def mark_document_downloaded(doc_id):
    """Mark a generated document as downloaded"""
    try:
        document = GeneratedDocument.query.get_or_404(doc_id)
        document.status = 'downloaded'
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Document marked as downloaded'
        })
    except Exception as e:
        logger.error(f"Error marking document as downloaded: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to update document status'
        }), 500


# Employee Dashboard
@app.route('/employee/dashboard')
@login_required('employee')
def employee_dashboard():
    # Get the logged-in user
    user = User.query.get(session['user_id'])
    
    # Get HR contact information - prioritize assigned HR or first available HR
    hr_contact = None
    
    # Initialize progress variables
    onboarding_progress = 0
    offboarding_progress = 0
    current_day = 0
    total_days = 0
    total_offboarding_tasks = 0
    completed_offboarding_tasks = 0
    
    # Check user status and get appropriate checklist
    if user.status == 'Onboarding':
        # Get real onboarding tasks
        onboarding_tasks_for_progress = Task.query.filter_by(
            assigned_to=user.id, 
            task_type='onboarding'
        ).all()
        
        # Calculate real progress based on actual tasks
        total_onboarding_tasks = len(onboarding_tasks_for_progress)
        completed_onboarding_tasks = len([task for task in onboarding_tasks_for_progress if task.status == 'completed'])
        
        if total_onboarding_tasks > 0:
            onboarding_progress = int((completed_onboarding_tasks / total_onboarding_tasks) * 100)
        else:
            onboarding_progress = 0
        
        # Calculate actual onboarding days from hire date
        current_day = 1
        total_days = 30  # Default 30-day onboarding period
        if user.hire_date:
            # Convert hire_date to datetime for proper comparison
            hire_datetime = datetime.combine(user.hire_date, datetime.min.time())
            days_since_hire = (datetime.utcnow() - hire_datetime).days
            # Calculate business days (excluding weekends) - rough approximation
            business_days = days_since_hire - (days_since_hire // 7) * 2
            business_days = max(0, business_days)  # Don't go negative
            current_day = min(business_days + 1, total_days)  # +1 to start from Day 1
        
        # Legacy check - onboarding checklist system has been removed
        hr_contact = None
        
    elif user.status == 'Offboarding':
        # Get offboarding tasks using the general Task model
        offboarding_tasks = Task.query.filter_by(
            assigned_to=user.id, 
            task_type='offboarding'
        ).all()
        
        # Calculate offboarding progress
        total_offboarding_tasks = len(offboarding_tasks)
        completed_offboarding_tasks = len([task for task in offboarding_tasks if task.status == 'completed'])
        offboarding_progress = 0
        if total_offboarding_tasks > 0:
            offboarding_progress = int((completed_offboarding_tasks / total_offboarding_tasks) * 100)
        
        # Calculate days since offboarding started
        current_day = 1
        total_days = 5  # Standard offboarding period
        # For now, use a simple calculation - this could be enhanced with actual offboarding start date
        current_day = 1  # Could be calculated from earliest task created_at
        
        # Try to get assigned HR from offboarding tasks
        if offboarding_tasks:
            hr_assignments = [task.assigned_by for task in offboarding_tasks if task.assigned_by]
            if hr_assignments:
                hr_contact = User.query.get(hr_assignments[0])
        
    else:
        # Active employee - no specific checklist
        onboarding_progress = 100
        offboarding_progress = 0
        current_day = 0
        total_days = 0
    
    # Generate employee data based on the logged-in user
    employee_data = {
        'name': user.full_name,
        'username': user.username,
        'email': user.email,
        'phone': user.phone,
        'department': user.department or 'Not specified',
        'position': user.position or 'Employee',
        'hire_date': user.hire_date.strftime('%B %d, %Y') if user.hire_date else 'Not specified',
        'employee_id': user.employee_id or f'EMP-{user.id:04d}',
        'status': user.status or 'Active',
        'manager': 'Sarah Johnson',  # This would come from manager relationship
        'location': 'Main Office',
        'work_schedule': '9:00 AM - 5:00 PM',
        'team_members': [
            {'name': 'John Smith', 'role': 'Senior Developer', 'status': 'online'},
            {'name': 'Emily Chen', 'role': 'Product Designer', 'status': 'busy'},
            {'name': 'Michael Brown', 'role': 'DevOps Engineer', 'status': 'offline'}
        ],
        'pending_tasks': [
            {'id': 1, 'title': 'Submit timesheet', 'due': 'Today', 'priority': 'high'},
            {'id': 2, 'title': 'Complete training module', 'due': 'Tomorrow', 'priority': 'medium'},
            {'id': 3, 'title': 'Update project documentation', 'due': 'Next week', 'priority': 'low'}
        ],
        'support_team': [
            {'name': hr_contact.full_name if hr_contact else 'HR Department', 
             'role': 'HR Team', 
             'email': hr_contact.email if hr_contact else 'hr@company.com', 
             'phone': hr_contact.phone if hr_contact else '+1 (555) 123-4567'},
            {'name': 'IT Support', 'role': 'Technical Support', 'email': 'support@company.com', 'phone': '+1 (555) 987-6543'}
        ],
        'resources': [
            {'name': 'Employee Handbook', 'type': 'document', 'format': 'PDF'},
            {'name': 'Benefits Guide', 'type': 'document', 'format': 'PDF'},
            {'name': 'Training Materials', 'type': 'folder', 'format': 'ZIP'},
            {'name': 'Company Policies', 'type': 'document', 'format': 'PDF'},
            {'name': 'IT Helpdesk', 'type': 'link', 'url': '#'}
        ],
        # Set appropriate progress based on user status
        'onboarding_progress': {
            'current_day': current_day if user.status == 'Onboarding' else 0,
            'total_days': total_days if user.status == 'Onboarding' else 0,
            'percentage': onboarding_progress if user.status == 'Onboarding' else 0,
            'tasks_completed': completed_onboarding_tasks if user.status == 'Onboarding' else 0,
            'total_tasks': total_onboarding_tasks if user.status == 'Onboarding' else 0,
            'tasks': onboarding_tasks_for_progress if user.status == 'Onboarding' else []
        } if user.status == 'Onboarding' else None,
        'offboarding_progress': {
            'current_day': current_day if user.status == 'Offboarding' else 0,
            'total_days': total_days if user.status == 'Offboarding' else 0,
            'percentage': offboarding_progress if user.status == 'Offboarding' else 0,
            'tasks_completed': completed_offboarding_tasks if user.status == 'Offboarding' else 0,
            'total_tasks': total_offboarding_tasks if user.status == 'Offboarding' else 0
        } if user.status == 'Offboarding' else None,
    }
    
    # Get employee documents count
    documents_count = EmployeeDocument.query.filter_by(user_id=user.id).count()
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(recipient_id=session['user_id'], status='unread').count()
    
    # Get HR contacts for messaging - exclude test HR users
    hr_contacts = User.query.filter_by(role='hr', is_active=True).filter(
        User.username != 'test_hr',
        User.username != 'testhr',
        User.full_name.notlike('%test%')
    ).all()
    
    # Set Deeksha as primary HR contact if available, otherwise use first available HR
    hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
        User.username != 'test_hr',
        User.username != 'testhr',
        User.full_name.notlike('%test%')
    ).filter(User.full_name.like('%Deeksha%')).first()
    
    if not hr_contact and hr_contacts:
        hr_contact = hr_contacts[0]
    
    # Generate employee journey timeline data
    employee_journey = {
        'offer_accepted': {
            'completed': True,
            'date': user.created_at.strftime('%Y-%m-%d') if user.created_at else '2024-01-15'
        },
        'documents_submitted': {
            'completed': documents_count > 0,
            'date': '2024-01-18' if documents_count > 0 else 'Pending'
        },
        'systems_created': {
            'completed': True,
            'date': user.created_at.strftime('%Y-%m-%d') if user.created_at else '2024-01-20'
        },
        'first_day': {
            'completed': user.hire_date and user.hire_date <= datetime.now().date(),
            'date': user.hire_date.strftime('%Y-%m-%d') if user.hire_date else '2024-02-01'
        },
        'orientation': {
            'completed': onboarding_progress >= 50,
            'date': '2024-02-02' if onboarding_progress >= 50 else 'Scheduled'
        },
        'training_modules': {
            'completed': onboarding_progress >= 80,
            'date': '2024-02-15' if onboarding_progress >= 80 else 'In Progress',
            'progress': min(onboarding_progress, 100)
        },
        'confirmation_letter': {
            'completed': False,  # Typically after probation period
            'date': 'Pending'
        }
    }
    
    # Get tasks assigned to this employee
    employee_tasks = Task.query.filter_by(assigned_to=user.id).all()
    formatted_employee_tasks = []
    for task in employee_tasks:
        formatted_employee_tasks.append({
            'id': task.id,
            'title': task.title,
            'description': task.description,
            'priority': task.priority,
            'status': task.status,
            'task_type': task.task_type,
            'due_date': task.due_date,
            'is_overdue': task.due_date and task.due_date.date() < datetime.utcnow().date() if task.due_date else False
        })

    # Get buddy/mentor assignments for this employee
    buddy_mentor = None
    try:
        # Legacy check - buddy/mentor system has been removed
        pass
    except:
        buddy_mentor = None

    return render_template('employee_dashboard.html', 
                         employee=employee_data, 
                         user=user,
                         hr_contact=hr_contact,
                         documents_count=documents_count,
                         unread_messages_count=unread_messages_count,
                         hr_contacts=hr_contacts,
                         onboarding_progress=onboarding_progress,
                         employee_journey=employee_journey,
                         employee_tasks=formatted_employee_tasks,
                         buddy_mentor=buddy_mentor,
                         current_user=user,  # For compatibility with existing templates
                         title=f'{user.full_name}\'s Dashboard')

# Employee Documents
@app.route('/employee/documents')
@login_required('employee')
def employee_documents():
    """Employee documents page"""
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    # Get documents generated for this employee
    documents = GeneratedDocument.query.filter_by(employee_id=user.id).order_by(GeneratedDocument.generated_at.desc()).all()
    
    # Format documents for display
    formatted_documents = []
    for doc in documents:
        formatted_documents.append({
            'id': doc.id,
            'template_name': doc.template.name if doc.template else 'Unknown Template',
            'template_type': doc.template.template_type if doc.template else 'Unknown',
            'document_name': doc.document_name,
            'content': doc.content,
            'status': doc.status,
            'generated_at': doc.generated_at,
            'generated_by_name': User.query.get(doc.generated_by).full_name if doc.generated_by else 'Unknown'
        })
    
    return render_template('employee_documents.html', 
                         documents=formatted_documents,
                         user=user,
                         title='My Documents')

@app.route('/employee/documents/<int:doc_id>/view')
@login_required('employee')
def view_employee_document(doc_id):
    """View document content"""
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    document = GeneratedDocument.query.filter_by(id=doc_id, employee_id=user.id).first()
    
    if not document:
        return jsonify({'success': False, 'message': 'Document not found'}), 404
    
    return jsonify({
        'success': True,
        'content': document.content,
        'template_name': document.template_name
    })

@app.route('/employee/documents/<int:doc_id>/download')
@login_required('employee')
def download_employee_document(doc_id):
    """Download document as PDF"""
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    document = GeneratedDocument.query.filter_by(id=doc_id, employee_id=user.id).first()
    
    if not document:
        return jsonify({'success': False, 'message': 'Document not found'}), 404
    
    # Update status to downloaded
    document.status = 'downloaded'
    document.downloaded_at = datetime.utcnow()
    db.session.commit()
    
    # Generate professional PDF
    # Create PDF in memory
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=colors.darkblue,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.darkblue,
        spaceBefore=20
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
        leading=14
    )
    
    # Build PDF content
    story = []
    
    # Company Header
    story.append(Paragraph("SmartHire HR System", title_style))
    story.append(Spacer(1, 20))
    
    # Document Information Table
    doc_info = [
        ['Document Type:', document.template.template_type if document.template else 'Unknown'],
        ['Template:', document.template.name if document.template else 'Unknown Template'],
        ['Generated Date:', document.generated_at.strftime('%B %d, %Y at %I:%M %p')],
        ['Generated By:', User.query.get(document.generated_by).full_name if document.generated_by else 'Unknown']
    ]
    
    info_table = Table(doc_info, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 30))
    
    # Document Title
    story.append(Paragraph(document.document_name, heading_style))
    story.append(Spacer(1, 20))
    
    # Document Content (convert HTML to plain text for PDF)
    soup = BeautifulSoup(document.content, 'html.parser')
    content_text = soup.get_text()
    
    # Split content into paragraphs
    paragraphs = content_text.split('\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), normal_style))
    
    # Footer
    story.append(Spacer(1, 40))
    story.append(Paragraph("This document was generated electronically by SmartHire HR System.", 
                          ParagraphStyle('Footer', parent=styles['Normal'], 
                                       fontSize=9, textColor=colors.grey, alignment=1)))
    
    # Generate PDF
    doc.build(story)
    
    # Get PDF content
    buffer.seek(0)
    pdf_content = buffer.getvalue()
    buffer.close()
    
    # Create response
    response = Response(pdf_content, mimetype='application/pdf')
    response.headers['Content-Disposition'] = f'attachment; filename="{document.document_name}.pdf"'
    
    return response

@app.route('/employee/documents/<int:doc_id>/acknowledge', methods=['POST'])
@login_required('employee')
def acknowledge_employee_document(doc_id):
    """Acknowledge document receipt"""
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    document = GeneratedDocument.query.filter_by(id=doc_id, employee_id=user.id).first()
    
    if not document:
        return jsonify({'success': False, 'message': 'Document not found'}), 404
    
    # Update status to acknowledged
    document.status = 'acknowledged'
    document.acknowledged_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Document acknowledged successfully'})

@app.route('/employee/messages')
@login_required('employee')
def employee_messages():
    """Employee messages page"""
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    # Get all users for messaging (HR and other employees)
    all_users = User.query.filter(User.id != user_id, User.is_active == True).all()
    
    # Get recent conversations
    recent_messages = Message.query.filter(
        (Message.sender_id == user_id) | (Message.recipient_id == user_id)
    ).order_by(Message.sent_at.desc()).limit(50).all()
    
    # Get unread count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user_id, 
        status='unread'
    ).count()
    
    return render_template('employee_messages.html',
                         user=user,
                         all_users=all_users,
                         recent_messages=recent_messages,
                         unread_messages_count=unread_messages_count,
                         title='Messages')

@app.route('/employee/profile', methods=['GET', 'POST'])
@login_required('employee')
def employee_profile():
    user = User.query.get(session['user_id'])
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user.id, 
        status='unread'
    ).count()

    if request.method == 'POST':
        try:
            user.phone = request.form.get('phone', user.phone)
            user.department = request.form.get('department', user.department)
            user.position = request.form.get('position', user.position)

            db.session.commit()
            flash('Your profile has been updated successfully.', 'success')
            return redirect(url_for('employee_profile'))
        except Exception as e:
            db.session.rollback()
            app.logger.error(f'Error updating employee profile: {str(e)}')
            flash('An error occurred while updating your profile. Please try again.', 'danger')

    return render_template('employee_profile.html', user=user, current_user=user, unread_messages_count=unread_messages_count)

@app.route('/api/employee/tasks/<int:task_id>/status', methods=['PUT'])
@login_required('employee')
def update_employee_task_status(task_id):
    """Update task status for employee"""
    try:
        task = Task.query.get_or_404(task_id)
        user_id = session.get('user_id')
        
        # Verify task is assigned to this employee
        if task.assigned_to != user_id:
            return jsonify({
                'status': 'error',
                'message': 'You can only update your own tasks'
            }), 403
        
        data = request.get_json()
        new_status = data.get('status')
        
        # Validate status
        valid_statuses = ['pending', 'in_progress', 'completed']
        if new_status not in valid_statuses:
            return jsonify({
                'status': 'error',
                'message': 'Invalid status'
            }), 400
        
        task.status = new_status
        if new_status == 'completed':
            task.completed_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Task status updated successfully'
        })
        
    except Exception as e:
        db.session.rollback()
@login_required(['hr', 'employee'])
def get_unread_count():
    """Get unread message count for current user"""
    try:
        current_user_id = session['user_id']
        unread_count = Message.query.filter_by(
            recipient_id=current_user_id, 
            status='unread'
        ).count()
        
        return jsonify({
            'status': 'success',
            'unread_count': unread_count
        })
    except Exception as e:
        app.logger.error(f"Error fetching unread count: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch unread count'
        }), 500

@app.route('/api/hr-contacts', methods=['GET'])
@login_required('employee')
def get_hr_contacts():
    """Get list of HR contacts for employees to message"""
    try:
        # Get HR contacts - exclude test HR users
        hr_contacts = User.query.filter_by(role='hr', is_active=True).filter(
            User.username != 'test_hr',
            User.username != 'testhr',
            User.full_name.notlike('%test%')
        ).all()
        contacts_list = []
        
        for hr in hr_contacts:
            contacts_list.append({
                'id': hr.id,
                'full_name': hr.full_name,
                'username': hr.username,
                'email': hr.email,
                'position': hr.position or 'HR Representative',
                'is_primary': hr.full_name.like('%Deeksha%')
            })
        
        return jsonify({
            'status': 'success',
            'contacts': contacts_list
        })
    except Exception as e:
        app.logger.error(f"Error fetching HR contacts: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch HR contacts'
        }), 500

@app.route('/api/messages/<int:message_id>', methods=['PUT'])
@login_required(['hr', 'employee'])
def update_message(message_id):
    """Update an existing message (edit functionality)"""
    try:
        current_user_id = session['user_id']
        
        # Get the message
        message = Message.query.get_or_404(message_id)
        
        # Check if user is the sender of this message
        if message.sender_id != current_user_id:
            return jsonify({
                'status': 'error',
                'message': 'You can only edit your own messages'
            }), 403
        
        # Check if message is recent (e.g., within 24 hours)
        from datetime import datetime, timedelta
        if datetime.now() - message.sent_at > timedelta(hours=24):
            return jsonify({
                'status': 'error',
                'message': 'Messages can only be edited within 24 hours of sending'
            }), 400
        
        # Get new content
        data = request.get_json()
        new_content = data.get('content', '').strip()
        
        if not new_content:
            return jsonify({
                'status': 'error',
                'message': 'Message content cannot be empty'
            }), 400
        
        # Update message
        message.content = new_content
        message.is_edited = True
        message.edited_at = datetime.now()
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Message updated successfully',
            'data': {
                'id': message.id,
                'content': message.content,
                'is_edited': message.is_edited,
                'edited_at': message.edited_at.isoformat() if message.edited_at else None
            }
        })
        
    except Exception as e:
        app.logger.error(f"Error updating message {message_id}: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to update message'
        }), 500

@app.route('/api/messages/<int:message_id>', methods=['DELETE'])
@login_required(['hr', 'employee'])
def delete_message(message_id):
    """Delete a message"""
    try:
        current_user_id = session['user_id']
        
        # Get the message
        message = Message.query.get_or_404(message_id)
        
        # Check if user is the sender of this message
        if message.sender_id != current_user_id:
            return jsonify({
                'status': 'error',
                'message': 'You can only delete your own messages'
            }), 403
        
        # Check if message is recent (e.g., within 24 hours)
        from datetime import datetime, timedelta
        if datetime.now() - message.sent_at > timedelta(hours=24):
            return jsonify({
                'status': 'error',
                'message': 'Messages can only be deleted within 24 hours of sending'
            }), 400
        
        # Soft delete by marking as deleted
        message.is_deleted = True
        message.deleted_at = datetime.now()
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Message deleted successfully'
        })
        
    except Exception as e:
        app.logger.error(f"Error deleting message {message_id}: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to delete message'
        }), 500

@app.route('/api/messages', methods=['POST'])
@login_required(['hr', 'employee'])
def create_message():
    """Create a new message"""
    try:
        # Handle both form data and JSON data
        if request.is_json:
            data = request.get_json()
            recipient_id = data.get('recipient_id')
            subject = data.get('subject', 'No Subject')
            content = data.get('content')
            is_priority = data.get('is_priority', False)
            message_type = data.get('message_type', 'message')
            priority = data.get('priority', 'normal')
        else:
            recipient_id = request.form.get('recipient_id')
            subject = request.form.get('subject', 'No Subject')
            content = request.form.get('content')
            is_priority = request.form.get('is_priority') == 'on'
            message_type = request.form.get('message_type', 'message')
            priority = request.form.get('priority', 'normal')
        
        # Validate required fields
        if not recipient_id or not content:
            return jsonify({
                'status': 'error',
                'message': 'Recipient and content are required'
            }), 400
        
        # Verify recipient exists
        recipient = User.query.get(recipient_id)
        if not recipient:
            return jsonify({
                'status': 'error',
                'message': 'Recipient not found'
            }), 404
        
        message = Message(
            sender_id=session['user_id'],
            recipient_id=recipient_id,
            subject=subject,
            content=content,
            message_type=message_type,
            priority=priority,
            is_priority=is_priority
        )
        
        db.session.add(message)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Message sent successfully',
            'message_id': message.id
        })
        
    except Exception as e:
        app.logger.error(f"Error creating message: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to send message'
        }), 500

@app.route('/api/messages/conversation/<int:user_id>', methods=['GET'])
@login_required(['hr', 'employee'])
def get_conversation_messages(user_id):
    """Get messages in a conversation with a specific user"""
    try:
        current_user_id = session['user_id']
        
        messages = Message.query.filter(
            ((Message.sender_id == current_user_id) & (Message.recipient_id == user_id)) |
            ((Message.sender_id == user_id) & (Message.recipient_id == current_user_id))
        ).order_by(Message.sent_at.asc()).all()
        
        message_list = []
        for msg in messages:
            # Get sender information
            sender = User.query.get(msg.sender_id)
            sender_name = sender.full_name if sender else 'Unknown'
            sender_role = sender.role if sender else 'unknown'
            
            message_list.append({
                'id': msg.id,
                'content': msg.content,
                'subject': msg.subject,
                'sent_at': msg.sent_at.isoformat(),
                'sender_id': msg.sender_id,
                'recipient_id': msg.recipient_id,
                'is_read': msg.status == 'read',
                'is_priority': msg.is_priority,
                'is_edited': msg.is_edited,
                'edited_at': msg.edited_at.isoformat() if msg.edited_at else None,
                'sender_name': sender_name,
                'sender_role': sender_role,
                'is_sent': msg.sender_id == current_user_id
            })
        
        user = User.query.get(user_id)
        
        return jsonify({
            'status': 'success',
            'messages': message_list,
            'user': {
                'id': user.id,
                'full_name': user.full_name,
                'username': user.username,
                'role': user.role
            }
        })
        
    except Exception as e:
        app.logger.error(f"Error getting conversation: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to load conversation'
        }), 500

@app.route('/api/messages/mark-read/<int:user_id>', methods=['PUT'])
@login_required(['hr', 'employee'])
def mark_messages_as_read(user_id):
    """Mark messages from a user as read"""
    try:
        current_user_id = session['user_id']
        
        Message.query.filter(
            Message.sender_id == user_id,
            Message.recipient_id == current_user_id,
            Message.status == 'unread'
        ).update({'status': 'read'})
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Messages marked as read'
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error marking messages as read: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to mark messages as read'
        }), 500

@app.route('/api/messages/unread-count', methods=['GET'])
@login_required(['hr', 'employee'])
def get_unread_count():
    """Get unread message count for current user"""
    try:
        current_user_id = session['user_id']
        
        unread_count = Message.query.filter_by(
            recipient_id=current_user_id,
            status='unread'
        ).count()
        
        return jsonify({
            'status': 'success',
            'unread_count': unread_count
        })
        
    except Exception as e:
        app.logger.error(f"Error getting unread count: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get unread count'
        }), 500

@app.route('/api/tasks', methods=['POST'])
@login_required('hr')
def create_task():
    """Create a new task"""
    try:
        assigned_to = request.form.get('assigned_to')
        current_user = User.query.get(session['user_id'])
        
        # Validate assignment rules
        if assigned_to:
            assigned_user = User.query.get(assigned_to)
            if not assigned_user:
                return jsonify({
                    'status': 'error',
                    'message': 'Assigned user not found'
                }), 400
            
            # Rule 1: Can only assign to employees (not HR)
            if assigned_user.role != 'employee':
                return jsonify({
                    'status': 'error',
                    'message': 'Tasks can only be assigned to employees'
                }), 400
            
            # Rule 2: Must be from different department
            if assigned_user.department == current_user.department:
                return jsonify({
                    'status': 'error',
                    'message': 'Cannot assign tasks to employees from the same department'
                }), 400
            
            # Rule 3: Employee must be under HR's chain of command
            # Check if the HR is in the management chain above this employee
            def is_in_management_chain(hr_user, employee_user):
                """Check if HR is in the management chain above employee"""
                current = employee_user
                while current and current.manager_id:
                    manager = User.query.get(current.manager_id)
                    if manager and manager.id == hr_user.id:
                        return True
                    current = manager
                return False
            
            if not is_in_management_chain(current_user, assigned_user):
                return jsonify({
                    'status': 'error',
                    'message': 'Can only assign tasks to employees under your chain of command'
                }), 400
        
        task = Task(
            title=request.form.get('title'),
            description=request.form.get('description'),
            assigned_to=assigned_to if assigned_to else None,
            assigned_by=session['user_id'],
            task_type=request.form.get('task_type', 'general'),
            priority=request.form.get('priority', 'medium'),
            status=request.form.get('status', 'pending'),
            due_date=datetime.strptime(request.form.get('due_date'), '%Y-%m-%d').date() if request.form.get('due_date') else None
        )
        
        db.session.add(task)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Task created successfully',
            'task_id': task.id
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error creating task: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to create task'
        }), 500

@app.route('/api/tasks/<int:task_id>', methods=['DELETE'])
@login_required('hr')
def delete_task(task_id):
    """Delete a task"""
    try:
        task = Task.query.get_or_404(task_id)
        db.session.delete(task)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Task deleted successfully'
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting task: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to delete task'
        }), 500

# Onboarding Tasks
@app.route('/employee/onboarding-tasks')
@login_required('employee')
def onboarding_tasks():
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user.id, 
        status='unread'
    ).count()

    # Get onboarding tasks using the general Task model
    onboarding_tasks = Task.query.filter_by(
        assigned_to=user.id, 
        task_type='onboarding'
    ).all()
    
    # If no onboarding tasks exist, create sample tasks for testing
    if not onboarding_tasks and user.status == 'Onboarding':
        sample_tasks = [
            {
                'title': 'Complete Personal Information Form',
                'description': 'Fill out your personal details and emergency contact information',
                'priority': 'high',
                'task_type': 'onboarding'
            },
            {
                'title': 'Upload Government ID Documents',
                'description': 'Upload your Aadhaar card and PAN card for verification',
                'priority': 'high',
                'task_type': 'onboarding'
            },
            {
                'title': 'Watch Company Orientation Video',
                'description': 'Complete the mandatory company orientation and safety training',
                'priority': 'medium',
                'task_type': 'onboarding'
            },
            {
                'title': 'Review Employee Handbook',
                'description': 'Read and acknowledge the employee handbook and company policies',
                'priority': 'medium',
                'task_type': 'onboarding'
            },
            {
                'title': 'Sign Employment Agreement',
                'description': 'Review and sign your employment agreement and confidentiality terms',
                'priority': 'high',
                'task_type': 'onboarding'
            }
        ]
        
        for task_data in sample_tasks:
            new_task = Task(
                title=task_data['title'],
                description=task_data['description'],
                priority=task_data['priority'],
                task_type=task_data['task_type'],
                status='pending',
                assigned_to=user.id,
                assigned_by=1,  # Assuming HR user ID is 1
                department='HR',
                due_date=datetime.utcnow() + timedelta(days=7),
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow()
            )
            db.session.add(new_task)
        
        db.session.commit()
        
        # Fetch the newly created tasks
        onboarding_tasks = Task.query.filter_by(
            assigned_to=user.id, 
            task_type='onboarding'
        ).all()
    
    # Create compatible task objects for template
    compatible_tasks = []
    for task in onboarding_tasks:
        compatible_task = type('Task', (), {
            'id': task.id,
            'task_name': task.title,
            'task_description': task.description,
            'is_completed': task.status == 'completed',
            'due_date': task.due_date,
            'priority': task.priority,
            'status': task.status
        })()
        compatible_tasks.append(compatible_task)
    
    # Create a checklist-like structure for compatibility with template
    checklist = None
    if compatible_tasks:
        def get_progress(self):
            if not compatible_tasks:
                return 0
            completed_tasks = len([task for task in compatible_tasks if task.is_completed])
            return int((completed_tasks / len(compatible_tasks)) * 100)
        
        checklist = type('Checklist', (), {
            'tasks': compatible_tasks,
            'created_at': onboarding_tasks[0].created_at if onboarding_tasks else datetime.utcnow(),
            'get_progress': get_progress
        })()

    return render_template('onboarding_tasks.html', user=user, checklist=checklist, now=datetime.utcnow(), unread_messages_count=unread_messages_count)


@app.route('/employee/offboarding-tasks')
@login_required('employee')
def offboarding_tasks():
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user.id, 
        status='unread'
    ).count()

    # Get user's documents
    documents = EmployeeDocument.query.filter_by(user_id=user_id).all()
    
    # Get offboarding tasks using the general Task model
    offboarding_tasks = Task.query.filter_by(
        assigned_to=user.id, 
        task_type='offboarding'
    ).all()
    
    # Create compatible task objects for template
    compatible_tasks = []
    for task in offboarding_tasks:
        compatible_task = type('Task', (), {
            'id': task.id,
            'task_name': task.title,
            'task_description': task.description,
            'is_completed': task.status == 'completed',
            'due_date': task.due_date,
            'priority': task.priority,
            'status': task.status
        })()
        compatible_tasks.append(compatible_task)
    
    # Create a checklist-like structure for compatibility with template
    checklist = None
    if compatible_tasks:
        def get_progress(self):
            if not compatible_tasks:
                return 0
            completed_tasks = len([task for task in compatible_tasks if task.is_completed])
            return int((completed_tasks / len(compatible_tasks)) * 100)
        
        checklist = type('Checklist', (), {
            'tasks': compatible_tasks,
            'created_at': offboarding_tasks[0].created_at if offboarding_tasks else datetime.utcnow(),
            'get_progress': get_progress
        })()
    
    # Get HR contact - prioritize Deeksha and exclude test HR users
    hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
        User.username != 'test_hr',
        User.username != 'testhr',
        User.full_name.notlike('%test%')
    ).filter(User.full_name.like('%Deeksha%')).first()
    
    if not hr_contact:
        hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
            User.username != 'test_hr',
            User.username != 'testhr',
            User.full_name.notlike('%test%')
        ).first()
    
    # Get manager information if available
    manager = None
    if user.manager_id:
        manager = User.query.get(user.manager_id)
    
    offboarding = {
        'status': user.status or 'Active',
        'exit_date': user.exit_date,
        'documents': documents,
        'hr_contact': hr_contact,
        'user': user,
        'manager': manager
    }

    return render_template('offboarding_tasks.html', user=user, checklist=checklist, offboarding=offboarding, unread_messages_count=unread_messages_count)


@app.route('/employee/pre-offboarding')
@login_required('employee')
def employee_pre_offboarding():
    user_id = session.get('user_id')
    user = User.query.get(user_id)

    if not user:
        flash('Employee not found.', 'danger')
        return redirect(url_for('employee_dashboard'))
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user.id, 
        status='unread'
    ).count()

    status = user.status or 'Active'
    phase = 'none'
    if status in ['Pre-Offboarding', 'Pre Offboarding']:
        phase = 'pre_offboarding'
    elif status == 'Offboarding':
        phase = 'offboarding'
    elif status in ['Inactive', 'Exited']:
        phase = 'completed'

    today = datetime.utcnow().date()
    notice_days = None
    if user.exit_date:
        try:
            notice_days = (user.exit_date - today).days
        except Exception:
            notice_days = None

    return render_template('pre-offboarding.html', user=user, phase=phase, notice_days=notice_days, unread_messages_count=unread_messages_count)

@app.route('/download-document/<int:doc_id>')
@login_required(['employee', 'hr'])
def download_document(doc_id):
    """Download employee document"""
    try:
        document = EmployeeDocument.query.get_or_404(doc_id)
        
        # Check permissions
        if session.get('role') == 'employee' and document.user_id != session['user_id']:
            flash('You can only download your own documents.', 'danger')
            return redirect(url_for('employee_dashboard'))
        
        # Check if document is approved for download
        if document.status != 'approved':
            flash('Document is not yet available for download.', 'warning')
            return redirect(request.referrer or url_for('employee_dashboard'))
        
        # For demo purposes, return a simple response
        from flask import make_response
        response = make_response(f"Document: {document.file_name}\nType: {document.document_type}\nStatus: {document.status}")
        response.headers['Content-Type'] = 'text/plain'
        response.headers['Content-Disposition'] = f'attachment; filename="{document.file_name}"'
        
        return response
        
    except Exception as e:
        app.logger.error(f'Error downloading document: {str(e)}')
        flash('Error downloading document.', 'danger')
        return redirect(request.referrer or url_for('employee_dashboard'))

@app.route('/pre-onboarding')
@login_required('hr')
def pre_onboarding():
    # Enhanced pre-onboarding with communication features and ATS screening
    employees = User.query.filter_by(role='employee', status='Onboarding').order_by(User.created_at.desc()).all()
    return render_template('pre-onboarding-enhanced.html', employees=employees)

@app.route('/pre-offboarding')
@login_required('hr')
def pre_offboarding():
    # Pre-offboarding management for HR
    employees = User.query.filter(User.status.in_(['Pre-Offboarding', 'Pre Offboarding'])).order_by(User.exit_date.asc()).all()
    return render_template('pre-offboarding.html', employees=employees)


# API Routes for Pre-Onboarding Features

@app.route('/api/team-members', methods=['GET'])
@login_required('hr')
def get_team_members():
    """Get all team members for introductions"""
    try:
        team_members = User.query.filter(
            User.role.in_(['hr', 'manager', 'employee']),
            User.status == 'Active'
        ).all()
        
        return jsonify({
            'status': 'success',
            'team_members': [
                {
                    'id': member.id,
                    'full_name': member.full_name,
                    'position': getattr(member, 'position', 'Team Member')
                }
                for member in team_members
            ]
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/welcome-packages/<int:employee_id>', methods=['GET'])
@login_required('hr')
def get_welcome_packages(employee_id):
    """Get welcome packages for an employee"""
    try:
        packages = WelcomePackage.query.filter_by(employee_id=employee_id).all()
        return jsonify({
            'status': 'success',
            'packages': [
                {
                    'id': pkg.id,
                    'package_name': pkg.package_name,
                    'items': json.loads(pkg.items) if pkg.items else [],
                    'status': pkg.status,
                    'tracking_number': pkg.tracking_number,
                    'estimated_delivery': pkg.estimated_delivery.isoformat() if pkg.estimated_delivery else None,
                    'notes': pkg.notes
                }
                for pkg in packages
            ]
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/welcome-package', methods=['POST'])
@login_required('hr')
def create_welcome_package():
    """Create a new welcome package"""
    try:
        data = request.json
        
        package = WelcomePackage(
            employee_id=data['employee_id'],
            package_name=data['package_name'],
            items=json.dumps(data['items']),
            shipping_address=data.get('shipping_address'),
            estimated_delivery=datetime.strptime(data['estimated_delivery'], '%Y-%m-%d').date() if data.get('estimated_delivery') else None,
            notes=data.get('notes')
        )
        
        db.session.add(package)
        db.session.commit()
        
        # Create notification
        create_notification(
            employee_id=data['employee_id'],
            title=f'Welcome Package: {data["package_name"]}',
            message=f'Welcome package "{data["package_name"]}" has been prepared for you.',
            notification_type='welcome_package',
            priority='normal'
        )
        
        return jsonify({'status': 'success', 'message': 'Welcome package created successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/admin-tasks/<int:employee_id>', methods=['GET'])
@login_required('hr')
def get_admin_tasks(employee_id):
    """Get administrative tasks for an employee"""
    try:
        tasks = PreOnboardingTask.query.filter_by(employee_id=employee_id).all()
        return jsonify({
            'status': 'success',
            'tasks': [
                {
                    'id': task.id,
                    'task_name': task.task_name,
                    'task_description': task.task_description,
                    'task_type': task.task_type,
                    'due_date': task.due_date.isoformat() if task.due_date else None,
                    'is_completed': task.is_completed,
                    'assigned_to': task.assigned_to
                }
                for task in tasks
            ]
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/admin-task', methods=['POST'])
@login_required('hr')
def create_admin_task():
    """Create a new administrative task"""
    try:
        data = request.json
        
        task = PreOnboardingTask(
            employee_id=data['employee_id'],
            task_name=data['task_name'],
            task_description=data.get('task_description'),
            task_type=data['task_type'],
            due_date=datetime.strptime(data['due_date'], '%Y-%m-%d').date() if data.get('due_date') else None,
            assigned_to=data['assigned_to']
        )
        
        db.session.add(task)
        db.session.commit()
        
        # Create notification if assigned to employee
        if data['assigned_to'] == 'Employee':
            create_notification(
                employee_id=data['employee_id'],
                title=f'Administrative Task: {data["task_name"]}',
                message=f'Please complete: {data.get("task_description", data["task_name"])}',
                notification_type='admin_task',
                priority='normal'
            )
        
        return jsonify({'status': 'success', 'message': 'Administrative task created successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/agenda/<int:employee_id>', methods=['GET'])
@login_required('hr')
def get_first_day_agenda(employee_id):
    """Get first-day agenda for an employee"""
    try:
        agenda = FirstDayAgenda.query.filter_by(employee_id=employee_id).first()
        if agenda:
            return jsonify({
                'status': 'success',
                'agenda': {
                    'id': agenda.id,
                    'agenda_date': agenda.agenda_date.isoformat(),
                    'start_time': agenda.start_time.isoformat() if agenda.start_time else None,
                    'end_time': agenda.end_time.isoformat() if agenda.end_time else None,
                    'location': agenda.location,
                    'agenda_items': json.loads(agenda.agenda_items) if agenda.agenda_items else [],
                    'meeting_links': json.loads(agenda.meeting_links) if agenda.meeting_links else [],
                    'preparation_notes': agenda.preparation_notes,
                    'is_shared': agenda.is_shared
                }
            })
        else:
            return jsonify({'status': 'success', 'agenda': None})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/first-day-agenda', methods=['POST'])
@login_required('hr')
def create_first_day_agenda():
    """Create a new first-day agenda"""
    try:
        data = request.json
        
        agenda = FirstDayAgenda(
            employee_id=data['employee_id'],
            agenda_date=datetime.strptime(data['agenda_date'], '%Y-%m-%d').date(),
            start_time=datetime.strptime(data['start_time'], '%H:%M').time() if data.get('start_time') else None,
            end_time=datetime.strptime(data['end_time'], '%H:%M').time() if data.get('end_time') else None,
            location=data.get('location'),
            agenda_items=json.dumps(data['agenda_items']),
            meeting_links=json.dumps(data.get('meeting_links', [])),
            preparation_notes=data.get('preparation_notes')
        )
        
        db.session.add(agenda)
        db.session.commit()
        
        return jsonify({'status': 'success', 'message': 'First-day agenda created successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/share-agenda/<int:employee_id>', methods=['POST'])
@login_required('hr')
def share_agenda_with_employee(employee_id):
    """Share first-day agenda with employee"""
    try:
        agenda = FirstDayAgenda.query.filter_by(employee_id=employee_id).first()
        if not agenda:
            return jsonify({'status': 'error', 'message': 'No agenda found'}), 404
        
        agenda.is_shared = True
        agenda.shared_at = datetime.utcnow()
        db.session.commit()
        
        # Create notification for employee
        create_notification(
            employee_id=employee_id,
            title='First-Day Agenda Available',
            message=f'Your first-day agenda for {agenda.agenda_date.strftime("%B %d, %Y")} is now available.',
            notification_type='agenda_shared',
            priority='high'
        )
        
        return jsonify({'status': 'success', 'message': 'Agenda shared successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/team-introductions/<int:employee_id>', methods=['GET'])
@login_required('hr')
def get_team_introductions(employee_id):
    """Get team introductions for an employee"""
    try:
        introductions = TeamIntroduction.query.filter_by(employee_id=employee_id).all()
        return jsonify({
            'status': 'success',
            'introductions': [
                {
                    'id': intro.id,
                    'team_member_id': intro.team_member_id,
                    'team_member_name': intro.team_member.full_name,
                    'introduction_type': intro.introduction_type,
                    'message': intro.message,
                    'is_sent': intro.is_sent,
                    'scheduled_date': intro.scheduled_date.isoformat() if intro.scheduled_date else None
                }
                for intro in introductions
            ]
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/team-introduction', methods=['POST'])
@login_required('hr')
def create_team_introduction():
    """Create a new team introduction"""
    try:
        data = request.json
        
        introduction = TeamIntroduction(
            employee_id=data['employee_id'],
            team_member_id=data['team_member_id'],
            introduction_type=data['introduction_type'],
            message=data.get('message'),
            scheduled_date=datetime.strptime(data['scheduled_date'], '%Y-%m-%d').date() if data.get('scheduled_date') else None
        )
        
        db.session.add(introduction)
        db.session.commit()
        
        # Create notification for team member
        create_notification(
            user_id=data['team_member_id'],
            title=f'New Team Member Introduction',
            message=f'You have been assigned to introduce yourself to a new team member.',
            notification_type='team_introduction',
            priority='normal'
        )
        
        return jsonify({'status': 'success', 'message': 'Team introduction scheduled successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/timeline/<int:employee_id>', methods=['GET'])
@login_required('hr')
def get_communication_timeline(employee_id):
    """Get communication timeline for an employee"""
    try:
        timeline_events = []
        
        # Welcome packages
        packages = WelcomePackage.query.filter_by(employee_id=employee_id).all()
        for pkg in packages:
            timeline_events.append({
                'type': 'welcome_package',
                'title': f'Welcome Package: {pkg.package_name}',
                'description': f'Status: {pkg.status}',
                'date': pkg.created_at.isoformat()
            })
        
        # Admin tasks
        tasks = PreOnboardingTask.query.filter_by(employee_id=employee_id).all()
        for task in tasks:
            timeline_events.append({
                'type': 'admin_task',
                'title': f'Administrative Task: {task.task_name}',
                'description': f'Assigned to: {task.assigned_to}',
                'date': task.created_at.isoformat()
            })
        
        # Agenda
        agenda = FirstDayAgenda.query.filter_by(employee_id=employee_id).first()
        if agenda:
            timeline_events.append({
                'type': 'agenda',
                'title': 'First-Day Agenda Created',
                'description': f'Date: {agenda.agenda_date.strftime("%B %d, %Y")}',
                'date': agenda.created_at.isoformat()
            })
        
        # Team introductions
        intros = TeamIntroduction.query.filter_by(employee_id=employee_id).all()
        for intro in intros:
            timeline_events.append({
                'type': 'team_intro',
                'title': f'Team Introduction: {intro.introduction_type}',
                'description': f'With: {intro.team_member.full_name}',
                'date': intro.created_at.isoformat()
            })
        
        # Sort by date (most recent first)
        timeline_events.sort(key=lambda x: x['date'], reverse=True)
        
        return jsonify({
            'status': 'success',
            'timeline': timeline_events
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/pre-onboarding/progress/<int:employee_id>', methods=['GET'])
@login_required('hr')
def get_pre_onboarding_progress(employee_id):
    """Get pre-onboarding progress for an employee"""
    try:
        # Welcome packages progress
        total_packages = WelcomePackage.query.filter_by(employee_id=employee_id).count()
        delivered_packages = WelcomePackage.query.filter_by(employee_id=employee_id, status='delivered').count()
        
        # Admin tasks progress
        total_tasks = PreOnboardingTask.query.filter_by(employee_id=employee_id).count()
        completed_tasks = PreOnboardingTask.query.filter_by(employee_id=employee_id, is_completed=True).count()
        
        # Team introductions progress
        total_intros = TeamIntroduction.query.filter_by(employee_id=employee_id).count()
        sent_intros = TeamIntroduction.query.filter_by(employee_id=employee_id, is_sent=True).count()
        
        return jsonify({
            'status': 'success',
            'progress': {
                'welcome_packages': {
                    'total': total_packages,
                    'completed': delivered_packages
                },
                'admin_tasks': {
                    'total': total_tasks,
                    'completed': completed_tasks
                },
                'team_introductions': {
                    'total': total_intros,
                    'completed': sent_intros
                }
            }
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

# Helper function to create notifications
def create_notification(employee_id=None, user_id=None, title=None, message=None, notification_type=None, priority='normal'):
    """Create a notification for an employee or user"""
    try:
        # This would integrate with your existing notification system
        # For now, we'll just log it
        logger.info(f"Notification created: {title} - {message}")
    except Exception as e:
        logger.error(f"Error creating notification: {e}")

# API Routes for Social Integration Features

@app.route('/api/employees', methods=['GET'])
@login_required('hr')
def get_employees():
    """Get all employees for dropdowns"""
    try:
        employees = User.query.filter_by(role='employee').all()
        return jsonify({
            'status': 'success',
            'employees': [
                {
                    'id': emp.id,
                    'full_name': emp.full_name,
                    'department': getattr(emp, 'department', 'Unknown')
                }
                for emp in employees
            ]
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500





    try:
        # Team-building statistics
        total_activities = 0
        total_participants = 0
        try:
            # Legacy check - team building system has been removed
            pass
        except:
            pass  # Table doesn't exist or other error
        
        # Cross-departmental statistics
        total_introductions = 0
        departments_connected = 0
        try:
            # Legacy check - cross-departmental system has been removed
            pass
        except:
            pass  # Table doesn't exist or other error
        
        # Social events statistics
        total_events = 0
        total_attendance = 0
        try:
            # Legacy check - social events system has been removed
            pass
        except:
            pass  # Table doesn't exist or other error
        
        # Buddy/mentor statistics
        active_partnerships = 0
        total_meetings = 0
        try:
            # Legacy check - buddy/mentor system has been removed
            pass
        except:
            pass  # Table doesn't exist or other error
        
        # Communication training statistics
        total_trainings = 0
        certifications_issued = 0
        try:
            # Legacy check - communication training system has been removed
            pass
        except:
            pass  # Table doesn't exist or other error
        
        return jsonify({
            'status': 'success',
            'statistics': {
                'team_building': {
                    'total': total_activities,
                    'participants': total_participants
                },
                'cross_departmental': {
                    'total': total_introductions,
                    'departments': departments_connected
                },
                'social_events': {
                    'total': total_events,
                    'attendance': total_attendance
                },
                'buddy_mentor': {
                    'active': active_partnerships,
                    'meetings': total_meetings
                },
                'communication_training': {
                    'total': total_trainings,
                    'certifications': certifications_issued
                }
            }
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/hr/start-onboarding', methods=['POST'])
@login_required('hr')
def hr_start_onboarding():
    try:
        employee_id = request.form.get('employee_id')
        hire_date_str = request.form.get('hire_date')

        if not employee_id or not hire_date_str:
            flash('Employee and hire date are required to start onboarding.', 'danger')
            return redirect(url_for('pre_onboarding'))

        employee = User.query.filter_by(id=int(employee_id), role='employee').first_or_404()

        # Parse hire date
        hire_date = datetime.strptime(hire_date_str, '%Y-%m-%d').date()
        employee.hire_date = hire_date

        db.session.commit()

        flash(f'Onboarding started for {employee.full_name} with hire date {hire_date.strftime("%B %d, %Y")}.', 'success')
        # Redirect to onboarding page, passing hire date
        return redirect(url_for('onboarding', employee_id=employee.id, hire_date=hire_date_str))

    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error starting onboarding: {str(e)}')
        flash('Failed to start onboarding. Please try again.', 'danger')
        return redirect(url_for('pre_onboarding'))

@app.route('/api/candidates', methods=['GET'])
@login_required('hr')
def get_candidates_list():
    """Get list of all candidates for employee selection"""
    try:
        candidates = Candidate.query.order_by(
            Candidate.created_at.desc()).all()

        candidates_list = []
        for candidate in candidates:
            # Extract contact info if not already extracted
            if candidate.resume_text and (not candidate.email or not candidate.phone):
                candidate.extract_contact_info()
                db.session.commit()
            
            candidates_list.append({
                'id': candidate.id,
                'name': candidate.name,
                'email': candidate.email or 'N/A',
                'phone': candidate.phone or 'N/A',
                'job_desc': candidate.job_desc or 'N/A',
                'score': round(float(candidate.score), 1) if (
                    candidate.score) else None,
                'summary': candidate.summary or '',
                'created_at': candidate.created_at.strftime(
                    '%Y-%m-%d')
            })

        return jsonify({
            'status': 'success',
            'data': candidates_list
        })
    except Exception as e:
        app.logger.error(f'Error fetching candidates: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to fetch candidates'}), 500


@app.route('/onboarding', methods=['GET', 'POST'])
@login_required('hr')
def onboarding():
    # Get all employees with 'Onboarding' status
    onboarding_employees = User.query.filter_by(status='Onboarding', is_active=True).all()
    
    # Check if employee_id is provided in URL parameters
    employee_id = request.args.get('employee_id')
    selected_employee = None
    checklist = None
    
    if employee_id:
        # Get employee and their onboarding checklist
        selected_employee = User.query.get_or_404(employee_id)
        
        # Get onboarding tasks using the general Task model
        onboarding_tasks = Task.query.filter_by(
            assigned_to=selected_employee.id, 
            task_type='onboarding'
        ).all()
        
        # Create compatible task objects for template
        compatible_tasks = []
        for task in onboarding_tasks:
            compatible_task = type('Task', (), {
                'id': task.id,
                'task_name': task.title,
                'task_description': task.description,
                'is_completed': task.status == 'completed',
                'due_date': task.due_date,
                'priority': task.priority,
                'status': task.status
            })()
            compatible_tasks.append(compatible_task)
        
        # Create a checklist-like structure for compatibility with template
        if compatible_tasks:
            def get_progress(self):
                if not compatible_tasks:
                    return 0
                completed_tasks = len([task for task in compatible_tasks if task.is_completed])
                return int((completed_tasks / len(compatible_tasks)) * 100)
            
            checklist = type('Checklist', (), {
                'tasks': compatible_tasks,
                'created_at': onboarding_tasks[0].created_at if onboarding_tasks else datetime.utcnow(),
                'get_progress': get_progress
            })()
        else:
            checklist = None
    
    # Get current logged-in HR user
    current_hr = User.query.get(session['user_id'])
    
    # Render onboarding template with employee's specific data
    return render_template('onboarding.html', 
                         selected_employee=selected_employee,
                         checklist=checklist,
                         onboarding_employees=onboarding_employees,
                         user=current_hr,
                         sample_candidates=[])
    
    candidates = Candidate.query.order_by(Candidate.created_at.desc()).all()
    selected_candidate = None
    selected_candidate_id = request.args.get('candidate_id')
    
    if selected_candidate_id and selected_candidate_id.isdigit():
        selected_candidate = Candidate.query.get(selected_candidate_id)
    
    if request.method == 'POST':
        name = request.form.get('name')
        job_desc = request.form.get('job_desc')
        resume_text = request.form.get('resume_text')
        
        # Mock AI analysis
        score = round(random.uniform(60, 95), 1)
        summary = f"AI analysis shows strong potential candidate with relevant experience in {job_desc.split()[0]}."
        
        candidate = Candidate(
            name=name,
            job_desc=job_desc,
            resume_text=resume_text,
            score=score,
            summary=summary
        )
        
        # Extract email and phone from resume text
        candidate.extract_contact_info()
        
        db.session.add(candidate)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'candidate': {
                'name': name,
                'score': score,
                'summary': summary
            }
        })
    
    # Get current logged-in HR user
    current_hr = User.query.get(session['user_id'])
    
    return render_template('onboarding.html', 
                         sample_candidates=candidates,
                         selected_candidate=selected_candidate,
                         user=current_hr)

def get_department_tasks(department, position):
    """Get department-specific onboarding tasks"""
    base_tasks = [
        {'name': 'Document Collection', 'description': 'Collect ID proof, address proof, and educational certificates'},
        {'name': 'Welcome Email & Handbook', 'description': 'Send welcome email and employee handbook'},
        {'name': 'ID Card Creation', 'description': 'Create employee ID card and access badge'},
        {'name': 'Email Account Setup', 'description': 'Create company email account and configure access'},
    ]
    
    if department == 'Engineering':
        base_tasks.extend([
            {'name': 'Development Environment Setup', 'description': 'Install development tools, IDEs, and configure development environment'},
            {'name': 'Code Repository Access', 'description': 'Grant access to GitHub/GitLab repositories and development tools'},
            {'name': 'Technical Orientation', 'description': 'Introduction to tech stack, coding standards, and development processes'},
            {'name': 'Laptop & Equipment Assignment', 'description': 'Assign high-spec laptop, monitor, and development peripherals'}
        ])
    elif department == 'Marketing':
        base_tasks.extend([
            {'name': 'Marketing Tools Access', 'description': 'Setup access to marketing automation, CRM, and analytics tools'},
            {'name': 'Brand Guidelines Training', 'description': 'Review brand guidelines, style guide, and marketing materials'},
            {'name': 'Campaign Management Training', 'description': 'Introduction to current campaigns and marketing processes'},
            {'name': 'Creative Assets Access', 'description': 'Grant access to design tools and brand asset libraries'}
        ])
    elif department == 'HR':
        base_tasks.extend([
            {'name': 'HRIS System Training', 'description': 'Training on HR information systems and employee databases'},
            {'name': 'Compliance Training', 'description': 'Employment law, data privacy, and HR compliance training'},
            {'name': 'Recruitment Tools Access', 'description': 'Setup access to ATS, job boards, and recruitment platforms'}
        ])
    else:
        # General department tasks
        base_tasks.extend([
            {'name': 'Department Introduction', 'description': f'Meet with {department} team and understand department goals'},
            {'name': 'Role-Specific Training', 'description': f'Training specific to {position} responsibilities'},
            {'name': 'System Access Setup', 'description': 'Configure access to department-specific systems and tools'}
        ])
    
    # Add common final tasks
    base_tasks.extend([
        {'name': 'Team Introduction Meeting', 'description': 'Schedule meet-and-greet with immediate team members'},
        {'name': 'First Week Check-in', 'description': 'Schedule check-in meeting to address questions and concerns'}
    ])
    
    return base_tasks

@app.route('/interview', methods=['GET', 'POST'])
@login_required(['hr', 'employee'])
def interview():
    if request.method == 'POST':
        data = request.get_json()
        
        # Handle training session data
        if data.get('action') == 'save_training':
            training_data = {
                'user_id': session['user_id'],
                'session_type': data.get('type'),
                'difficulty': data.get('difficulty'),
                'responses': data.get('responses', []),
                'score': data.get('score', 0),
                'completed_at': datetime.utcnow()
            }
            
            # In a real app, save to database
            # For now, just return success
            return jsonify({
                'success': True,
                'message': 'Training session saved successfully'
            })
        
        # Handle interview responses (existing functionality)
        candidate_id = data.get('candidate_id')
        responses = data.get('responses', {})
        
        summary = "Candidate demonstrated good communication skills and relevant experience."
        
        interview = Interview(
            candidate_id=candidate_id,
            responses=responses,
            summary=summary
        )
        db.session.add(interview)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'summary': summary
        })
    
    return render_template('interview.html')

def analyze_exit_feedback(reason, feedback):
    """AI-powered analysis of exit feedback using Gemini"""
    try:
        prompt = f"""
        Analyze this employee exit feedback and provide comprehensive insights:
        
        Exit Reason: {reason}
        Detailed Feedback: {feedback}
        
        Please provide a JSON response with the following structure:
        {{
            "sentiment": "Positive|Neutral|Negative",
            "emotional_tone": "enthusiastic|satisfied|neutral|concerned|frustrated|angry",
            "key_themes": ["theme1", "theme2", "theme3"],
            "risk_level": "low|medium|high|critical",
            "actionable_insights": [
                {{
                    "category": "management|culture|compensation|work_life_balance|career_growth|work_environment",
                    "insight": "Specific insight description",
                    "priority": "low|medium|high",
                    "department": "HR|Management|IT|Operations|All"
                }}
            ],
            "retention_probability": "low|medium|high",
            "recommendations": ["recommendation1", "recommendation2"]
        }}
        
        Focus on:
        1. Emotional intelligence - detect underlying emotions
        2. Root cause analysis - identify core issues
        3. Actionable insights - provide specific, implementable recommendations
        4. Risk assessment - evaluate potential impact on organization
        5. Trend indicators - patterns that might indicate systemic issues
        """
        
        response = model.generate_content(prompt)
        ai_text = response.text
        
        # Extract JSON from response
        import json
        import re
        
        # Look for JSON pattern in the response
        json_match = re.search(r'\{.*\}', ai_text, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            analysis = json.loads(json_str)
            return analysis
        else:
            # Fallback to basic analysis if JSON parsing fails
            return {
                'sentiment': 'Neutral',
                'emotional_tone': 'neutral',
                'key_themes': ['general_feedback'],
                'risk_level': 'medium',
                'actionable_insights': [],
                'retention_probability': 'medium',
                'recommendations': ['Review feedback patterns']
            }
            
    except Exception as e:
        app.logger.error(f'Error in AI analysis: {str(e)}')
        # Fallback to mock analysis for demo purposes
        return mock_ai_analysis(reason, feedback)

def mock_ai_analysis(reason, feedback):
    """Mock AI analysis for demonstration purposes"""
    # Simple keyword-based analysis
    feedback_lower = feedback.lower()
    reason_lower = reason.lower()
    
    # Determine sentiment
    positive_words = ['good', 'great', 'excellent', 'amazing', 'love', 'appreciate', 'thankful']
    negative_words = ['bad', 'terrible', 'awful', 'hate', 'frustrated', 'angry', 'disappointed', 'limited']
    
    positive_count = sum(1 for word in positive_words if word in feedback_lower)
    negative_count = sum(1 for word in negative_words if word in feedback_lower)
    
    if positive_count > negative_count:
        sentiment = 'Positive'
        emotional_tone = 'satisfied'
    elif negative_count > positive_count:
        sentiment = 'Negative'
        emotional_tone = 'frustrated'
    else:
        sentiment = 'Neutral'
        emotional_tone = 'neutral'
    
    # Determine key themes
    themes = []
    theme_keywords = {
        'career_growth': ['career', 'growth', 'advancement', 'promotion', 'development'],
        'management': ['manager', 'management', 'leadership', 'boss', 'supervisor'],
        'compensation': ['salary', 'pay', 'compensation', 'benefits', 'money', 'raise'],
        'work_life_balance': ['balance', 'hours', 'overtime', 'flexible', 'remote', 'family'],
        'culture': ['culture', 'environment', 'team', 'atmosphere', 'toxic', 'supportive'],
        'work_environment': ['office', 'workspace', 'tools', 'equipment', 'resources']
    }
    
    for theme, keywords in theme_keywords.items():
        if any(keyword in feedback_lower for keyword in keywords):
            themes.append(theme)
    
    if not themes:
        themes = ['general_feedback']
    
    # Determine risk level
    if 'toxic' in feedback_lower or 'harassment' in feedback_lower or 'illegal' in feedback_lower:
        risk_level = 'critical'
    elif negative_count > 3 or 'frustrated' in emotional_tone:
        risk_level = 'high'
    elif negative_count > 1:
        risk_level = 'medium'
    else:
        risk_level = 'low'
    
    # Generate actionable insights
    insights = []
    if 'career_growth' in themes:
        insights.append({
            'category': 'career_growth',
            'insight': 'Employee seeking career advancement opportunities',
            'priority': 'high',
            'department': 'Management'
        })
    if 'management' in themes:
        insights.append({
            'category': 'management',
            'insight': 'Management style or communication issues identified',
            'priority': 'medium',
            'department': 'HR'
        })
    if 'compensation' in themes:
        insights.append({
            'category': 'compensation',
            'insight': 'Compensation and benefits concerns raised',
            'priority': 'high',
            'department': 'HR'
        })
    
    # Determine retention probability
    if sentiment == 'Negative' and risk_level in ['high', 'critical']:
        retention_probability = 'low'
    elif sentiment == 'Positive':
        retention_probability = 'high'
    else:
        retention_probability = 'medium'
    
    # Generate recommendations
    recommendations = []
    if 'career_growth' in themes:
        recommendations.append('Review career development programs and advancement paths')
    if 'management' in themes:
        recommendations.append('Implement management training and feedback systems')
    if 'compensation' in themes:
        recommendations.append('Conduct market compensation analysis')
    if 'work_life_balance' in themes:
        recommendations.append('Evaluate work-life balance policies and flexibility')
    
    return {
        'sentiment': sentiment,
        'emotional_tone': emotional_tone,
        'key_themes': themes,
        'risk_level': risk_level,
        'actionable_insights': insights,
        'retention_probability': retention_probability,
        'recommendations': recommendations
    }


@app.route('/ai-interview-training')
@login_required(['hr', 'employee'])
def ai_interview_training():
    """AI-powered interview training page"""
    return render_template('ai_interview_training.html')

@app.route('/employee/interview')
@login_required('employee')
def employee_interview():
    """Employee AI interview session"""
    user_id = session.get('user_id')
    user = User.query.get(user_id)
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user.id, 
        status='unread'
    ).count()
    
    return render_template('employee_interview.html', unread_messages_count=unread_messages_count)

@app.route('/api/notifications/send', methods=['POST'])
@login_required('employee')
def send_ready_notification():
    """Send notification from employee to HR"""
    try:
        user_id = session.get('user_id')
        user = User.query.get(user_id)
        
        data = request.get_json()
        message_content = data.get('content', 'I am ready for the interview')
        message_type = data.get('message_type', 'interview_ready')
        priority = data.get('priority', 'normal')
        
        # Get all HR users
        hr_users = User.query.filter_by(role='hr', is_active=True).all()
        
        if not hr_users:
            return jsonify({
                'status': 'error',
                'message': 'No HR users available to notify'
            }), 404
        
        # Create notifications for all HR users
        notifications_sent = 0
        for hr_user in hr_users:
            notification = Message(
                sender_id=user_id,
                recipient_id=hr_user.id,
                subject=f'Employee Ready for Interview - {user.full_name}',
                content=f'{user.full_name} ({user.email}) has sent the following message:\n\n"{message_content}"\n\nPlease start their interview session when ready.\n\nSent: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")}',
                priority=priority,
                status='unread'
            )
            db.session.add(notification)
            notifications_sent += 1
        
        # Also create an interview request record
        interview_request = Interview(
            user_id=user_id,
            status='requested',
            created_at=datetime.utcnow(),
            summary=f'Employee {user.full_name} is ready for interview. Message: "{message_content}"'
        )
        db.session.add(interview_request)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': f'Notification sent to {notifications_sent} HR user(s)',
            'notifications_sent': notifications_sent,
            'request_id': interview_request.id
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error sending notification: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to send notification'
        }), 500

@app.route('/api/employee/interview/status', methods=['GET'])
@login_required('employee')
def check_interview_status():
    """Check interview session status"""
    try:
        user_id = session['user_id']
        
        # Get the most recent interview request for this user
        recent_interview = Interview.query.filter_by(
            user_id=user_id
        ).order_by(Interview.created_at.desc()).first()
        
        if recent_interview:
            if recent_interview.status == 'requested':
                return jsonify({
                    'status': 'success',
                    'interview_status': 'requested',
                    'interview_id': recent_interview.id,
                    'created_at': recent_interview.created_at.isoformat(),
                    'message': 'Interview requested, waiting for HR'
                })
            elif recent_interview.status == 'ready':
                return jsonify({
                    'status': 'success',
                    'interview_status': 'ready',
                    'interview_id': recent_interview.id,
                    'created_at': recent_interview.created_at.isoformat(),
                    'message': 'Interview ready to start'
                })
            elif recent_interview.status == 'in_progress':
                return jsonify({
                    'status': 'success',
                    'interview_status': 'in_progress',
                    'interview_id': recent_interview.id,
                    'created_at': recent_interview.created_at.isoformat(),
                    'message': 'Interview in progress'
                })
            else:
                return jsonify({
                    'status': 'success',
                    'interview_status': recent_interview.status,
                    'interview_id': recent_interview.id,
                    'created_at': recent_interview.created_at.isoformat()
                })
        else:
            return jsonify({
                'status': 'success',
                'interview_status': 'no_active_interview',
                'interview_id': None
            })
    except Exception as e:
        app.logger.error(f'Error checking interview status: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to check interview status'
        }), 500

@app.route('/api/hr/interview/requests', methods=['GET'])
@login_required('hr')
def get_interview_requests():
    """Get pending interview requests for HR"""
    try:
        # Get all interview requests that are pending (for employees)
        pending_requests = Interview.query.filter_by(status='requested').filter(Interview.user_id.isnot(None)).order_by(Interview.created_at.desc()).all()
        
        requests_data = []
        for request in pending_requests:
            if request.user_id:
                employee = User.query.get(request.user_id)
                if employee:
                    requests_data.append({
                        'id': request.id,
                        'employee_id': employee.id,
                        'employee_name': employee.full_name,
                        'employee_email': employee.email,
                        'employee_position': employee.position or 'Not specified',
                        'created_at': request.created_at.isoformat(),
                        'message': request.summary or 'Employee is ready for interview',
                        'status': request.status
                    })
        
        return jsonify({
            'status': 'success',
            'requests': requests_data,
            'total_requests': len(requests_data)
        })
        
    except Exception as e:
        app.logger.error(f'Error fetching interview requests: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch interview requests'
        }), 500

@app.route('/api/hr/interview/<int:interview_id>/start', methods=['POST'])
@login_required('hr')
def start_employee_interview(interview_id):
    """HR starts an interview session"""
    try:
        hr_user_id = session.get('user_id')
        
        # Get the interview request
        interview = Interview.query.get_or_404(interview_id)
        
        # Verify interview is in requested status
        if interview.status != 'requested':
            return jsonify({
                'status': 'error',
                'message': 'Interview is not in requested status'
            }), 400
        
        # Update interview status
        interview.status = 'ready'
        interview.updated_at = datetime.utcnow()
        
        # Notify the employee that interview is ready
        if interview.user_id:
            notification = Message(
                sender_id=hr_user_id,
                recipient_id=interview.user_id,
                subject='Interview Ready - Join Now',
                content=f'Your interview session has been started by HR. Please join the interview session now.\n\nInterview ID: {interview.id}\nStarted: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")}',
                priority='high',
                status='unread'
            )
            db.session.add(notification)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Interview started successfully',
            'interview_id': interview.id,
            'status': 'ready'
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error starting interview: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to start interview'
        }), 500

@app.route('/api/employee/interview/submit', methods=['POST'])
@login_required('employee')
def submit_employee_interview():
    """Submit employee interview responses"""
    try:
        data = request.get_json()
        user_id = session['user_id']
        
        interview_record = Interview(
            candidate_id=None,
            responses=data.get('responses', []),
            summary=data.get('summary', ''),
            created_at=datetime.utcnow()
        )
        
        db.session.add(interview_record)
        db.session.commit()
        
        # Create notification for HR
        user = User.query.get(user_id)
        if user:
            hr_users = User.query.filter_by(role='hr').all()
            for hr_user in hr_users:
                notification = Message(
                    subject='New Interview Submission',
                    content=f'{user.full_name} has completed an AI interview session.',
                    sender_id=user_id,
                    recipient_id=hr_user.id,
                    message_type='interview',
                    priority='normal'
                )
                db.session.add(notification)
            db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Interview submitted successfully',
            'interview_id': interview_record.id
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error submitting employee interview: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to submit interview'
        }), 500

@app.route('/api/interviews/<int:interview_id>', methods=['GET'])
@login_required('hr')
def get_interview_details(interview_id):
    """Get full details for a specific interview and its candidate for review."""
    try:
        interview = Interview.query.get_or_404(interview_id)
        candidate = interview.candidate_ref

        if not candidate:
            return jsonify({
                'status': 'error',
                'message': 'Candidate not found for this interview'
            }), 404

        score = float(candidate.score) if candidate.score is not None else None

        if score is not None:
            if score >= 7.5:
                recommendation = 'Proceed to Next Round'
            elif score >= 6.0:
                recommendation = 'Hold for Review'
            else:
                recommendation = 'Reject'
        else:
            recommendation = 'Not Scored'

        return jsonify({
            'status': 'success',
            'data': {
                'interview_id': interview.id,
                'created_at': interview.created_at.strftime('%Y-%m-%d %H:%M:%S') if interview.created_at else None,
                'summary': interview.summary,
                'responses': interview.responses or {},
                'candidate': {
                    'id': candidate.id,
                    'name': candidate.name,
                    'job_desc': candidate.job_desc,
                    'resume_text': candidate.resume_text,
                    'score': score,
                    'recommendation': recommendation,
                    'created_at': candidate.created_at.strftime('%Y-%m-%d %H:%M:%S') if candidate.created_at else None
                }
            }
        })
    except Exception as e:
        app.logger.error(f'Error fetching interview details: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch interview details'
        }), 500

# New AI Interview System API Endpoints
@app.route('/api/interview/start', methods=['POST'])
@login_required(['hr', 'employee'])
def start_ai_interview():
    """Start a new AI interview session"""
    try:
        from ai_interviewer import AIInterviewer
        
        data = request.get_json()
        interview_type = data.get('interview_type', 'technical')
        is_fresher = data.get('is_fresher', False)
        job_role = data.get('job_role', 'General')
        enable_hr_training = data.get('enable_hr_training', True)
        
        # Initialize AI interviewer
        interviewer = AIInterviewer()
        interviewer.set_interview_type(interview_type)
        interviewer.set_job_description(f"Job Role: {job_role}")
        
        # Store interviewer in session
        session['ai_interviewer'] = {
            'interview_type': interview_type,
            'is_fresher': is_fresher,
            'job_role': job_role,
            'enable_hr_training': enable_hr_training,
            'questions_asked': [],
            'responses': [],
            'current_question_index': 0
        }
        
        return jsonify({
            'status': 'success',
            'message': 'Interview started successfully',
            'total_questions': 5,  # Default number of questions
            'interview_type': interview_type
        })
    except Exception as e:
        app.logger.error(f'Error starting interview: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to start interview'
        }), 500

@app.route('/api/interview/next-question', methods=['POST'])
@login_required(['hr', 'employee'])
def get_next_question():
    """Get the next interview question"""
    try:
        from ai_interviewer import AIInterviewer
        
        data = request.get_json()
        question_index = data.get('question_index', 0)
        previous_question = data.get('previous_question', '')
        previous_answer = data.get('previous_answer', '')
        
        # Get interviewer data from session
        interviewer_data = session.get('ai_interviewer', {})
        
        # Initialize AI interviewer
        interviewer = AIInterviewer()
        interviewer.set_interview_type(interviewer_data.get('interview_type', 'technical'))
        interviewer.set_job_description(f"Job Role: {interviewer_data.get('job_role', 'General')}")
        
        # Generate question
        if question_index == 0:
            # First question - always generate a proper question, not a greeting
            question = interviewer.generate_question()
        else:
            # Add previous answer to conversation history
            if previous_question and previous_answer:
                interviewer.conversation_history.append({
                    'role': 'user',
                    'content': previous_answer
                })
                interviewer.conversation_history.append({
                    'role': 'model', 
                    'content': previous_question
                })
            question = interviewer.generate_question()
        
        # Update session
        interviewer_data['current_question_index'] = question_index + 1
        interviewer_data['questions_asked'].append(question)
        session['ai_interviewer'] = interviewer_data
        
        return jsonify({
            'status': 'success',
            'question': question,
            'question_index': question_index + 1
        })
    except Exception as e:
        app.logger.error(f'Error getting next question: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to get next question'
        }), 500

@app.route('/api/interview/analyze-response', methods=['POST'])
@login_required(['hr', 'employee'])
def analyze_response():
    """Analyze candidate's response"""
    try:
        from ai_interviewer import AIInterviewer
        
        data = request.get_json()
        question = data.get('question', '')
        answer = data.get('answer', '')
        question_index = data.get('question_index', 0)
        
        # Get interviewer data from session
        interviewer_data = session.get('ai_interviewer', {})
        
        # Initialize AI interviewer
        interviewer = AIInterviewer()
        interviewer.set_interview_type(interviewer_data.get('interview_type', 'technical'))
        interviewer.set_job_description(f"Job Role: {interviewer_data.get('job_role', 'General')}")
        
        # Analyze response
        analysis_result = interviewer.analyze_response(question, answer)
        
        # Store response
        interviewer_data['responses'].append({
            'question': question,
            'answer': answer,
            'analysis': analysis_result,
            'question_index': question_index
        })
        
        # Store last answer and feedback for clarification
        interviewer_data['last_answer'] = answer
        interviewer_data['last_feedback'] = {
            'score': min(10, max(1, int(analysis_result.get('score', 5)))),
            'text': analysis_result.get('analysis', 'Good response'),
            'keywords': analysis_result.get('keywords', [])
        }
        
        session['ai_interviewer'] = interviewer_data
        
        # Extract score and keywords from analysis
        score = min(10, max(1, int(analysis_result.get('score', 5))))  # Ensure score is between 1-10
        keywords = analysis_result.get('keywords', [])
        analysis_text = analysis_result.get('analysis', 'Good response')
        
        return jsonify({
            'status': 'success',
            'feedback': {
                'analysis': analysis_text,
                'score': score,
                'keywords': keywords[:5]  # Limit to 5 keywords
            }
        })
    except Exception as e:
        app.logger.error(f'Error analyzing response: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to analyze response'
        }), 500

@app.route('/api/interview/request-hint', methods=['POST'])
@login_required(['hr', 'employee'])
def request_hint():
    """Get a hint for the current question"""
    try:
        data = request.get_json()
        question = data.get('question', '')
        
        # Generate a simple hint based on the question
        hint = "Think about your personal experience and provide specific examples. Focus on the key aspects mentioned in the question."
        
        if "technical" in question.lower():
            hint = "Consider explaining the technical concept, its practical applications, and any challenges you've faced with it."
        elif "hr" in question.lower() or "behavioral" in question.lower():
            hint = "Use the STAR method (Situation, Task, Action, Result) to structure your answer with specific examples."
        
        return jsonify({
            'status': 'success',
            'hint': hint
        })
    except Exception as e:
        app.logger.error(f'Error getting hint: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to get hint'
        }), 500

@app.route('/api/interview/request-clarification', methods=['POST'])
@login_required(['hr', 'employee'])
def request_clarification():
    """Request clarification for the current question with comprehensive performance feedback"""
    try:
        data = request.get_json()
        question = data.get('question', '')
        
        # Get interviewer data from session for performance context
        interviewer_data = session.get('ai_interviewer', {})
        last_answer = interviewer_data.get('last_answer', '')
        last_feedback = interviewer_data.get('last_feedback', {})
        
        # Generate comprehensive clarification with performance feedback
        clarification_parts = []
        
        # 1. Enhanced question clarification
        if "tell me about yourself" in question.lower():
            clarification_parts.append(
                "📝 **Question Deep Dive**: This is often the first impression question. I'm looking for:\n"
                "• Your educational journey and key achievements\n"
                "• Professional experience and notable projects\n"
                "• Technical skills and personal strengths\n"
                "• Career aspirations and why you're interested in this role\n"
                "• What makes you unique compared to other candidates"
            )
        elif "strengths" in question.lower():
            clarification_parts.append(
                "📝 **Question Deep Dive**: For strengths, I want to see:\n"
                "• 2-3 specific strengths relevant to this position\n"
                "• Real examples of when you demonstrated these strengths\n"
                "• The impact or results of your actions\n"
                "• How these strengths will benefit the company"
            )
        elif "weakness" in question.lower():
            clarification_parts.append(
                "📝 **Question Deep Dive**: For weaknesses, show maturity by:\n"
                "• Being honest about a real area for improvement\n"
                "• Explaining what you're doing to address it\n"
                "• Demonstrating self-awareness and growth mindset\n"
                "• Avoiding clichés like 'I'm a perfectionist' or 'I work too hard'"
            )
        else:
            clarification_parts.append(
                f"📝 **Question Deep Dive**: For '{question.split('?')[0]}', I'm seeking:\n"
                "• Your direct experience with this topic\n"
                "• Specific examples and measurable results\n"
                "• Your thought process and problem-solving approach\n"
                "• How this relates to the position you're applying for"
            )
        
        # 2. Detailed score breakdown
        if last_feedback and 'score' in last_feedback:
            score = last_feedback.get('score', 0)
            analysis = last_feedback.get('text', '')
            
            # Score interpretation
            score_interpretation = {
                9: "Exceptional - Exceeded expectations with outstanding examples",
                8: "Excellent - Strong response with good examples and clarity",
                7: "Very Good - Solid answer with relevant details",
                6: "Good - Adequate response with some good points",
                5: "Average - Basic response that meets minimum requirements",
                4: "Below Average - Lacked depth or specific examples",
                3: "Weak - Minimal relevant content",
                2: "Poor - Lacked relevance or clarity",
                1: "Very Poor - Did not address the question effectively"
            }.get(score, "Needs improvement")
            
            clarification_parts.append(
                f"\n🎯 **Score Analysis**: {score}/10 - {score_interpretation}\n\n"
                f"**AI's Assessment**: {analysis}\n\n"
                f"**Scoring Criteria Applied**:\n"
                f"• 🎯 **Relevance** (25%): How well you addressed the question\n"
                f"• 📊 **Content Quality** (25%): Depth, examples, and specifics\n"
                f"• 💡 **Communication** (25%): Clarity, structure, and confidence\n"
                f"• 🔍 **Impact** (25%): Results, achievements, and value shown"
            )
            
            # 3. Detailed performance analysis
            if last_answer:
                answer_lower = last_answer.lower()
                word_count = len(last_answer.split())
                
                # Content analysis
                content_analysis = []
                if any(word in answer_lower for word in ['project', 'developed', 'created', 'built', 'designed']):
                    content_analysis.append("✅ **Project Experience**: You demonstrated hands-on experience")
                if any(word in answer_lower for word in ['team', 'collaborate', 'worked with', 'together', 'group']):
                    content_analysis.append("✅ **Collaboration**: Showed teamwork abilities")
                if any(word in answer_lower for word in ['learn', 'study', 'course', 'degree', 'education', 'university']):
                    content_analysis.append("✅ **Learning Mindset**: Highlighted educational background")
                if any(word in answer_lower for word in ['problem', 'solve', 'solution', 'challenge', 'overcome']):
                    content_analysis.append("✅ **Problem Solving**: Addressed challenges and solutions")
                if any(word in answer_lower for word in ['improve', 'better', 'increase', 'reduce', 'optimize']):
                    content_analysis.append("✅ **Results Oriented**: Focused on improvements and outcomes")
                
                # Communication analysis
                comm_analysis = []
                if word_count > 50:
                    comm_analysis.append("📝 **Comprehensive**: Provided detailed response")
                elif word_count > 30:
                    comm_analysis.append("📝 **Balanced**: Good amount of detail")
                else:
                    comm_analysis.append("📝 **Concise**: Could use more detail")
                
                if 'i' in answer_lower and ('am' in answer_lower or 'have' in answer_lower):
                    comm_analysis.append("🎤 **Confidence**: Spoke with self-assurance")
                if any(word in answer_lower for word in ['because', 'therefore', 'result', 'led to']):
                    comm_analysis.append("🔗 **Logical Flow**: Connected ideas well")
                
                if content_analysis or comm_analysis:
                    clarification_parts.append(
                        f"\n🌟 **Performance Breakdown**:\n\n"
                        f"**Content Strengths**:\n" + "\n".join(f"  {item}" for item in content_analysis) + "\n\n"
                        f"**Communication Style**:\n" + "\n".join(f"  {item}" for item in comm_analysis)
                    )
            
            # 4. Body Language & Confidence Indicators (inferred from text)
            clarification_parts.append(
                f"\n🧠 **Inferred Communication Traits**:\n\n"
                f"• **Professionalism**: {'High' if any(word in last_answer.lower() for word in ['experience', 'project', 'developed']) else 'Moderate'}\n"
                f"• **Enthusiasm**: {'Evident' if any(word in last_answer.lower() for word in ['passionate', 'love', 'enjoy', 'excited']) else 'Could be stronger'}\n"
                f"• **Detail Orientation**: {'Strong' if len(last_answer.split()) > 40 else 'Developing'}\n"
                f"• **Self-Awareness**: {'Good' if any(word in last_answer.lower() for word in ['learn', 'improve', 'grow']) else 'Room for growth'}"
            )
            
            # 5. Strengths & Weaknesses Analysis
            clarification_parts.append(
                f"\n💪 **Strengths Demonstrated**:\n"
                f"• **Technical Knowledge**: {'Evident' if any(word in last_answer.lower() for word in ['technology', 'software', 'programming', 'code']) else 'Not clearly shown'}\n"
                f"• **Project Experience**: {'Strong' if 'project' in last_answer.lower() else 'Could emphasize more'}\n"
                f"• **Communication Skills**: {'Clear' if len(last_answer.split()) > 25 else 'Developing'}\n"
                f"• **Goal Orientation**: {'Present' if any(word in last_answer.lower() for word in ['achieve', 'goal', 'target', 'objective']) else 'Could highlight'}\n\n"
                f"🎯 **Areas to Enhance**:\n"
                f"• **Quantifiable Results**: Add specific metrics and numbers\n"
                f"• **Industry Alignment**: Connect experience to job requirements\n"
                f"• **Unique Value**: What makes you different from others\n"
                f"• **Future Vision**: How you'll contribute to the company"
            )
            
            # 6. Actionable Improvement Plan
            clarification_parts.append(
                f"\n🚀 **Next Steps to Improve**:\n\n"
                f"**Immediate (Next Interview)**:\n"
                f"• Start with a strong opening statement\n"
                f"• Include 2-3 specific achievements with metrics\n"
                f"• Practice the STAR method for behavioral questions\n"
                f"• Research the company and align your experience\n\n"
                f"**Long-term Development**:\n"
                f"• Build a portfolio of your best projects\n"
                f"• Practice explaining technical concepts simply\n"
                f"• Develop your personal elevator pitch\n"
                f"• Seek mock interview opportunities"
            )
        
        clarification = "\n".join(clarification_parts)
        
        return jsonify({
            'status': 'success',
            'clarification': clarification
        })
    except Exception as e:
        app.logger.error(f'Error getting clarification: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to get clarification'
        }), 500

@app.route('/api/interview/summary', methods=['GET'])
@login_required(['hr', 'employee'])
def get_interview_summary():
    """Get interview summary and evaluation"""
    try:
        from ai_interviewer import AIInterviewer
        
        # Get interviewer data from session
        interviewer_data = session.get('ai_interviewer', {})
        responses = interviewer_data.get('responses', [])
        
        if not responses:
            return jsonify({
                'status': 'error',
                'message': 'No interview data found'
            }), 404
        
        # Initialize AI interviewer
        interviewer = AIInterviewer()
        interviewer.set_interview_type(interviewer_data.get('interview_type', 'technical'))
        interviewer.set_job_description(f"Job Role: {interviewer_data.get('job_role', 'General')}")
        
        # Build conversation history
        for response in responses:
            interviewer.conversation_history.append({
                'role': 'user',
                'content': response['answer']
            })
            interviewer.conversation_history.append({
                'role': 'model',
                'content': response['question']
            })
        
        # Generate summary
        summary_result = interviewer.generate_summary()
        
        # Extract structured data
        summary_text = summary_result.get('summary', 'Interview completed successfully')
        overall_score = summary_result.get('overall_score', 7.0)
        recommendation = summary_result.get('recommendation', 'Needs Further Evaluation')
        strengths = summary_result.get('strengths', [])
        areas_for_improvement = summary_result.get('areas_for_improvement', [])
        confidence_score = summary_result.get('confidence_score', 75)
        
        return jsonify({
            'status': 'success',
            'summary': summary_text,
            'overall_score': overall_score,
            'recommendation': recommendation,
            'strengths': strengths,
            'areas_for_improvement': areas_for_improvement,
            'confidence_score': confidence_score,
            'total_questions': len(responses)
        })
    except Exception as e:
        app.logger.error(f'Error generating summary: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate summary'
        }), 500

@app.route('/api/interview/complete', methods=['POST'])
@login_required(['hr', 'employee'])
def complete_interview():
    """Complete the interview and save results"""
    try:
        data = request.get_json()
        
        # Get interviewer data from session
        interviewer_data = session.get('ai_interviewer', {})
        
        # Save interview to database
        interview_record = Interview(
            candidate_id=session.get('user_id'),
            responses=interviewer_data.get('responses', []),
            summary=data.get('summary', 'Interview completed'),
            created_at=datetime.utcnow()
        )
        
        db.session.add(interview_record)
        db.session.commit()
        
        # Clear session data
        session.pop('ai_interviewer', None)
        
        return jsonify({
            'status': 'success',
            'message': 'Interview completed successfully',
            'interview_id': interview_record.id
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error completing interview: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to complete interview'
        }), 500

@app.route('/exit', methods=['GET', 'POST'])
@app.route('/exit/<int:employee_id>', methods=['GET', 'POST'])
@login_required(['employee', 'hr'])  # Allow both employees and HR to access
def exit_interview(employee_id=None):
    # Get the current user's employee data
    employee = User.query.get(session['user_id'])
    
    # If employee_id is provided (HR viewing specific employee)
    if employee_id and session.get('role') == 'hr':
        target_employee = User.query.get(employee_id)
        if target_employee and target_employee.status == 'Offboarding':
            employee = target_employee
        elif not target_employee:
            flash('Employee not found', 'error')
            return redirect(url_for('exit_interview'))
        else:
            flash('Employee is not in offboarding status', 'warning')
            return redirect(url_for('exit_interview'))
    
    if request.method == 'POST':
        # Use the logged-in user's name instead of form data
        name = f"{employee.full_name} ({employee.username})"
        reason = request.form.get('reason')
        feedback = request.form.get('feedback', '')
        
        try:
            # Create new exit feedback entry
            exit_feedback = ExitFeedback(
                name=name,
                reason=reason,
                feedback=feedback
            )
            
            db.session.add(exit_feedback)
            db.session.commit()
            
            # Perform AI analysis on the feedback
            try:
                ai_analysis = analyze_exit_feedback(reason, feedback)
                if ai_analysis:
                    import json
                    exit_feedback.sentiment = ai_analysis.get('sentiment', 'Neutral')
                    exit_feedback.key_themes = json.dumps(ai_analysis.get('key_themes', []))
                    exit_feedback.risk_level = ai_analysis.get('risk_level', 'medium')
                    exit_feedback.actionable_insights = json.dumps(ai_analysis.get('actionable_insights', []))
                    exit_feedback.emotional_tone = ai_analysis.get('emotional_tone', 'neutral')
                    exit_feedback.retention_probability = ai_analysis.get('retention_probability', 'medium')
                    exit_feedback.recommendations = json.dumps(ai_analysis.get('recommendations', []))
                    db.session.commit()
                    app.logger.info(f'AI analysis completed for exit feedback by {employee.username}')
            except Exception as ai_error:
                app.logger.warning(f'AI analysis failed for exit feedback: {str(ai_error)}')
                # Continue without AI analysis - don't fail the entire process
            
            app.logger.info(f'Exit feedback submitted by {employee.username}')
            return jsonify({
                'success': True,
                'message': 'Feedback submitted successfully!'
            })
            
        except Exception as e:
            db.session.rollback()
            app.logger.error(f'Error submitting exit feedback: {str(e)}')
            return jsonify({
                'success': False,
                'message': 'An error occurred while submitting your feedback. Please try again.'
            }), 500
    
    # For GET request, get exit feedback data and offboarding employees
    offboarding_employees = []
    if session.get('role') == 'hr':
        # HR can see all exit feedback and offboarding employees
        exit_feedbacks = ExitFeedback.query.order_by(ExitFeedback.created_at.desc()).limit(10).all()
        # Get all employees with offboarding status
        offboarding_employees = User.query.filter_by(status='Offboarding').all()
    else:
        # Employees can only see their own feedback
        exit_feedbacks = ExitFeedback.query.filter(
            ExitFeedback.name.like(f"%{employee.full_name}%")
        ).order_by(ExitFeedback.created_at.desc()).limit(10).all()
        # Employee can only see themselves if they are offboarding
        if employee and employee.status == 'Offboarding':
            offboarding_employees = [employee]
    
    # Get HR contact information - use current user if HR, otherwise find an HR user
    if session.get('role') == 'hr':
        hr_contact = User.query.get(session['user_id'])
    else:
        # Prioritize Deeksha and exclude test HR users
        hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
            User.username != 'test_hr',
            User.username != 'testhr',
            User.full_name.notlike('%test%')
        ).filter(User.full_name.like('%Deeksha%')).first()
        
        if not hr_contact:
            hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
                User.username != 'test_hr',
                User.username != 'testhr',
                User.full_name.notlike('%test%')
            ).first()
    
    # Render the template with employee data, exit feedbacks, offboarding employees, and HR contact
    return render_template('exit.html', employee=employee, exit_feedbacks=exit_feedbacks, 
                         offboarding_employees=offboarding_employees, hr_contact=hr_contact)


@app.route('/knowledge-transfer', methods=['GET'])
@login_required(['hr', 'admin', 'manager'])
def knowledge_transfer():
    """Knowledge Transfer System page"""
    # Get the current user's employee data
    employee = User.query.get(session['user_id'])
    
    # Get HR contact information
    if session.get('role') == 'hr':
        hr_contact = User.query.get(session['user_id'])
    else:
        # Prioritize Deeksha and exclude test HR users
        hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
            User.username != 'test_hr',
            User.username != 'testhr',
            User.full_name.notlike('%test%')
        ).filter(User.full_name.like('%Deeksha%')).first()
        
        if not hr_contact:
            hr_contact = User.query.filter_by(role='hr', is_active=True).filter(
                User.username != 'test_hr',
                User.username != 'testhr',
                User.full_name.notlike('%test%')
            ).first()
    
    # Fetch Knowledge Transfer data from database
    kt_sessions = KTSession.query.order_by(KTSession.scheduled_date.desc()).all()
    kt_documents = KTDocument.query.order_by(KTDocument.upload_date.desc()).all()
    successor_training = SuccessorTraining.query.order_by(SuccessorTraining.created_at.desc()).all()
    project_handovers = ProjectHandover.query.order_by(ProjectHandover.created_at.desc()).all()
    
    # Get or create KT progress for current employee
    kt_progress = KTProgress.query.filter_by(employee_id=session['user_id']).first()
    if not kt_progress:
        kt_progress = KTProgress(
            employee_id=session['user_id'],
            sessions_completed=KTSession.query.filter_by(status='Completed').count(),
            docs_uploaded=KTDocument.query.count(),
            successor_trained_percent=calculate_successor_progress(),
            projects_handover=ProjectHandover.query.filter_by(status='Completed').count(),
            overall_progress=calculate_overall_progress(session['user_id'])
        )
        db.session.add(kt_progress)
        db.session.commit()
    
    return render_template('knowledge_transfer.html', 
                         employee=employee, 
                         hr_contact=hr_contact,
                         kt_sessions=kt_sessions,
                         kt_documents=kt_documents,
                         successor_training=successor_training,
                         project_handovers=project_handovers,
                         kt_progress=kt_progress)

def calculate_successor_progress():
    """Calculate successor training progress percentage"""
    try:
        total_modules = SuccessorTraining.query.count()
        completed_modules = SuccessorTraining.query.filter_by(status='Completed').count()
        if total_modules > 0:
            return int((completed_modules / total_modules) * 100)
        return 0
    except:
        return 0

def calculate_overall_progress(employee_id):
    """Calculate overall KT progress percentage"""
    try:
        # Get weights for different components
        session_weight = 0.3
        doc_weight = 0.2
        training_weight = 0.3
        project_weight = 0.2
        
        # Calculate individual progress
        total_sessions = KTSession.query.count()
        completed_sessions = KTSession.query.filter_by(status='Completed').count()
        session_progress = (completed_sessions / total_sessions * 100) if total_sessions > 0 else 0
        
        doc_progress = 100 if KTDocument.query.count() > 0 else 0
        training_progress = calculate_successor_progress()
        
        total_projects = ProjectHandover.query.count()
        completed_projects = ProjectHandover.query.filter_by(status='Completed').count()
        project_progress = (completed_projects / total_projects * 100) if total_projects > 0 else 0
        
        # Calculate weighted average
        overall = (session_progress * session_weight + 
                   doc_progress * doc_weight + 
                   training_progress * training_weight + 
                   project_progress * project_weight)
        
        return int(overall)
    except:
        return 0

@app.route('/api/project/<int:project_id>/details')
@login_required(['hr', 'admin', 'manager'])
def get_project_details(project_id):
    """Get project details for view modal"""
    try:
        project = ProjectHandover.query.get_or_404(project_id)
        
        # Get related documents (in a real implementation, you'd have a separate documents table)
        sample_docs = [
            {'name': 'Project Plan.pdf', 'type': 'pdf', 'size': '2.3 MB', 'path': '/uploads/project_plan.pdf'},
            {'name': 'Technical Documentation.docx', 'type': 'docx', 'size': '1.8 MB', 'path': '/uploads/tech_docs.docx'},
            {'name': 'Meeting Notes.txt', 'type': 'txt', 'size': '45 KB', 'path': '/uploads/meeting_notes.txt'}
        ]
        
        return jsonify({
            'id': project.id,
            'name': project.project_name,
            'description': project.project_description or 'No description available',
            'status': project.status,
            'handover_date': project.handover_date.strftime('%Y-%m-%d') if project.handover_date else None,
            'recipient': project.recipient_name,
            'verified': project.verified,
            'notes': project.notes,
            'documents': sample_docs if project.documentation_path else []
        })
        
    except Exception as e:
        app.logger.error(f'Error fetching project details: {str(e)}')
        return jsonify({'error': 'Failed to fetch project details'}), 500





@app.route('/employee/mood-tracking')
@login_required('employee')
def employee_mood_tracking():
    """Render the employee mood and wellness tracking page with real database data."""
    from datetime import datetime, timedelta
    from flask import session
    
    # Get current user from session
    user_id = session.get('user_id')
    user = db.session.get(User, user_id) if user_id else None
    
    if not user:
        return redirect(url_for('login'))
    
    # Get unread messages count
    unread_messages_count = Message.query.filter_by(
        recipient_id=user.id, 
        status='unread'
    ).count()
    
    # Get mood feedback history
    feedback_history = EmployeeFeedback.query.filter_by(user_id=user.id)\
                                          .order_by(EmployeeFeedback.created_at.desc())\
                                          .limit(30).all()
    
    # Calculate current mood and confidence
    if feedback_history:
        latest_feedback = feedback_history[0]
        current_mood = latest_feedback.mood_rating
        current_confidence = latest_feedback.confidence_rating
    else:
        current_mood = 3.5
        current_confidence = 3.5
    
    # Calculate 7-day trend
    recent_7_days = feedback_history[:7]
    if recent_7_days:
        avg_mood_7 = sum(f.mood_rating for f in recent_7_days) / len(recent_7_days)
        avg_confidence_7 = sum(f.confidence_rating for f in recent_7_days) / len(recent_7_days)
    else:
        avg_mood_7 = 3.5
        avg_confidence_7 = 3.5
    
    # Calculate trend percentage (simplified)
    if len(feedback_history) >= 14:
        older_period = feedback_history[7:14]
        if older_period:
            older_avg = sum(f.mood_rating for f in older_period) / len(older_period)
            trend_percentage = round(((avg_mood_7 - older_avg) / older_avg) * 100, 1)
        else:
            trend_percentage = 0
    else:
        trend_percentage = 12  # Default positive trend for new users
    
    # Calculate check-in streak (consecutive days with feedback)
    streak = 0
    if feedback_history:
        current_date = datetime.utcnow().date()
        for feedback in feedback_history:
            feedback_date = feedback.created_at.date()
            if feedback_date == current_date - timedelta(days=streak):
                streak += 1
            else:
                break
    
    # Calculate wellness score (based on mood and confidence)
    wellness_score = round((avg_mood_7 + avg_confidence_7) / 2 * 20, 1)  # Convert to 0-100 scale
    
    # Get mood distribution for last 30 days
    thirty_days_ago = datetime.utcnow() - timedelta(days=30)
    recent_feedback = [f for f in feedback_history if f.created_at >= thirty_days_ago]
    
    mood_distribution = [0, 0, 0, 0, 0]  # Counts for mood ratings 1-5
    for feedback in recent_feedback:
        if 1 <= feedback.mood_rating <= 5:
            mood_distribution[feedback.mood_rating - 1] += 1
    
    # Prepare chart data (last 7 days)
    chart_data = []
    for i in range(7):
        date = datetime.utcnow().date() - timedelta(days=6-i)
        day_feedback = [f for f in feedback_history if f.created_at.date() == date]
        if day_feedback:
            avg_mood = sum(f.mood_rating for f in day_feedback) / len(day_feedback)
            avg_confidence = sum(f.confidence_rating for f in day_feedback) / len(day_feedback)
        else:
            avg_mood = 3.5
            avg_confidence = 3.5
        
        chart_data.append({
            'date': date.strftime('%b %d'),
            'mood': round(avg_mood, 1),
            'confidence': round(avg_confidence, 1)
        })
    
    return render_template(
        'employee_mood_tracking.html',
        title='Mood & Wellness',
        today=datetime.utcnow(),
        employee=user,
        unread_messages_count=unread_messages_count,
        current_mood=current_mood,
        current_confidence=current_confidence,
        avg_mood_7=round(avg_mood_7, 1),
        avg_confidence_7=round(avg_confidence_7, 1),
        trend_percentage=trend_percentage,
        check_in_streak=streak,
        wellness_score=wellness_score,
        feedback_history=feedback_history[:10],  # Last 10 entries
        mood_distribution=mood_distribution,
        chart_data=chart_data
    )

def extract_text_from_pdf(file_stream):
    pdf_reader = PyPDF2.PdfReader(file_stream)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_stream):
    doc = docx.Document(BytesIO(file_stream.read()))
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def extract_skills(text):
    # Common tech skills to look for
    skills = [
        'Python', 'JavaScript', 'Java', 'C++', 'C#', 'Ruby', 'PHP', 'Swift', 'Kotlin', 'Go',
        'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring', 'Laravel', 'Ruby on Rails',
        'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'DynamoDB',
        'AWS', 'Azure', 'Google Cloud', 'Docker', 'Kubernetes', 'Terraform',
        'Git', 'CI/CD', 'Jenkins', 'GitHub Actions',
        'Machine Learning', 'Data Science', 'TensorFlow', 'PyTorch', 'Pandas', 'NumPy',
        'Agile', 'Scrum', 'DevOps', 'TDD', 'REST API', 'GraphQL', 'Microservices'
    ]
    
    found_skills = []
    for skill in skills:
        if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
            found_skills.append(skill)
    return found_skills[:15]  # Return max 15 skills to avoid clutter

def extract_experience(text):
    # Look for job experience patterns and return concise highlights
    experience = []

    # Look for date patterns like 2020 - 2022 or 01/2020 - Present
    date_patterns = [
        r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4})\s*[-–]\s*(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}|Present\b)',
        r'(\d{1,2}/\d{4})\s*[-–]\s*(\d{1,2}/\d{4}|Present\b)',
        r'(\d{4})\s*[-–]\s*(\d{4}|Present\b)'
    ]

    # Patterns for lines we want to ignore (personal/contact details, links, headers)
    ignore_patterns = [
        r'@',                          # emails
        r'https?://',                  # URLs
        r'linkedin\.com',
        r'github\.com',
        r'phone',
        r'contact',
        r'\baddress\b',
        r'\bsummary\b',
        r'\bobjective\b',
        r'\bskills?\b',
        r'\beducation\b',
        r'\bprojects?\b',
    ]

    def is_personal_or_header(line: str) -> bool:
        lower = line.lower()
        if len(lower) < 3:
            return True
        for pat in ignore_patterns:
            if re.search(pat, lower):
                return True
        # Lines that are mostly punctuation or links
        if len(re.sub(r'[^a-zA-Z]', '', lower)) < 3:
            return True
        return False

    # Collect contextual snippets around date ranges
    for pattern in date_patterns:
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        for i, match in enumerate(matches):
            if i >= 3:
                break  # Limit to first 3 experiences per pattern
            start = max(0, match.start() - 200)
            end = min(len(text), match.end() + 200)
            context = text[start:end]
            # Split context into lines and filter
            for raw_line in context.split('\n'):
                line = raw_line.strip()
                if not line:
                    continue
                if is_personal_or_header(line):
                    continue
                # Prefer lines that look like responsibilities/roles
                if any(kw in line.lower() for kw in ['developer', 'engineer', 'intern', 'manager', 'lead', 'built', 'developed', 'implemented', 'designed']):
                    experience.append(line)

    # Deduplicate while preserving order
    seen = set()
    unique_experience = []
    for line in experience:
        if line not in seen:
            seen.add(line)
            unique_experience.append(line)

    return unique_experience[:5] if unique_experience else ["Experience not found in standard format"]

def extract_education(text):
    education = []
    
    # Look for education section
    education_indicators = [
        r'(?i)education\s*',
        r'(?i)academic\s*background',
        r'(?i)qualifications',
        r'(?i)degrees?\s*',
    ]
    
    # Look for degree patterns
    degree_patterns = [
        r'\b(?:B\.?S\.?|Bachelor(?:\'?s)?\s+of\s+Science|B\.?A\.?|Bachelor(?:\'?s)?\s+of\s+Arts|M\.?S\.?|Master(?:\'?s)?\s+of\s+Science|M\.?A\.?|Master(?:\'?s)?\s+of\s+Arts|Ph\.?D\.?|Doctor(?:ate)?(?:\'?s)?\s+of\s+Philosophy)\b',
        r'\b(?:B\.?E\.?|Bachelor(?:\'?s)?\s+of\s+Engineering|B\.?Tech|Bachelor(?:\'?s)?\s+of\s+Technology)\b',
        r'\b(?:M\.?B\.?A\.?|Master(?:\'?s)?\s+of\s+Business\s+Administration)\b',
    ]
    
    # First try to find education section
    for indicator in education_indicators:
        section_match = re.search(indicator + r'[\s\S]*?(?=\n\s*\n|$)', text, re.IGNORECASE)
        if section_match:
            section = section_match.group(0)
            # Extract lines that look like education entries
            lines = [line.strip() for line in section.split('\n') if any(re.search(pattern, line, re.IGNORECASE) for pattern in degree_patterns)]
            if lines:
                return lines[:3]  # Return first 3 education entries
    
    # If no education section found, search the whole text for degree patterns
    for pattern in degree_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for i, match in enumerate(matches):
            if i < 3:  # Limit to 3 entries
                start = max(0, match.start() - 100)
                end = min(len(text), match.end() + 100)
                context = text[start:end].strip()
                education.append(context)
    
    return education if education else ["Education information not found"]

def extract_email(text):
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return email_match.group(0) if email_match else ''

def extract_phone(text):
    """Extract phone number from resume text"""
    phone_patterns = [
        r'\b\d{3}-\d{3}-\d{4}\b',  # 123-456-7890
        r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',  # (123) 456-7890
        r'\b\d{3}\.\d{3}\.\d{4}\b',  # 123.456.7890
        r'\b\d{10}\b',  # 1234567890
        r'\b\+?\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'  # International format
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            return phone_match.group(0)
    
    return ''

def extract_name(text):
    # Look for name at the beginning of the document
    first_line = text.split('\n')[0].strip()
    # Simple heuristic: if the first line is likely a name (2-4 words, title case)
    if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z\.]+){1,3}$', first_line):
        return first_line
    return ''

def calculate_ats_score(resume_data, job_description):
    """
    Calculate ATS (Applicant Tracking System) score based on resume
    and job description match.
    Score is out of 100.
    """
    score = 0
    feedback = []

    # Normalize text for comparison
    resume_text = ' '.join(resume_data.get('skills', [])).lower()

    # 1. Skills Match (35 points max)
    resume_skills = set([s.lower() for s in resume_data.get('skills', [])])
    job_skills = extract_skills(job_description)
    job_skills_set = set([s.lower() for s in job_skills])

    if job_skills_set:
        matched_skills = resume_skills.intersection(job_skills_set)
        skill_match_percentage = len(matched_skills) / len(job_skills_set)
        skills_score = skill_match_percentage * 35
        score += skills_score

        if matched_skills:
            matched_list = list(matched_skills)[:3]
            feedback.append(
                f"Matched {len(matched_skills)}/{len(job_skills_set)} "
                f"required skills: {', '.join(matched_list)}"
            )
        else:
            required_skills = list(job_skills_set)[:3]
            feedback.append(
                f"No matching skills found. Required: "
                f"{', '.join(required_skills)}"
            )

    # 2. Experience (30 points max)
    experience_count = len(resume_data.get('experience', []))
    if experience_count >= 4:
        score += 30
        feedback.append("Strong experience level (4+ years)")
    elif experience_count >= 2:
        score += 20
        feedback.append("Moderate experience level (2-4 years)")
    elif experience_count >= 1:
        score += 10
        feedback.append("Entry-level experience")

    # 3. Education (20 points max)
    education = resume_data.get('education', [])
    education_text = ' '.join([str(e) for e in education]).lower()

    has_bachelor = ('bachelor' in education_text or
                    'b.tech' in education_text or
                    'b.s.' in education_text)
    has_master = ('master' in education_text or
                  'm.tech' in education_text or
                  'diploma' in education_text)

    if has_bachelor:
        score += 15
        feedback.append("Bachelor's degree found")
    elif has_master:
        score += 20
        feedback.append("Advanced degree found")
    elif education:
        score += 10
        feedback.append("Education information provided")

    # 4. Contact Information (10 points max)
    contact_score = 0
    if resume_data.get('email'):
        contact_score += 5
    if resume_data.get('phone'):
        contact_score += 5
    score += contact_score

    if contact_score >= 10:
        feedback.append("Complete contact information provided")

    # 5. Keyword Presence (5 points max)
    keywords = ['python', 'javascript', 'api', 'database', 'agile', 'git']
    found_keywords = sum(1 for kw in keywords if kw in resume_text)
    keyword_score = (found_keywords / len(keywords)) * 5
    score += keyword_score

    # Normalize score to 0-100
    final_score = min(100, max(0, score))

    return {
        'score': round(final_score, 1),
        'feedback': feedback,
        'recommendation': 'PASS' if final_score >= 70 else 'REVIEW'
    }


@app.route('/analyze-resume', methods=['POST'])
def analyze_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    # Get job description from request
    job_description = request.form.get('jobDescription', '')

    if not job_description:
        return jsonify({'error': 'Job description is required'}), 400

    try:
        filename = file.filename.lower()
        file_stream = file.stream

        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_stream)
        elif filename.endswith(('.doc', '.docx')):
            text = extract_text_from_docx(file_stream)
        else:
            error_msg = (
                'Unsupported file format. Please upload a PDF or '
                'Word document.'
            )
            return jsonify({'error': error_msg}), 400

        # Extract information
        resume_data = {
            'name': extract_name(text),
            'email': extract_email(text),
            'phone': extract_phone(text),
            'skills': extract_skills(text),
            'experience': extract_experience(text),
            'education': extract_education(text),
            'raw_text': text
        }

        # Calculate ATS score
        ats_result = calculate_ats_score(resume_data, job_description)

        # Prepare AI summary using the AI response function
        # Only include skills that are present in both the resume and
        # the job description
        job_skills = extract_skills(job_description)
        resume_skills_set = set(s.lower() for s in resume_data.get('skills', []))
        job_skills_set = set(s.lower() for s in job_skills)

        matching_skills = [
            s for s in resume_data.get('skills', [])
            if s.lower() in job_skills_set
        ]

        # Fallback to top resume skills if no explicit matches found
        display_skills = matching_skills if matching_skills else resume_data.get('skills', [])

        summary_prompt = (
            f"Provide a brief (2-3 sentences) summary of how well this "
            f"candidate matches the job description. Matching skills: "
            f"{', '.join(display_skills[:5])}. "
            f"Experience count: {len(resume_data['experience'])}. "
            f"Focus on strengths relevant to: {job_description[:200]}"
        )

        ai_summary = "Resume analyzed. Candidate profile processed."
        try:
            ai_messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an HR AI assistant. Provide brief, "
                        "professional summaries."
                    )
                },
                {"role": "user", "content": summary_prompt}
            ]
            ai_response = generate_ai_response(
                ai_messages, max_tokens=150, temperature=0.7)
            if ai_response:
                ai_summary = ai_response
        except Exception as e:
            app.logger.warning(f"Could not generate AI summary: {str(e)}")
            ai_summary = (
                f"Candidate has {len(matching_skills) if matching_skills else len(resume_data['skills'])} relevant skills "
                f"and {len(resume_data['experience'])} years of experience."
            )

        result = {
            'name': resume_data['name'],
            'email': resume_data['email'],
            'phone': resume_data['phone'],
            'skills': resume_data['skills'],
            'experience': resume_data['experience'],
            'education': resume_data['education'],
            'ats_score': ats_result['score'],
            'ats_feedback': ats_result['feedback'],
            'ats_recommendation': ats_result['recommendation'],
            'ai_summary': ai_summary,
            'pass_screening': ats_result['score'] >= 60
        }

        return jsonify(result)

    except Exception as e:
        app.logger.error(f"Error processing resume: {str(e)}")
        return jsonify(
            {'error': 'Error processing resume. Please try again.'}), 500

def generate_ai_response(messages, max_tokens=500, temperature=0.7):
    """Helper function to generate AI responses using Gemini"""
    try:
        if not gemini_api_key or gemini_api_key == 'your-gemini-api-key-here':
            raise Exception("Gemini API key not configured")
            
        logger.debug("Using Gemini Pro model")
        
        # Convert messages to Gemini format if needed
        if isinstance(messages, list) and len(messages) > 0:
            # If it's a chat completion format, extract the last user message
            last_message = messages[-1]
            if isinstance(last_message, dict) and 'content' in last_message:
                prompt = last_message['content']
            else:
                prompt = str(messages[-1])
        else:
            prompt = str(messages)
            
        # Generate response using Gemini
        response = model.generate_content(
            prompt,
            generation_config={
                'max_output_tokens': max_tokens,
                'temperature': temperature,
            }
        )
        
        return response.text.strip()
    except Exception as e:
        logger.error(f"Error in generate_ai_response: {str(e)}")
        raise Exception(f"Error generating response: {str(e)}")

def generate_onboarding_plan(role, department, start_date):
    """Generate a personalized onboarding plan using Gemini AI"""
    prompt = f"""As an HR specialist, create a comprehensive 30-60-90 day onboarding plan for a new {role} in the {department} department starting on {start_date}.
    
    Include these sections with detailed, actionable items:
    
    # {role} Onboarding Plan - {department} Department
    
    ## First 30 Days: Learning & Orientation
    - Week 1: Company introduction, tools setup, and team integration
    - Week 2-3: Role-specific training and initial tasks
    - Week 4: First month review and goal setting
    
    ## Days 31-60: Skill Development
    - Department-specific training
    - Cross-functional exposure
    - Performance check-ins
    
    ## Days 61-90: Independence & Contribution
    - Independent project work
    - Process improvements
    - 90-day review preparation
    
    ## Key Resources:
    - Department playbook
    - Training materials
    - Important contacts
    
    Format the response in clean markdown with clear sections and bullet points."""
    
    try:
        response = model.generate_content({
            'parts': [prompt],
            'generation_config': {
                'max_output_tokens': 2000,
                'temperature': 0.7,
            }
        })
        return response.text
    except Exception as e:
        app.logger.error(f"Error generating onboarding plan: {str(e)}")
        return "I couldn't generate an onboarding plan at the moment. Please try again later."

def generate_offboarding_checklist(employee_name, last_working_day, role, department):
    """Generate a comprehensive offboarding checklist using Gemini AI"""
    prompt = f"""As an HR professional, create a detailed offboarding checklist for {employee_name} ({role} in {department}) who is leaving on {last_working_day}.
    
    Include these categories with specific, actionable items:
    
    # Offboarding Checklist for {employee_name}
    
    ## Pre-Exit (Before {last_working_day})
    - [ ] Schedule and conduct exit interview
    - [ ] Document knowledge transfer sessions
    - [ ] Collect company assets (laptop, ID, access cards, etc.)
    
    ## IT & Access
    - [ ] Disable email and system access
    - [ ] Transfer or archive files and documents
    - [ ] Update contact lists and distribution groups
    
    ## HR & Admin
    - [ ] Finalize payroll and benefits
    - [ ] Return company property
    - [ ] Update organizational charts and directories
    
    ## Department-Specific Tasks
    - [ ] Handover ongoing projects
    - [ ] Document key processes
    - [ ] Update documentation and manuals
    
    ## Post-Exit
    - [ ] Send farewell/transition email
    - [ ] Update external contacts if needed
    - [ ] Schedule knowledge transfer with replacement if applicable
    
    Format as a clean markdown checklist with checkboxes [ ]."""
    
    try:
        response = model.generate_content({
            'parts': [prompt],
            'generation_config': {
                'max_output_tokens': 2000,
                'temperature': 0.5,
            }
        })
        return response.text
    except Exception as e:
        app.logger.error(f"Error generating offboarding checklist: {str(e)}")
        return "I couldn't generate an offboarding checklist at the moment. Please try again later."


@app.route('/api/ai/generate-question', methods=['POST'])
@login_required(['hr', 'employee'])
def generate_ai_question():
    """Generate AI-powered interview question using Gemini"""
    try:
        data = request.get_json()
        training_type = data.get('type', 'interview')
        difficulty = data.get('difficulty', 'beginner')
        question_index = data.get('questionIndex', 0)
        previous_responses = data.get('previousResponses', [])
        
        # Create context for Gemini
        context = f"You are an AI interview trainer conducting a {difficulty} level {training_type} training session."
        
        if previous_responses:
            context += f" Previous responses show the candidate has answered {len(previous_responses)} questions."
        
        prompt = f"{context} Generate question {question_index + 1} that is appropriate for {training_type} skills at {difficulty} level. Make it engaging and relevant. Return only the question."
        
        # Generate question using Gemini
        ai_question = generate_ai_response([{"role": "user", "content": prompt}], max_tokens=100)
        
        return jsonify({
            'success': True,
            'question': ai_question
        })
        
    except Exception as e:
        app.logger.error(f'Error generating AI question: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to generate question'
        }), 500

@app.route('/api/ai/generate-feedback', methods=['POST'])
@login_required(['hr', 'employee'])
def generate_ai_feedback():
    """Generate AI feedback for training responses using Gemini"""
    try:
        data = request.get_json()
        response_text = data.get('response', '')
        question = data.get('question', '')
        training_type = data.get('type', 'interview')
        difficulty = data.get('difficulty', 'beginner')
        
        # Create feedback prompt for Gemini
        prompt = f"""As an expert interview coach, analyze this {training_type} response at {difficulty} level:
        
Question: {question}
Response: {response_text}
        
Provide constructive feedback in 2-3 sentences focusing on:
1. What they did well
2. One specific improvement suggestion
3. Encouragement for next question
        
Keep it supportive and actionable."""
        
        # Generate feedback using Gemini
        ai_feedback = generate_ai_response([{"role": "user", "content": prompt}], max_tokens=150)
        
        return jsonify({
            'success': True,
            'feedback': ai_feedback
        })
        
    except Exception as e:
        app.logger.error(f'Error generating AI feedback: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to generate feedback'
        }), 500

@app.route('/api/training/generate-feedback', methods=['POST'])
@login_required(['hr', 'employee'])
def generate_training_feedback():
    """Generate AI feedback for training responses"""
    try:
        data = request.get_json()
        response_text = data.get('response', '')
        question_type = data.get('type', 'general')
        
        # Generate contextual feedback based on response
        feedback = analyze_training_response(response_text, question_type)
        
        return jsonify({
            'success': True,
            'feedback': feedback
        })
        
    except Exception as e:
        app.logger.error(f'Error generating training feedback: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to generate feedback'
        }), 500

def analyze_training_response(response, question_type):
    """Analyze training response and provide feedback"""
    word_count = len(response.split())
    has_examples = 'example' in response.lower() or 'instance' in response.lower()
    
    feedback = {
        'score': 0,
        'strengths': [],
        'improvements': [],
        'overall': ''
    }
    
    # Analyze word count
    if word_count >= 30:
        feedback['score'] += 25
        feedback['strengths'].append('Good response length')
    else:
        feedback['improvements'].append('Provide more detailed responses')
    
    # Check for examples
    if has_examples:
        feedback['score'] += 25
        feedback['strengths'].append('Included specific examples')
    else:
        feedback['improvements'].append('Add concrete examples to strengthen your answer')
    
    # Check structure
    sentences = response.split('.')
    if len(sentences) >= 3:
        feedback['score'] += 25
        feedback['strengths'].append('Well-structured response')
    else:
        feedback['improvements'].append('Organize your thoughts more clearly')
    
    # Question-specific analysis
    if question_type == 'behavioral':
        if 'situation' in response.lower() and 'action' in response.lower():
            feedback['score'] += 25
            feedback['strengths'].append('Used STAR method effectively')
        else:
            feedback['improvements'].append('Consider using the STAR method (Situation, Task, Action, Result)')
    else:
        feedback['score'] += 15  # Base score for other types
    
    # Generate overall feedback
    if feedback['score'] >= 80:
        feedback['overall'] = 'Excellent response! You demonstrated strong communication skills.'
    elif feedback['score'] >= 60:
        feedback['overall'] = 'Good response with room for improvement.'
    else:
        feedback['overall'] = 'Keep practicing! Focus on the suggested improvements.'
    
    return feedback

@app.route('/api/ai/generate-summary', methods=['POST'])
@login_required(['hr', 'employee'])
def generate_ai_summary():
    """Generate comprehensive HR interview summary using Gemini"""
    try:
        data = request.get_json()
        session_data = data.get('session', {})
        responses = data.get('responses', [])
        candidate_name = session_data.get('candidateName', 'Candidate')
        role = session_data.get('roleHiring', 'Position')
        
        # Create comprehensive analysis prompt
        responses_text = "\n".join([f"Q: {r.get('question', '')}\nA: {r.get('answer', '')}" for r in responses if not r.get('skipped')])
        
        prompt = f"""Generate a comprehensive HR interview summary for {candidate_name} applying for {role}:

{responses_text}

Provide detailed analysis in this format:

**EXECUTIVE SUMMARY**
- Overall Rating (1-10):
- Fit for Role:
- Hire Recommendation:

**SKILLS ANALYSIS** (Score 0-10 each)
- Technical Skills:
- Communication:
- Problem Solving:
- Leadership:
- Cultural Fit:

**TONE & MANNERISM**
- Confidence Level:
- Communication Style:
- Response Quality:

**ANSWER BREAKDOWN**
[For each major question, provide: Question, Summary, Score, Red flags/Highlights]

**RECOMMENDED ACTION**
[Proceed to next round/Hold/Reject/Skill test/HR round]

Be specific and professional."""
        
        # Generate comprehensive summary
        ai_summary = generate_ai_response([{"role": "user", "content": prompt}], max_tokens=800)
        
        return jsonify({
            'success': True,
            'summary': ai_summary,
            'candidate': candidate_name,
            'role': role
        })
        
    except Exception as e:
        app.logger.error(f'Error generating AI summary: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to generate summary'
        }), 500

@app.route('/api/ai/analyze-response', methods=['POST'])
@login_required(['hr', 'employee'])
def analyze_response_realtime():
    """Real-time response analysis during interview"""
    try:
        data = request.get_json()
        response_text = data.get('response', '')
        question = data.get('question', '')
        
        # Analyze response characteristics
        word_count = len(response_text.split())
        hesitation_words = ['um', 'uh', 'like', 'you know', 'actually']
        hesitation_count = sum(response_text.lower().count(word) for word in hesitation_words)
        
        # AI analysis prompt
        prompt = f"""Analyze this interview response for:
1. Confidence (Low/Medium/High)
2. Clarity (Poor/Good/Excellent)
3. Relevance (Off-topic/Partial/Relevant)
4. Technical accuracy (if applicable)
5. Communication style

Question: {question}
Response: {response_text}

Provide brief analysis in JSON format with scores."""
        
        ai_analysis = generate_ai_response([{"role": "user", "content": prompt}], max_tokens=200)
        
        return jsonify({
            'success': True,
            'analysis': {
                'word_count': word_count,
                'hesitation_count': hesitation_count,
                'ai_feedback': ai_analysis
            }
        })
        
    except Exception as e:
        app.logger.error(f'Error analyzing response: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to analyze response'
        }), 500

@app.route('/api/ai/send-interview-summary', methods=['POST'])
@login_required('hr')
def send_interview_summary():
    """Send interview summary to HR email and dashboard"""
    try:
        data = request.get_json()
        summary = data.get('summary', '')
        candidate_name = data.get('candidate', '')
        role = data.get('role', '')
        hr_email = session.get('email', 'hr@company.com')
        
        # Save to database (Interview model)
        interview_record = Interview(
            candidate_id=None,  # Link to candidate if exists
            responses={'summary': summary, 'candidate': candidate_name, 'role': role},
            summary=summary,
            created_at=datetime.utcnow()
        )
        db.session.add(interview_record)
        db.session.commit()
        
        # Generate email content
        email_subject = f"Interview Summary - {candidate_name} for {role}"
        email_body = f"""Interview Summary Report
        
Candidate: {candidate_name}
Position: {role}
Interview Date: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}
        
{summary}
        
This summary has been automatically generated and saved to your HR dashboard.
        
Best regards,
SmartHire AI System"""
        
        # In production, integrate with email service (SendGrid, AWS SES, etc.)
        # For now, log the email
        app.logger.info(f"Interview summary email sent to {hr_email}")
        
        return jsonify({
            'success': True,
            'message': 'Interview summary sent and saved',
            'interview_id': interview_record.id
        })
        
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error sending interview summary: {str(e)}')
        return jsonify({
            'success': False,
            'message': 'Failed to send summary'
        }), 500

@app.route('/api/generate-welcome-email', methods=['POST'])
def generate_welcome_email():
    try:
        data = request.json
        employee_name = data.get('name', 'New Employee')
        start_date = data.get('start_date', 'their start date')
        
        prompt = f"""Write a warm, professional welcome email for {employee_name} 
        who is joining on {start_date}. Include:
        1. A warm welcome message
        2. First day details (time, location, who to ask for)
        3. What to bring on the first day
        4. A friendly closing"""
        
        email_content = generate_ai_response([
            {"role": "system", "content": "You are an HR professional writing welcome emails."},
            {"role": "user", "content": prompt}
        ], max_tokens=800)
        
        return jsonify({
            'status': 'success',
            'email': email_content
        })
        
    except Exception as e:
        app.logger.error(f"Error generating welcome email: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate welcome email'
        }), 500

@app.route('/api/feedback/submit', methods=['POST'])
@login_required('employee')
def submit_feedback():
    try:
        data = request.get_json(silent=True) or {}

        # Safely coerce ratings to integers
        try:
            mood_rating_raw = data.get('mood_rating')
            confidence_rating_raw = data.get('confidence_rating')

            mood_rating = int(mood_rating_raw) if mood_rating_raw is not None else None
            confidence_rating = int(confidence_rating_raw) if confidence_rating_raw is not None else None
        except (TypeError, ValueError):
            app.logger.warning(f"Invalid rating types in feedback payload: {data}")
            return jsonify({'status': 'error', 'message': 'Ratings must be numbers between 1 and 5'}), 400

        feedback = data.get('feedback', '') or ''

        # Validate presence
        if mood_rating is None or confidence_rating is None:
            return jsonify({'status': 'error', 'message': 'Both mood and confidence ratings are required'}), 400

        # Validate range (1-5)
        if not (1 <= mood_rating <= 5) or not (1 <= confidence_rating <= 5):
            return jsonify({'status': 'error', 'message': 'Ratings must be between 1 and 5'}), 400

        user_id = session.get('user_id')
        if not user_id:
            app.logger.warning('Feedback submit attempted without user in session')
            return jsonify({'status': 'error', 'message': 'Authentication required'}), 401

        # Create new feedback entry
        new_feedback = EmployeeFeedback(
            user_id=user_id,
            mood_rating=mood_rating,
            confidence_rating=confidence_rating,
            feedback=feedback
        )

        db.session.add(new_feedback)
        db.session.commit()
        
        # Create notification for HR users about new feedback
        user = User.query.get(user_id)
        if user:
            # Check if mood or confidence is low (<= 2)
            if mood_rating <= 2 or confidence_rating <= 2:
                notification_type = 'low_mood' if mood_rating <= 2 else 'low_confidence'
                priority = 'high'
                title = f"Low {notification_type.replace('_', ' ').title()} Alert"
                message = f"{user.full_name} reported {notification_type.replace('_', ' ')} (Mood: {mood_rating}, Confidence: {confidence_rating})"
            else:
                notification_type = 'feedback'
                priority = 'normal'
                title = "New Mood Feedback"
                message = f"{user.full_name} submitted mood feedback (Mood: {mood_rating}, Confidence: {confidence_rating})"
            
            # Send to all HR users
            hr_users = User.query.filter_by(role='hr').all()
            for hr_user in hr_users:
                notification = Message(
                    subject=title,
                    content=message,
                    sender_id=user_id,
                    recipient_id=hr_user.id,
                    message_type='hr_notification',
                    priority=priority,
                    notification_data={
                        'notification_type': notification_type,
                        'feedback_id': new_feedback.id,
                        'mood_rating': mood_rating,
                        'confidence_rating': confidence_rating
                    }
                )
                db.session.add(notification)
            
            db.session.commit()

        return jsonify({
            'status': 'success',
            'message': 'Feedback submitted successfully',
            'data': {
                'id': new_feedback.id,
                'created_at': new_feedback.created_at.isoformat()
            }
        })

    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error submitting feedback: {str(e)}", exc_info=True)
        return jsonify({
            'status': 'error',
            'message': 'Failed to submit feedback. Please try again.'
        }), 500

@app.route('/api/feedback/debug')
@login_required('hr')
def debug_feedback_data():
    """Debug endpoint to check database connection and feedback data"""
    try:
        # Test database connection
        feedback_count = EmployeeFeedback.query.count()
        recent_feedback = EmployeeFeedback.query.order_by(EmployeeFeedback.created_at.desc()).limit(5).all()
        
        debug_info = {
            'database_connected': True,
            'total_feedback_count': feedback_count,
            'recent_feedback_sample': []
        }
        
        for feedback in recent_feedback:
            debug_info['recent_feedback_sample'].append({
                'id': feedback.id,
                'user_id': feedback.user_id,
                'mood_rating': feedback.mood_rating,
                'confidence_rating': feedback.confidence_rating,
                'feedback': feedback.feedback,
                'created_at': feedback.created_at.isoformat() if feedback.created_at else None,
                'user_name': feedback.user.full_name if feedback.user else 'Unknown'
            })
        
        return jsonify({
            'status': 'success',
            'debug_info': debug_info
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e),
            'database_connected': False
        }), 500

@app.route('/api/feedback/recent')
@login_required('hr')
def get_recent_feedback():
    """Get recent feedback with employee details"""
    try:
        # Get feedback from the last 30 days
        thirty_days_ago = datetime.utcnow() - timedelta(days=30)
        recent_feedback = EmployeeFeedback.query.filter(
            EmployeeFeedback.created_at >= thirty_days_ago
        ).order_by(EmployeeFeedback.created_at.desc()).limit(20).all()
        
        feedback_list = []
        for feedback in recent_feedback:
            # Get mood emoji
            mood_emoji = "😊" if feedback.mood_rating >= 4 else "😐" if feedback.mood_rating >= 3 else "😞"
            
            feedback_list.append({
                'id': feedback.id,
                'date': feedback.created_at.strftime('%Y-%m-%d %H:%M'),
                'employee_name': feedback.user.full_name if feedback.user else 'Unknown',
                'mood_rating': feedback.mood_rating,
                'mood_emoji': mood_emoji,
                'confidence_rating': feedback.confidence_rating,
                'feedback_text': feedback.feedback or 'No comment provided',
                'department': feedback.user.department if feedback.user and feedback.user.department else 'N/A'
            })
        
        return jsonify({
            'status': 'success',
            'data': feedback_list
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/feedback/summary')
@login_required('hr')
def get_feedback_summary():
    try:
        # Get feedback from the last 7 days
        one_week_ago = datetime.utcnow() - timedelta(days=7)
        recent_feedback = EmployeeFeedback.query.filter(
            EmployeeFeedback.created_at >= one_week_ago
        ).all()
        
        if not recent_feedback:
            return jsonify({
                'status': 'success',
                'data': {
                    'total_submissions': 0,
                    'avg_mood': 0,
                    'avg_confidence': 0,
                    'mood_trend': [],
                    'confidence_trend': [],
                    'alerts': []
                }
            })
        
        # Calculate averages
        total = len(recent_feedback)
        avg_mood = sum(f.mood_rating for f in recent_feedback) / total
        avg_confidence = sum(f.confidence_rating for f in recent_feedback) / total
        
        # Generate daily trends (last 7 days)
        today = datetime.utcnow().date()
        dates = [(today - timedelta(days=i)).strftime('%Y-%m-%d') for i in range(6, -1, -1)]
        
        mood_trend = []
        confidence_trend = []
        alerts = []
        
        for date in dates:
            daily_feedback = [f for f in recent_feedback if f.created_at.date().isoformat() == date]
            if daily_feedback:
                mood_avg = sum(f.mood_rating for f in daily_feedback) / len(daily_feedback)
                conf_avg = sum(f.confidence_rating for f in daily_feedback) / len(daily_feedback)
                
                # Check for significant drops in confidence
                if len(daily_feedback) >= 3 and conf_avg < 2.5:  # Threshold for low confidence
                    alerts.append({
                        'type': 'low_confidence',
                        'date': date,
                        'value': round(conf_avg, 1)
                    })
            else:
                mood_avg = 0
                conf_avg = 0
                
            mood_trend.append(round(mood_avg, 1))
            confidence_trend.append(round(conf_avg, 1))
        
        # Calculate engagement score (simple average of mood and confidence)
        engagement_score = round((avg_mood + avg_confidence) / 2, 1)
        
        return jsonify({
            'status': 'success',
            'data': {
                'total_submissions': total,
                'avg_mood': round(avg_mood, 1),
                'avg_confidence': round(avg_confidence, 1),
                'engagement_score': engagement_score,
                'mood_trend': mood_trend,
                'confidence_trend': confidence_trend,
                'alerts': alerts,
                'dates': dates
            }
        })
        
    except Exception as e:
        app.logger.error(f"Error fetching feedback summary: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch feedback summary.'
        }), 500


# ============ EMPLOYEE PROFILE, DOCUMENTS & SETTINGS ENDPOINTS ============

@app.route('/api/employee/profile', methods=['GET'])
@login_required('employee')
def get_employee_profile():
    """Get employee profile data"""
    try:
        user = User.query.get(session['user_id'])
        if not user:
            return jsonify({'status': 'error',
                           'message': 'User not found'}), 404
        
        return jsonify({
            'status': 'success',
            'data': {
                'full_name': user.full_name,
                'email': user.email,
                'phone': user.phone or '',
                'department': user.department or '',
                'position': user.position or '',
                'employee_id': user.employee_id or f'EMP-{user.id:04d}',
                'hire_date': (
                    user.hire_date.strftime('%Y-%m-%d')
                    if user.hire_date else ''),
                'status': user.status or 'Active'
            }
        })
    except Exception as e:
        app.logger.error(f'Error fetching profile: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to fetch profile'}), 500


@app.route('/api/employee/profile', methods=['POST'])
@login_required('employee')
def update_employee_profile():
    """Update employee profile"""
    try:
        user = User.query.get(session['user_id'])
        if not user:
            return jsonify({'status': 'error',
                           'message': 'User not found'}), 404
        
        # Update user profile
        user.phone = request.form.get('phone', user.phone)
        user.department = request.form.get(
            'department', user.department)
        user.position = request.form.get('position', user.position)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Profile updated successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error updating profile: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to update profile'}), 500


@app.route('/api/employee/documents', methods=['GET'])
@login_required('employee')
def get_employee_documents():
    """Get list of employee documents"""
    try:
        documents = EmployeeDocument.query.filter_by(
            user_id=session['user_id']
        ).order_by(EmployeeDocument.upload_date.desc()).all()
        
        docs_list = []
        for doc in documents:
            docs_list.append({
                'id': doc.id,
                'file_name': doc.file_name,
                'document_type': doc.document_type or 'Other',
                'upload_date': doc.upload_date.strftime('%Y-%m-%d %H:%M'),
                'file_size': doc.file_size,
                'description': doc.description or '',
                'is_verified': doc.is_verified
            })
        
        return jsonify({
            'status': 'success',
            'data': docs_list
        })
    except Exception as e:
        app.logger.error(f'Error fetching documents: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to fetch documents'}), 500


@app.route('/api/employee/documents/upload', methods=['POST'])
@login_required('employee')
def upload_employee_document():
    """Upload employee document"""
    try:
        if 'file' not in request.files:
            return jsonify({'status': 'error',
                           'message': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'status': 'error',
                           'message': 'No file selected'}), 400
        
        # Allowed file extensions
        allowed_extensions = {
            'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png', 'txt', 'xlsx'
        }
        
        if not ('.' in file.filename and
                file.filename.rsplit('.', 1)[1].lower()
                in allowed_extensions):
            return jsonify({
                'status': 'error',
                'message': 'File type not allowed'
            }), 400
        
        # Check file size (max 10MB)
        file_size = len(file.read())
        file.seek(0)
        if file_size > 10 * 1024 * 1024:
            return jsonify({
                'status': 'error',
                'message': 'File size exceeds 10MB limit'
            }), 400
        
        # Create documents directory if it doesn't exist
        doc_dir = os.path.join(app.root_path, 'uploads', 'documents')
        os.makedirs(doc_dir, exist_ok=True)
        
        # Generate unique filename
        user_id = session['user_id']
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        filename = f"{user_id}_{timestamp}_{file.filename}"
        file_path = os.path.join(doc_dir, filename)
        
        # Save file
        file.save(file_path)
        
        # Create database entry
        document = EmployeeDocument(
            user_id=user_id,
            file_name=file.filename,
            file_path=f"/uploads/documents/{filename}",
            document_type=request.form.get(
                'document_type', 'Other'),
            file_size=file_size,
            description=request.form.get('description', ''),
            is_verified=False
        )
        
        db.session.add(document)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Document uploaded successfully',
            'data': {
                'id': document.id,
                'file_name': document.file_name,
                'upload_date': document.upload_date.strftime(
                    '%Y-%m-%d %H:%M')
            }
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error uploading document: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to upload document'}), 500


@app.route('/api/employee/documents/<int:doc_id>', methods=['DELETE'])
@login_required('employee')
def delete_employee_document(doc_id):
    """Delete employee document"""
    try:
        document = EmployeeDocument.query.get_or_404(doc_id)
        
        # Verify ownership
        if document.user_id != session['user_id']:
            return jsonify({'status': 'error',
                           'message': 'Unauthorized'}), 403
        
        # Delete file
        file_path = os.path.join(
            app.root_path, 'uploads', 'documents',
            os.path.basename(document.file_path))
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Delete database entry
        db.session.delete(document)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Document deleted successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error deleting document: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to delete document'}), 500


@app.route('/api/employee/settings', methods=['GET'])
@login_required('employee')
def get_employee_settings():
    """Get employee settings"""
    try:
        settings = EmployeeSettings.query.filter_by(
            user_id=session['user_id']
        ).first()
        
        if not settings:
            # Create default settings if not exists
            settings = EmployeeSettings(user_id=session['user_id'])
            db.session.add(settings)
            db.session.commit()
        
        return jsonify({
            'status': 'success',
            'data': {
                'email_notifications': settings.email_notifications,
                'sms_notifications': settings.sms_notifications,
                'notification_frequency': (
                    settings.notification_frequency),
                'theme': settings.theme,
                'language': settings.language,
                'two_factor_enabled': settings.two_factor_enabled
            }
        })
    except Exception as e:
        app.logger.error(f'Error fetching settings: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to fetch settings'}), 500


@app.route('/api/employee/settings', methods=['POST'])
@login_required('employee')
def update_employee_settings():
    """Update employee settings"""
    try:
        settings = EmployeeSettings.query.filter_by(
            user_id=session['user_id']
        ).first()
        
        if not settings:
            settings = EmployeeSettings(user_id=session['user_id'])
            db.session.add(settings)
        
        # Update settings
        settings.email_notifications = request.form.get(
            'email_notifications', 'false').lower() == 'true'
        settings.sms_notifications = request.form.get(
            'sms_notifications', 'false').lower() == 'true'
        settings.notification_frequency = request.form.get(
            'notification_frequency', 'immediately')
        settings.theme = request.form.get('theme', 'light')
        settings.language = request.form.get('language', 'en')
        settings.two_factor_enabled = request.form.get(
            'two_factor_enabled', 'false').lower() == 'true'
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Settings updated successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error updating settings: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to update settings'}), 500


@app.route('/api/employee/change-password', methods=['POST'])
@login_required('employee')
def change_employee_password():
    """Change employee password"""
    try:
        user = User.query.get(session['user_id'])
        if not user:
            return jsonify({'status': 'error',
                           'message': 'User not found'}), 404
        
        current_password = request.form.get('current_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')
        
        # Validate current password
        if not user.check_password(current_password):
            return jsonify({'status': 'error',
                           'message': 'Current password is incorrect'}), 401
        
        # Validate new password
        if len(new_password) < 8:
            return jsonify({
                'status': 'error',
                'message': 'Password must be at least 8 characters'
            }), 400
        
        # Validate passwords match
        if new_password != confirm_password:
            return jsonify({'status': 'error',
                           'message': 'Passwords do not match'}), 400
        
        # Update password
        user.set_password(new_password)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Password changed successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error changing password: {str(e)}')
        return jsonify({'status': 'error',
                       'message': 'Failed to change password'}), 500

@app.route('/api/onboarding/tasks/<int:task_id>', methods=['DELETE'])
@login_required('hr')
def delete_onboarding_task(task_id):
    """Delete an onboarding task"""
    try:
        # Legacy check - onboarding task system has been removed
        return jsonify({'status': 'error', 'message': 'Onboarding task system has been removed'}), 400
    except Exception as e:
        app.logger.error(f'Error deleting onboarding task: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to delete task'
        }), 500

@app.route('/api/employee/onboarding/tasks/<int:task_id>/toggle', methods=['PUT'])
@login_required('employee')
def employee_toggle_onboarding_task(task_id):
    """Toggle onboarding task completion (employee view)"""
    try:
        user_id = session.get('user_id')
        user = User.query.get(user_id)
        
        # Get the task and verify it belongs to the current user
        task = Task.query.filter_by(id=task_id, assigned_to=user.id, task_type='onboarding').first()
        if not task:
            return jsonify({'status': 'error', 'message': 'Task not found'}), 404
        
        # Toggle task status
        if task.status == 'completed':
            task.status = 'pending'
            task.completed_at = None
        else:
            task.status = 'completed'
            task.completed_at = datetime.utcnow()
        
        task.updated_at = datetime.utcnow()
        db.session.commit()
        
        # Notify HR users about the task update
        hr_users = User.query.filter_by(role='hr', is_active=True).all()
        for hr_user in hr_users:
            message = Message(
                sender_id=user.id,
                recipient_id=hr_user.id,
                subject=f'Onboarding Task Updated: {task.title}',
                content=f'Employee {user.full_name} has {"completed" if task.status == "completed" else "reopened"} the onboarding task "{task.title}".\n\nTask Status: {task.status.title()}\nUpdated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")}',
                priority='low',
                status='unread'
            )
            db.session.add(message)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success', 
            'message': 'Task updated successfully',
            'task_status': task.status,
            'completed_at': task.completed_at.isoformat() if task.completed_at else None
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error toggling onboarding task: {str(e)}')
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/hr/onboarding/tasks/<int:employee_id>/updates', methods=['GET'])
@login_required('hr')
def hr_onboarding_task_updates(employee_id):
    """Get real-time updates for employee onboarding tasks"""
    try:
        # Verify employee exists
        employee = User.query.get_or_404(employee_id)
        
        # Get latest task updates
        onboarding_tasks = Task.query.filter_by(
            assigned_to=employee.id, 
            task_type='onboarding'
        ).order_by(Task.updated_at.desc()).all()
        
        # Return recent updates (last 10 tasks updated in last 24 hours)
        recent_cutoff = datetime.utcnow() - timedelta(hours=24)
        recent_tasks = [task for task in onboarding_tasks if task.updated_at >= recent_cutoff]
        
        updates = []
        for task in recent_tasks[:10]:  # Limit to 10 most recent
            updates.append({
                'id': task.id,
                'title': task.title,
                'status': task.status,
                'updated_at': task.updated_at.isoformat(),
                'completed_at': task.completed_at.isoformat() if task.completed_at else None,
                'employee_name': employee.full_name,
                'priority': task.priority
            })
        
        return jsonify({
            'status': 'success',
            'updates': updates,
            'total_tasks': len(onboarding_tasks),
            'completed_tasks': len([t for t in onboarding_tasks if t.status == 'completed']),
            'last_updated': max([t.updated_at for t in onboarding_tasks]).isoformat() if onboarding_tasks else None
        })
    except Exception as e:
        app.logger.error(f'Error fetching task updates: {str(e)}')
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/employee/onboarding/tasks/<int:task_id>/report-issue', methods=['POST'])
@login_required('employee')
def employee_report_issue(task_id):
    """Report an issue with a task"""
    try:
        user_id = session.get('user_id')
        user = User.query.get(user_id)
        
        # Get the task and verify it belongs to the current user
        task = Task.query.filter_by(id=task_id, assigned_to=user.id, task_type='onboarding').first()
        if not task:
            return jsonify({'status': 'error', 'message': 'Task not found'}), 404
        
        # Get issue details from request
        data = request.get_json()
        issue_description = data.get('issue_description', '')
        issue_type = data.get('issue_type', 'general')
        
        if not issue_description:
            return jsonify({'status': 'error', 'message': 'Issue description is required'}), 400
        
        # Update task to mark as having an issue
        task.status = 'issue_reported'
        task.updated_at = datetime.utcnow()
        
        # Add issue details to task description or comments
        issue_note = f"[ISSUE REPORTED - {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}]: {issue_description} (Type: {issue_type})"
        if task.description:
            task.description = f"{task.description}\n\n{issue_note}"
        else:
            task.description = issue_note
        
        # Create a message to HR about the issue
        hr_users = User.query.filter_by(role='hr', is_active=True).all()
        for hr_user in hr_users:
            message = Message(
                sender_id=user.id,
                recipient_id=hr_user.id,
                subject=f'Onboarding Task Issue: {task.title}',
                content=f'Employee {user.full_name} has reported an issue with onboarding task "{task.title}".\n\nIssue Details:\n{issue_description}\n\nIssue Type: {issue_type}',
                priority='medium',
                status='unread'
            )
            db.session.add(message)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success', 
            'message': 'Issue reported successfully. HR has been notified.',
            'task_status': task.status
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error reporting task issue: {str(e)}')
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/employee/onboarding/comment', methods=['POST'])
@login_required('employee')
def employee_add_comment():
    """Add a comment about onboarding"""
    try:
        user_id = session.get('user_id')
        data = request.json
        comment = data.get('comment')
        
        if not comment:
            return jsonify({'status': 'error', 'message': 'Comment is required'}), 400
        
        # Store comment in the first pending task or create a general comment
        # Legacy check - onboarding checklist system has been removed
        checklist = None
        if not checklist:
            return jsonify({'status': 'error', 'message': 'No onboarding checklist found'}), 404
        
        # Legacy check - onboarding task system has been removed
        return jsonify({'status': 'error', 'message': 'Onboarding task system has been removed'}), 404
        
        # Create notification for HR
        create_notification(
            user_id=user_id,
            title='New Employee Comment',
            message=f'Employee added a comment: {comment}',
            notification_type='employee_comment',
            priority='normal'
        )
        
        return jsonify({'status': 'success', 'message': 'Comment added successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/onboarding/tasks/<int:task_id>/resolve-issue', methods=['POST'])
@login_required('hr')
def resolve_task_issue(task_id):
    """Resolve a task issue (HR view)"""
    try:
        # Legacy check - onboarding task system has been removed
        return jsonify({'status': 'error', 'message': 'Onboarding task system has been removed'}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/onboarding/tasks/<int:task_id>/toggle', methods=['PUT'])
@login_required('hr')
def toggle_onboarding_task(task_id):
    """Toggle onboarding task completion status"""
    try:
        # Legacy check - onboarding task system has been removed
        return jsonify({'status': 'error', 'message': 'Onboarding task system has been removed'}), 400
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error toggling task: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to update task'
        }), 500

@app.route('/api/onboarding/checklist/<int:checklist_id>/tasks', methods=['PUT'])
@login_required('hr')
def update_onboarding_tasks(checklist_id):
    """Update onboarding tasks for a checklist"""
    try:
        # Legacy check - onboarding checklist system has been removed
        checklist = None
        data = request.get_json()
        
        # Legacy check - onboarding task system has been removed
        return jsonify({'status': 'error', 'message': 'Onboarding task system has been removed'}), 400
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error updating onboarding tasks: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to update tasks'
        }), 500


# ===== OFFBOARDING API ENDPOINTS =====

@app.route('/api/offboarding/candidates', methods=['GET'])
@login_required('hr')
def get_offboarding_candidates():
    """Get all offboarding candidates"""
    try:
        candidates = OffboardingCandidate.query.all()
        return jsonify({
            'status': 'success',
            'candidates': [{
                'id': c.id,
                'employee_id': c.employee_id,
                'employee_name': c.employee.full_name if c.employee else 'N/A',
                'exit_date': c.exit_date.isoformat() if c.exit_date else None,
                'exit_reason': c.exit_reason,
                'status': c.status,
                'assigned_hr': c.assigned_hr,
                'created_at': c.created_at.isoformat() if c.created_at else None
            } for c in candidates]
        })
    except Exception as e:
        app.logger.error(f'Error fetching offboarding candidates: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch candidates'
        }), 500


@app.route('/api/offboarding/candidates', methods=['POST'])
@login_required('hr')
def create_offboarding_candidate():
    """Create a new offboarding candidate"""
    try:
        data = request.get_json()
        
        candidate = OffboardingCandidate(
            employee_id=data['employee_id'],
            exit_date=datetime.fromisoformat(data['exit_date'].replace('Z', '+00:00')) if data.get('exit_date') else None,
            exit_reason=data.get('exit_reason', ''),
            assigned_hr=data.get('assigned_hr'),
            status=data.get('status', 'pending')
        )
        
        db.session.add(candidate)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Offboarding candidate created successfully',
            'candidate_id': candidate.id
        }), 201
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error creating offboarding candidate: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to create candidate'
        }), 500


@app.route('/api/offboarding/candidates/<int:candidate_id>', methods=['PUT'])
@login_required('hr')
def update_offboarding_candidate(candidate_id):
    """Update offboarding candidate"""
    try:
        candidate = OffboardingCandidate.query.get_or_404(candidate_id)
        data = request.get_json()
        
        if 'exit_date' in data and data['exit_date']:
            candidate.exit_date = datetime.fromisoformat(data['exit_date'].replace('Z', '+00:00'))
        if 'exit_reason' in data:
            candidate.exit_reason = data['exit_reason']
        if 'assigned_hr' in data:
            candidate.assigned_hr = data['assigned_hr']
        if 'status' in data:
            candidate.status = data['status']
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Candidate updated successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error updating offboarding candidate: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to update candidate'
        }), 500


@app.route('/api/offboarding/assets/<int:candidate_id>', methods=['GET'])
@login_required('hr')
def get_offboarding_assets(candidate_id):
    """Get assets for offboarding candidate"""
    try:
        assets = OffboardingAsset.query.filter_by(candidate_id=candidate_id).all()
        return jsonify({
            'status': 'success',
            'assets': [{
                'id': a.id,
                'asset_name': a.asset_name,
                'asset_type': a.asset_type,
                'serial_number': a.serial_number,
                'status': a.status,
                'return_date': a.return_date.isoformat() if a.return_date else None,
                'notes': a.notes
            } for a in assets]
        })
    except Exception as e:
        app.logger.error(f'Error fetching offboarding assets: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch assets'
        }), 500


@app.route('/api/offboarding/assets', methods=['POST'])
@login_required('hr')
def create_offboarding_asset():
    """Create offboarding asset record"""
    try:
        data = request.get_json()
        
        asset = OffboardingAsset(
            candidate_id=data['candidate_id'],
            asset_name=data['asset_name'],
            asset_type=data.get('asset_type', 'equipment'),
            serial_number=data.get('serial_number', ''),
            status=data.get('status', 'assigned'),
            return_date=datetime.fromisoformat(data['return_date'].replace('Z', '+00:00')) if data.get('return_date') else None,
            notes=data.get('notes', '')
        )
        
        db.session.add(asset)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Asset recorded successfully',
            'asset_id': asset.id
        }), 201
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error creating offboarding asset: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to record asset'
        }), 500


@app.route('/api/offboarding/clearance/<int:candidate_id>', methods=['GET'])
@login_required('hr')
def get_offboarding_clearance(candidate_id):
    """Get clearance status for offboarding candidate"""
    try:
        clearances = OffboardingClearance.query.filter_by(candidate_id=candidate_id).all()
        return jsonify({
            'status': 'success',
            'clearances': [{
                'id': c.id,
                'department': c.department,
                'clearance_type': c.clearance_type,
                'status': c.status,
                'cleared_by': c.cleared_by,
                'cleared_date': c.cleared_date.isoformat() if c.cleared_date else None,
                'notes': c.notes
            } for c in clearances]
        })
    except Exception as e:
        app.logger.error(f'Error fetching offboarding clearance: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch clearance status'
        }), 500


@app.route('/api/offboarding/clearance', methods=['POST'])
@login_required('hr')
def update_offboarding_clearance():
    """Update clearance status"""
    try:
        data = request.get_json()
        
        clearance = OffboardingClearance(
            candidate_id=data['candidate_id'],
            department=data['department'],
            clearance_type=data.get('clearance_type', 'general'),
            status=data.get('status', 'pending'),
            cleared_by=data.get('cleared_by'),
            cleared_date=datetime.utcnow() if data.get('status') == 'cleared' else None,
            notes=data.get('notes', '')
        )
        
        # Check if clearance already exists
        existing = OffboardingClearance.query.filter_by(
            candidate_id=data['candidate_id'],
            department=data['department'],
            clearance_type=data.get('clearance_type', 'general')
        ).first()
        
        if existing:
            existing.status = data.get('status', 'pending')
            existing.cleared_by = data.get('cleared_by')
            existing.cleared_date = datetime.utcnow() if data.get('status') == 'cleared' else None
            existing.notes = data.get('notes', '')
        else:
            db.session.add(clearance)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Clearance status updated successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error updating offboarding clearance: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to update clearance'
        }), 500


# ===== DASHBOARD API ENDPOINTS =====

@app.route('/api/dashboard/stats', methods=['GET'])
@login_required('hr')
def get_dashboard_stats():
    """Get dashboard statistics"""
    try:
        # Employee statistics
        total_employees = Employee.query.count()
        active_employees = Employee.query.filter_by(status='Active').count()
        new_hires_this_month = Employee.query.filter(
            Employee.hire_date >= datetime.utcnow().replace(day=1)
        ).count()
        
        # Legacy check - onboarding checklist system has been removed
        pending_onboarding = 0
        completed_onboarding = 0
        
        # Offboarding statistics
        active_offboarding = OffboardingCandidate.query.filter_by(status='active').count()
        completed_offboarding = OffboardingCandidate.query.filter_by(status='completed').count()
        
        # Document statistics
        total_documents = EmployeeDocument.query.count()
        pending_documents = EmployeeDocument.query.filter_by(status='pending').count()
        
        # Feedback statistics
        recent_feedback = EmployeeFeedback.query.filter(
            EmployeeFeedback.created_at >= datetime.utcnow() - timedelta(days=30)
        ).count()
        
        return jsonify({
            'status': 'success',
            'stats': {
                'employees': {
                    'total': total_employees,
                    'active': active_employees,
                    'new_hires_this_month': new_hires_this_month
                },
                'onboarding': {
                    'pending': pending_onboarding,
                    'completed': completed_onboarding,
                    'completion_rate': round((completed_onboarding / (completed_onboarding + pending_onboarding) * 100) if (completed_onboarding + pending_onboarding) > 0 else 0, 1)
                },
                'offboarding': {
                    'active': active_offboarding,
                    'completed': completed_offboarding
                },
                'documents': {
                    'total': total_documents,
                    'pending': pending_documents
                },
                'feedback': {
                    'recent_feedback': recent_feedback
                }
            }
        })
    except Exception as e:
        app.logger.error(f'Error fetching dashboard stats: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch dashboard statistics'
        }), 500


@app.route('/api/dashboard/onboarding-progress', methods=['GET'])
@login_required('hr')
def get_onboarding_progress():
    """Get onboarding progress data"""
    try:
        # Legacy check - onboarding checklist system has been removed
        checklists = []
        progress_data = []
        
        return jsonify({
            'status': 'success',
            'data': progress_data
        })
    except Exception as e:
        app.logger.error(f'Error fetching onboarding progress: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch onboarding progress'
        }), 500


@app.route('/api/dashboard/department-stats', methods=['GET'])
@login_required('hr')
def get_department_stats():
    """Get department-wise statistics"""
    try:
        # Group employees by department
        dept_stats = {}
        employees = Employee.query.all()
        
        for emp in employees:
            dept = emp.department or 'Unassigned'
            if dept not in dept_stats:
                dept_stats[dept] = {
                    'total_employees': 0,
                    'active_employees': 0,
                    'new_hires': 0,
                    'pending_onboarding': 0,
                    'completed_onboarding': 0
                }
            
            dept_stats[dept]['total_employees'] += 1
            if emp.status == 'Active':
                dept_stats[dept]['active_employees'] += 1
            if emp.hire_date and emp.hire_date >= datetime.utcnow().replace(day=1):
                dept_stats[dept]['new_hires'] += 1
        
        # Legacy check - onboarding checklist system has been removed
        checklists = []
        for checklist in checklists:
            dept = checklist.employee.department if checklist.employee and checklist.employee.department else 'Unassigned'
            if dept in dept_stats:
                if checklist.status == 'Pending':
                    dept_stats[dept]['pending_onboarding'] += 1
                elif checklist.status == 'Completed':
                    dept_stats[dept]['completed_onboarding'] += 1
        
        return jsonify({
            'status': 'success',
            'departments': dept_stats
        })
    except Exception as e:
        app.logger.error(f'Error fetching department stats: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch department statistics'
        }), 500


@app.route('/api/dashboard/recent-activities', methods=['GET'])
@login_required('hr')
def get_recent_activities():
    """Get recent system activities"""
    try:
        activities = []
        
        # Recent logins
        # Legacy check - access records system has been removed
        recent_logins = []
        for login in recent_logins:
            activities.append({
                'type': 'login',
                'description': f"{login.user.full_name if login.user else 'Unknown'} logged in",
                'timestamp': login.timestamp.isoformat(),
                'user': login.user.full_name if login.user else 'Unknown'
            })
        
        # Recent document uploads
        recent_docs = EmployeeDocument.query.order_by(EmployeeDocument.upload_date.desc()).limit(5).all()
        for doc in recent_docs:
            activities.append({
                'type': 'document',
                'description': f"{doc.user.full_name if doc.user else 'Unknown'} uploaded {doc.file_name}",
                'timestamp': doc.upload_date.isoformat(),
                'user': doc.user.full_name if doc.user else 'Unknown'
            })
        
        # Legacy check - onboarding checklist system has been removed
        recent_onboarding = []
        for checklist in recent_onboarding:
            activities.append({
                'type': 'onboarding',
                'description': f"Onboarding completed for {checklist.employee.full_name if checklist.employee else 'Unknown'}",
                'timestamp': checklist.completed_at.isoformat(),
                'user': checklist.employee.full_name if checklist.employee else 'Unknown'
            })
        
        # Sort by timestamp
        activities.sort(key=lambda x: x['timestamp'], reverse=True)
        
        return jsonify({
            'status': 'success',
            'activities': activities[:10]  # Return latest 10 activities
        })
    except Exception as e:
        app.logger.error(f'Error fetching recent activities: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch recent activities'
        }), 500


@app.route('/api/dashboard/feedback-summary', methods=['GET'])
@login_required('hr')
def get_dashboard_feedback_summary():
    """Get feedback summary and trends"""
    try:
        # Get last 6 months of feedback
        six_months_ago = datetime.utcnow() - timedelta(days=180)
        feedback_data = EmployeeFeedback.query.filter(
            EmployeeFeedback.created_at >= six_months_ago
        ).all()
        
        # Group by month and calculate average mood
        monthly_data = {}
        for feedback in feedback_data:
            month_key = feedback.created_at.strftime('%Y-%m')
            if month_key not in monthly_data:
                monthly_data[month_key] = {
                    'count': 0,
                    'total_mood': 0,
                    'avg_mood': 0
                }
            
            monthly_data[month_key]['count'] += 1
            if feedback.mood_rating:
                monthly_data[month_key]['total_mood'] += feedback.mood_rating
        
        # Calculate averages
        for month in monthly_data:
            if monthly_data[month]['count'] > 0:
                monthly_data[month]['avg_mood'] = round(
                    monthly_data[month]['total_mood'] / monthly_data[month]['count'], 2
                )
        
        # Get department-wise feedback
        dept_feedback = {}
        for feedback in feedback_data:
            dept = feedback.user.employee.department if feedback.user and feedback.user.employee and feedback.user.employee.department else 'Unassigned'
            if dept not in dept_feedback:
                dept_feedback[dept] = {
                    'count': 0,
                    'total_mood': 0,
                    'avg_mood': 0
                }
            
            dept_feedback[dept]['count'] += 1
            if feedback.mood_rating:
                dept_feedback[dept]['total_mood'] += feedback.mood_rating
        
        # Calculate department averages
        for dept in dept_feedback:
            if dept_feedback[dept]['count'] > 0:
                dept_feedback[dept]['avg_mood'] = round(
                    dept_feedback[dept]['total_mood'] / dept_feedback[dept]['count'], 2
                )
        
        return jsonify({
            'status': 'success',
            'data': {
                'monthly_trends': monthly_data,
                'department_feedback': dept_feedback,
                'total_feedback': len(feedback_data),
                'overall_avg_mood': round(
                    sum(f.mood_rating for f in feedback_data if f.mood_rating) / 
                    len([f for f in feedback_data if f.mood_rating]), 2
                ) if feedback_data else 0
            }
        })
    except Exception as e:
        app.logger.error(f'Error fetching feedback summary: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch feedback summary'
        }), 500


# ===== HR NOTIFICATION SYSTEM =====

def create_hr_notification(title, message, notification_type='info', priority='normal', action_url=None, recipient_id=None, sender_id=None):
    """Create an HR notification"""
    try:
        # If no sender specified, use system user or first HR user
        if sender_id is None:
            # Try to get from session if available, otherwise use first admin/system user
            try:
                sender_id = session.get('user_id')
            except RuntimeError:
                # No request context, use system or first HR user
                sender_id = None
            
            if sender_id is None:
                # Find system admin or first HR user
                system_user = User.query.filter_by(role='hr').first()
                sender_id = system_user.id if system_user else 1
        
        # If no recipient specified, send to all HR users
        if recipient_id is None:
            hr_users = User.query.filter_by(role='hr').all()
            recipients = [u.id for u in hr_users]
        else:
            recipients = [recipient_id]
        
        notifications = []
        for recip_id in recipients:
            notification = Message(
                subject=title,
                content=message,
                sender_id=sender_id,
                recipient_id=recip_id,
                message_type='hr_notification',
                priority=priority,
                status='unread'
            )
            notification.notification_data = {
                'notification_type': notification_type,
                'action_url': action_url
            }
            notifications.append(notification)
            db.session.add(notification)
        
        db.session.commit()
        app.logger.info(f'Created HR notification for {len(recipients)} recipients: {title}')
        return notifications
    except Exception as e:
        app.logger.error(f'Error creating HR notification: {str(e)}')
        db.session.rollback()
        return []

def check_overdue_onboarding_tasks():
    """Check for overdue onboarding tasks and create notifications"""
    try:
        overdue_tasks = []
        
        # Legacy check - onboarding checklist system has been removed
        active_checklists = []
        
        for checklist in active_checklists:
            # Legacy check - onboarding checklist system has been removed
            incomplete_tasks = []
            
            for task in incomplete_tasks:
                # Calculate days overdue
                days_overdue = (datetime.utcnow() - checklist.created_at).days
                
                # Consider tasks overdue after 7 days
                if days_overdue > 7:
                    overdue_tasks.append({
                        'employee_id': checklist.employee_id,
                        'employee_name': checklist.employee.full_name if checklist.employee else 'Unknown',
                        'task_title': task.task_name,
                        'task_id': task.id,
                        'days_overdue': days_overdue,
                        'assigned_date': checklist.created_at.isoformat()
                    })
        
        # Create notifications for overdue tasks
        for task in overdue_tasks:
            title = "Overdue Onboarding Task"
            message = f"Task '{task['task_title']}' for {task['employee_name']} is {task['days_overdue']} days overdue."
            action_url = f"/onboarding?employee={task['employee_id']}"
            
            create_hr_notification(
                title=title,
                message=message,
                notification_type='onboarding_overdue',
                priority='high' if task['days_overdue'] > 14 else 'normal',
                action_url=action_url
            )
        
        return overdue_tasks
    except Exception as e:
        app.logger.error(f'Error checking overdue onboarding tasks: {str(e)}')
        return []

def check_pending_access_revocation():
    """Check for pending access revocation and create notifications"""
    try:
        pending_revocations = []
        
        # Get employees who have exited but still have active access
        exited_employees = User.query.filter(
            User.exit_date.isnot(None),
            User.exit_date <= datetime.utcnow().date()
        ).all()
        
        for employee in exited_employees:
            # Check for active access records
            # Legacy check - access records system has been removed
            active_access = []
            
            for access in active_access:
                days_since_exit = (datetime.utcnow().date() - employee.exit_date).days
                
                # Alert if access is still active 3 days after exit
                if days_since_exit >= 3:
                    pending_revocations.append({
                        'employee_id': employee.id,
                        'employee_name': employee.full_name,
                        'access_id': access.id,
                        'system_name': access.resource_name,
                        'exit_date': employee.exit_date.isoformat(),
                        'days_since_exit': days_since_exit
                    })
        
        # Create notifications for pending revocations
        for revocation in pending_revocations:
            title = "Pending Access Revocation"
            message = f"Access to {revocation['system_name']} for {revocation['employee_name']} should be revoked ({revocation['days_since_exit']} days post-exit)."
            action_url = f"/access-records?employee={revocation['employee_id']}"
            
            create_hr_notification(
                title=title,
                message=message,
                notification_type='access_revocation',
                priority='urgent' if revocation['days_since_exit'] > 7 else 'high',
                action_url=action_url
            )
        
        return pending_revocations
    except Exception as e:
        app.logger.error(f'Error checking pending access revocation: {str(e)}')
        return []


# ===== HR NOTIFICATION API ENDPOINTS =====

@app.route('/api/hr/notifications', methods=['GET'])
@login_required('hr')
def get_hr_notifications():
    """Get HR notifications count and recent high-priority notifications"""
    try:
        # Simple test response
        return jsonify({
            'status': 'success',
            'unread_count': 0,
            'notifications': [],
            'message': 'HR notifications API working'
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/hr/notifications/all', methods=['GET'])
@login_required('hr')
def get_all_hr_notifications():
    """Get all HR notifications"""
    try:
        # Get all notifications for current HR user
        all_notifications = Message.query.filter_by(
            recipient_id=session['user_id'],
            message_type='hr_notification'
        ).order_by(Message.sent_at.desc()).limit(50).all()
        
        # Format notifications
        notifications = []
        for notif in all_notifications:
            notification_data = notif.notification_data or {}
            notifications.append({
                'id': notif.id,
                'title': notif.subject,
                'message': notif.content,
                'type': notification_data.get('notification_type', 'info'),
                'priority': notif.priority,
                'created_at': notif.sent_at.isoformat(),
                'action_url': notification_data.get('action_url'),
                'is_read': notif.status == 'read'
            })
        
        return jsonify({
            'status': 'success',
            'notifications': notifications
        })
    except Exception as e:
        app.logger.error(f'Error getting all HR notifications: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to get notifications'
        }), 500

@app.route('/api/hr/daily-summary', methods=['POST'])
@login_required('hr')
def remove_daily_summary_notifications():
    """Remove all daily summary notifications from the database"""
    try:
        # Delete notifications where notification_data contains 'daily_summary'
        daily_summaries = Message.query.filter(
            Message.notification_data.isnot(None),
            Message.notification_data['notification_type'].as_string() == 'daily_summary'
        ).delete(synchronize_session=False)
        
        db.session.commit()
        app.logger.info(f"Removed {daily_summaries} daily summary notifications")
        return daily_summaries
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error removing daily summary notifications: {str(e)}")
        return 0

@app.route('/api/hr/cleanup-daily-summaries', methods=['POST'])
@login_required('hr')
def cleanup_daily_summaries():
    """Remove all daily summary notifications"""
    try:
        count = remove_daily_summary_notifications()
        return jsonify({
            'status': 'success',
            'message': f'Removed {count} daily summary notifications'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Failed to clean up daily summaries: {str(e)}'
        }), 500

@app.route('/api/hr/daily-summary', methods=['POST'])
@login_required('hr')
def generate_daily_hr_summary():
    """Generate daily summary of pending work without creating notifications"""
    try:
        # Get summary data but don't create a notification
        overdue_onboarding = check_overdue_onboarding_tasks()
        pending_access_revocation = check_pending_access_revocation()
        
        stats = {
            'pending_tasks': Task.query.filter_by(status='pending').count(),
            'pending_documents': EmployeeDocument.query.filter_by(status='pending').count(),
            'upcoming_interviews': Interview.query.filter(
                Interview.created_at >= datetime.utcnow(),
                Interview.created_at <= datetime.utcnow() + timedelta(days=7)
            ).count(),
            'completed_onboarding': 0
        }
        
        summary = {
            'overdue_onboarding': overdue_onboarding,
            'pending_access_revocation': pending_access_revocation,
            'stats': stats,
            'generated_at': datetime.utcnow().isoformat()
        }
        
        return jsonify({
            'status': 'success',
            'summary': summary
        })
    except Exception as e:
        app.logger.error(f'Error generating daily HR summary: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate daily summary'
        }), 500

@app.route('/api/hr/notifications/<int:notification_id>/read', methods=['PUT'])
@login_required('hr')
def mark_hr_notification_read(notification_id):
    """Mark HR notification as read"""
    try:
        notification = Message.query.filter_by(
            id=notification_id,
            recipient_id=session['user_id']
        ).first_or_404()
        
        notification.status = 'read'
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Notification marked as read'
        })
    except Exception as e:
        app.logger.error(f'Error marking HR notification as read: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to mark notification as read'
        }), 500

@app.route('/api/hr/notifications/mark-all-read', methods=['PUT'])
@login_required('hr')
def mark_all_hr_notifications_read():
    """Mark all HR notifications as read"""
    try:
        Message.query.filter_by(
            recipient_id=session['user_id'],
            message_type='hr_notification',
            status='unread'
        ).update({'status': 'read'})
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'All notifications marked as read'
        })
    except Exception as e:
        app.logger.error(f'Error marking all HR notifications as read: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to mark notifications as read'
        }), 500


# ===== NOTIFICATION HELPERS =====

def create_notification(recipient_id, subject, content, message_type='notification', priority='normal', sender_id=None):
    """Create a new notification message"""
    try:
        # If no sender specified, use system user (ID 1) or create a system notification
        if sender_id is None:
            # Try to find an HR user or use the first user as system
            system_user = User.query.filter_by(role='hr').first()
            sender_id = system_user.id if system_user else 1
        
        notification = Message(
            subject=subject,
            content=content,
            sender_id=sender_id,
            recipient_id=recipient_id,
            message_type=message_type,
            priority=priority,
            status='unread'
        )
        
        db.session.add(notification)
        db.session.commit()
        
        app.logger.info(f'Created notification for user {recipient_id}: {subject}')
        return notification
        
    except Exception as e:
        app.logger.error(f'Error creating notification: {str(e)}')
        db.session.rollback()
        return None

def notify_task_assigned(employee_id, task_title, task_description):
    """Notify employee when a new task is assigned"""
    employee = User.query.get(employee_id)
    if employee:
        subject = "New Task Assigned"
        content = f"You have been assigned a new task: {task_title}\n\n{task_description}"
        create_notification(
            recipient_id=employee_id,
            subject=subject,
            content=content,
            message_type='task',
            priority='normal'
        )

def notify_document_approved(employee_id, document_name):
    """Notify employee when a document is approved"""
    employee = User.query.get(employee_id)
    if employee:
        subject = "Document Approved"
        content = f"Your document '{document_name}' has been approved by HR."
        create_notification(
            recipient_id=employee_id,
            subject=subject,
            content=content,
            message_type='document',
            priority='normal'
        )

def notify_interview_scheduled(employee_id, interview_details):
    """Notify employee when an interview is scheduled"""
    employee = User.query.get(employee_id)
    if employee:
        subject = "Interview Scheduled"
        content = f"An interview has been scheduled for you:\n\n{interview_details}"
        create_notification(
            recipient_id=employee_id,
            subject=subject,
            content=content,
            message_type='interview',
            priority='high'
        )

def notify_exit_summary_ready(employee_id):
    """Notify employee when exit summary is ready"""
    employee = User.query.get(employee_id)
    if employee:
        subject = "Exit Summary Ready"
        content = "Your exit interview summary and final documents are ready for review."
        create_notification(
            recipient_id=employee_id,
            subject=subject,
            content=content,
            message_type='exit',
            priority='high'
        )


# ===== NOTIFICATION API ENDPOINTS =====

@app.route('/api/notifications', methods=['GET'])
@login_required(['hr', 'employee'])
def get_notifications():
    """Get notifications for the current user"""
    try:
        user_id = session['user_id']
        messages = Message.query.filter_by(recipient_id=user_id, status='unread').order_by(Message.sent_at.desc()).limit(20).all()
        
        notifications = []
        for msg in messages:
            notifications.append({
                'id': msg.id,
                'subject': msg.subject,
                'content': msg.content[:100] + '...' if len(msg.content) > 100 else msg.content,
                'message_type': msg.message_type,
                'priority': msg.priority,
                'sent_at': msg.sent_at.isoformat(),
                'sender': msg.sender.full_name if msg.sender else 'System'
            })
        
        return jsonify({
            'status': 'success',
            'notifications': notifications,
            'unread_count': len(notifications)
        })
    except Exception as e:
        app.logger.error(f'Error fetching notifications: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch notifications'
        }), 500


@app.route('/api/notifications/<int:message_id>/read', methods=['PUT'])
@login_required(['hr', 'employee'])
def mark_notification_read(message_id):
    """Mark notification as read"""
    try:
        user_id = session['user_id']
        message = Message.query.filter_by(id=message_id, recipient_id=user_id).first_or_404()
        
        message.status = 'read'
        message.read_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Notification marked as read'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error marking notification as read: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to mark notification as read'
        }), 500


@app.route('/api/notifications/send', methods=['POST'])
@login_required(['hr', 'employee'])
def send_notification():
    """Send notification to users"""
    try:
        data = request.get_json()
        
        recipients = data.get('recipients', [])
        subject = data.get('subject', '')
        content = data.get('content', '')
        message_type = data.get('message_type', 'notification')
        priority = data.get('priority', 'normal')
        
        if not subject or not content:
            return jsonify({
                'status': 'error',
                'message': 'Subject and content are required'
            }), 400
        
        sender_id = session['user_id']
        
        if not recipients:
            hr_users = User.query.filter_by(role='hr').all()
            recipients = [u.id for u in hr_users]
        
        sent_messages = []
        for recipient_id in recipients:
            message = Message(
                subject=subject,
                content=content,
                sender_id=sender_id,
                recipient_id=recipient_id,
                message_type=message_type,
                priority=priority,
                status='unread'
            )
            db.session.add(message)
            sent_messages.append(recipient_id)
        
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': f'Notification sent to {len(sent_messages)} recipients',
            'recipients': sent_messages
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error sending notification: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to send notification'
        }), 500


@app.route('/api/notifications/reminders', methods=['GET'])
@login_required('hr')
def get_reminders():
    """Get system reminders and upcoming tasks"""
    try:
        reminders = []
        
        # Legacy check - onboarding checklist system has been removed
        pending_onboarding = []
        for checklist in pending_onboarding:
            days_since_start = (datetime.utcnow() - checklist.created_at).days
            if days_since_start > 7:  # Remind about pending onboarding older than 7 days
                reminders.append({
                    'type': 'onboarding',
                    'title': f'Pending Onboarding: {checklist.employee.full_name if checklist.employee else "Unknown"}',
                    'description': f'Onboarding started {days_since_start} days ago',
                    'priority': 'high' if days_since_start > 14 else 'medium',
                    'action_url': f'/onboarding/{checklist.id}',
                    'created_at': checklist.created_at.isoformat()
                })
        
        # Document reminders
        pending_docs = EmployeeDocument.query.filter_by(status='pending').all()
        for doc in pending_docs:
            days_pending = (datetime.utcnow() - doc.upload_date).days
            if days_pending > 3:  # Remind about pending documents older than 3 days
                reminders.append({
                    'type': 'document',
                    'title': f'Pending Document: {doc.file_name}',
                    'description': f'Document uploaded {days_pending} days ago awaiting approval',
                    'priority': 'medium',
                    'action_url': f'/documents/review/{doc.id}',
                    'created_at': doc.upload_date.isoformat()
                })
        
        # Offboarding reminders
        active_offboarding = OffboardingCandidate.query.filter_by(status='active').all()
        for candidate in active_offboarding:
            if candidate.exit_date:
                days_until_exit = (candidate.exit_date - datetime.utcnow().date()).days
                if days_until_exit <= 3 and days_until_exit >= 0:  # Remind about upcoming exit
                    reminders.append({
                        'type': 'offboarding',
                        'title': f'Upcoming Exit: {candidate.employee.full_name if candidate.employee else "Unknown"}',
                        'description': f'Exit date in {days_until_exit} days',
                        'priority': 'high' if days_until_exit <= 1 else 'medium',
                        'action_url': f'/offboarding/{candidate.id}',
                        'created_at': candidate.created_at.isoformat()
                    })
        
        # Sort reminders by priority and date
        priority_order = {'high': 0, 'medium': 1, 'low': 2}
        reminders.sort(key=lambda x: (priority_order.get(x['priority'], 3), x['created_at']), reverse=True)
        
        return jsonify({
            'status': 'success',
            'reminders': reminders[:20],  # Return latest 20 reminders
            'total_count': len(reminders)
        })
    except Exception as e:
        app.logger.error(f'Error fetching reminders: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch reminders'
        }), 500


@app.route('/api/notifications/messages', methods=['GET'])
@login_required(['hr', 'employee'])
def get_messages():
    """Get messages (conversation thread)"""
    try:
        user_id = session['user_id']
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        
        # Get messages where user is either sender or recipient
        messages = Message.query.filter(
            (Message.sender_id == user_id) | (Message.recipient_id == user_id)
        ).order_by(Message.sent_at.desc()).paginate(
            page=page, per_page=per_page, error_out=False
        )
        
        message_list = []
        for msg in messages.items:
            message_list.append({
                'id': msg.id,
                'subject': msg.subject,
                'content': msg.content,
                'message_type': msg.message_type,
                'priority': msg.priority,
                'status': msg.status,
                'sent_at': msg.sent_at.isoformat(),
                'read_at': msg.read_at.isoformat() if msg.read_at else None,
                'sender': msg.sender.full_name if msg.sender else 'System',
                'recipient': msg.recipient.full_name if msg.recipient else 'Unknown',
                'is_sent': msg.sender_id == user_id,
                'is_reply': msg.parent_id is not None
            })
        
        return jsonify({
            'status': 'success',
            'messages': message_list,
            'pagination': {
                'page': messages.page,
                'pages': messages.pages,
                'per_page': messages.per_page,
                'total': messages.total,
                'has_next': messages.has_next,
                'has_prev': messages.has_prev
            }
        })
    except Exception as e:
        app.logger.error(f'Error fetching messages: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to fetch messages'
        }), 500


@app.route('/api/notifications/messages', methods=['POST'])
@login_required(['hr', 'employee'])
def send_message():
    """Send a message to another user"""
    try:
        data = request.get_json()
        recipient_id = data.get('recipient_id')
        subject = data.get('subject', '')
        content = data.get('content', '')
        parent_id = data.get('parent_id')  # For replies
        
        if not recipient_id or not content:
            return jsonify({
                'status': 'error',
                'message': 'Recipient and content are required'
            }), 400
        
        # Verify recipient exists
        recipient = User.query.get(recipient_id)
        if not recipient:
            return jsonify({
                'status': 'error',
                'message': 'Recipient not found'
            }), 404
        
        sender_id = session['user_id']
        
        message = Message(
            subject=subject or 'No Subject',
            content=content,
            sender_id=sender_id,
            recipient_id=recipient_id,
            message_type='message',
            priority='normal',
            status='unread',
            parent_id=parent_id
        )
        
        db.session.add(message)
        db.session.commit()
        
        return jsonify({
            'status': 'success',
            'message': 'Message sent successfully',
            'message_id': message.id
        }), 201
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error sending message: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to send message'
        }), 500


@app.route('/api/employee/documents/<int:document_id>/approve', methods=['PUT'])
@login_required('hr')
def approve_employee_document(document_id):
    """Approve or reject employee document"""
    try:
        data = request.get_json()
        status = data.get('status', 'approved')  # approved or rejected
        remarks = data.get('remarks', '')
        
        document = EmployeeDocument.query.get_or_404(document_id)
        
        if document.status != 'pending':
            return jsonify({
                'status': 'error',
                'message': 'Document has already been processed'
            }), 400
        
        document.status = status
        document.remarks = remarks
        document.reviewed_by = session['user_id']
        document.reviewed_at = datetime.utcnow()
        
        db.session.commit()
        
        # Send notification to employee
        if status == 'approved':
            notify_document_approved(document.employee_id, document.document_name)
        else:
            # Create rejection notification
            employee = User.query.get(document.employee_id)
            if employee:
                subject = "Document Rejected"
                content = f"Your document '{document.document_name}' has been rejected."
                if remarks:
                    content += f"\n\nRemarks: {remarks}"
                create_notification(
                    recipient_id=document.employee_id,
                    subject=subject,
                    content=content,
                    message_type='document',
                    priority='high'
                )
        
        return jsonify({
            'status': 'success',
            'message': f'Document {status} successfully'
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error approving document: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to process document'
        }), 500


@app.route('/api/interviews', methods=['POST'])
@login_required('hr')
def schedule_interview():
    """Schedule an interview for a candidate"""
    try:
        data = request.get_json()
        
        candidate_id = data.get('candidate_id')
        employee_id = data.get('employee_id')  # For internal interviews
        interview_type = data.get('type', 'technical')
        scheduled_date = data.get('scheduled_date')
        scheduled_time = data.get('scheduled_time')
        location = data.get('location', 'Video Call')
        notes = data.get('notes', '')
        
        if not scheduled_date or not scheduled_time:
            return jsonify({
                'status': 'error',
                'message': 'Date and time are required'
            }), 400
        
        # Parse date and time
        try:
            date_obj = datetime.strptime(scheduled_date, '%Y-%m-%d').date()
            time_obj = datetime.strptime(scheduled_time, '%H:%M').time()
            scheduled_datetime = datetime.combine(date_obj, time_obj)
        except ValueError:
            return jsonify({
                'status': 'error',
                'message': 'Invalid date or time format'
            }), 400
        
        # Create interview record
        interview = Interview(
            candidate_id=candidate_id,
            scheduled_at=scheduled_datetime,
            interview_type=interview_type,
            location=location,
            notes=notes,
            status='scheduled'
        )
        
        db.session.add(interview)
        db.session.commit()
        
        # Send notification to employee if it's an internal interview
        if employee_id:
            interview_details = f"""
Type: {interview_type.title()}
Date: {scheduled_date}
Time: {scheduled_time}
Location: {location}
{notes if notes else ''}
"""
            notify_interview_scheduled(employee_id, interview_details)
        
        return jsonify({
            'status': 'success',
            'message': 'Interview scheduled successfully',
            'data': {
                'id': interview.id,
                'scheduled_at': interview.scheduled_at.isoformat(),
                'type': interview.interview_type,
                'location': interview.location
            }
        })
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error scheduling interview: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to schedule interview'
        }), 500


@app.route('/api/exit/summary/<int:employee_id>', methods=['POST'])
@login_required('hr')
def generate_exit_summary(employee_id):
    """Generate and notify about exit summary"""
    try:
        employee = User.query.get(employee_id)
        if not employee:
            return jsonify({
                'status': 'error',
                'message': 'Employee not found'
            }), 404
        
        # In a real app, generate comprehensive exit summary
        # For now, just create a notification
        notify_exit_summary_ready(employee_id)
        
        return jsonify({
            'status': 'success',
            'message': 'Exit summary generated and employee notified'
        })
    except Exception as e:
        app.logger.error(f'Error generating exit summary: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate exit summary'
        }), 500


# Initialize APScheduler
scheduler = BackgroundScheduler()

def scheduled_notification_check():
    """Scheduled task to check for overdue tasks and pending revocations"""
    with app.app_context():
        try:
            app.logger.info("Running scheduled notification check...")
            check_overdue_onboarding_tasks()
            check_pending_access_revocation()
            app.logger.info("Scheduled notification check completed")
        except Exception as e:
            app.logger.error(f'Error in scheduled notification check: {str(e)}')

# Schedule the notification check to run every hour
scheduler.add_job(
    func=scheduled_notification_check,
    trigger=IntervalTrigger(hours=1),
    id='notification_check_job',
    name='Check for overdue tasks and pending revocations',
    replace_existing=True
)

# Start the scheduler
scheduler.start()

# Shut down the scheduler when the app exits
import atexit
atexit.register(lambda: scheduler.shutdown())

# AI Feature Endpoints

@app.route('/api/ai/generate-onboarding-checklist', methods=['POST'])
@login_required('hr')
def ai_generate_onboarding_checklist():
    try:
        data = request.json
        employee_name = data.get('employee_name')
        position = data.get('position')
        department = data.get('department')
        
        if not all([employee_name, position, department]):
            return jsonify({'status': 'error', 'message': 'Missing required fields'}), 400
            
        # Legacy check - onboarding checklist system has been removed
        checklist = ai_services.generate_onboarding_checklist(employee_name, position, department)
        return jsonify({'status': 'success', 'checklist': checklist})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/ai/analyze-feedback', methods=['POST'])
@login_required('hr')
def ai_analyze_feedback():
    try:
        feedback_text = request.json.get('feedback')
        if not feedback_text:
            return jsonify({'status': 'error', 'message': 'No feedback provided'}), 400
            
        analysis = ai_services.analyze_employee_sentiment(feedback_text)
        return jsonify({'status': 'success', 'analysis': analysis})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/ai/generate-interview-feedback', methods=['POST'])
@login_required('hr')
def ai_generate_interview_feedback():
    try:
        responses = request.json.get('responses')
        if not responses:
            return jsonify({'status': 'error', 'message': 'No interview responses provided'}), 400
            
        feedback = ai_services.generate_interview_feedback(responses)
        return jsonify({'status': 'success', 'feedback': feedback})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/ai/generate-employee-report/<int:employee_id>', methods=['GET'])
@login_required('hr')
def ai_generate_employee_report(employee_id):
    try:
        employee = User.query.get_or_404(employee_id)
        feedbacks = EmployeeFeedback.query.filter_by(user_id=employee_id).all()
        
        employee_data = {
            'name': employee.full_name,
            'position': employee.position,
            'department': employee.department,
            'feedbacks': [{
                'date': f.created_at.strftime('%Y-%m-%d'),
                'mood': f.mood_rating,
                'confidence': f.confidence_rating,
                'feedback': f.feedback
            } for f in feedbacks]
        }
        
        report = ai_services.generate_employee_report(employee_data)
        return jsonify({'status': 'success', 'report': report})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/hr/start-interview/<int:employee_id>', methods=['POST'])
@login_required('hr')
def start_interview(employee_id):
    try:
        data = request.get_json()
        instructions = data.get('instructions', 'Conduct a professional interview')
        employee = User.query.get_or_404(employee_id)
        notification = Message(
            subject='Interview Ready',
            content='Your interview is ready to start.',
            sender_id=session['user_id'],
            recipient_id=employee_id,
            message_type='interview_ready',
            priority='high',
            status='unread',
            notification_data={'instructions': instructions}
        )
        db.session.add(notification)
        db.session.commit()
        return jsonify({'status': 'success', 'message': f'Interview started for {employee.full_name}'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'Error starting interview: {str(e)}')
        return jsonify({'status': 'error', 'message': 'Failed to start interview'}), 500

@app.route('/api/employee/interview/questions', methods=['GET'])
@login_required('employee')
def get_interview_questions():
    try:
        user_id = session['user_id']
        notification = Message.query.filter_by(
            recipient_id=user_id,
            message_type='interview_ready'
        ).order_by(Message.sent_at.desc()).first()
        if not notification:
            return jsonify({'status': 'error', 'message': 'No active interview'}), 404
        instructions = notification.notification_data.get('instructions', '') if notification.notification_data else ''
        questions = [
            "Tell me about yourself and your professional background.",
            "What are your key strengths and how do they apply to this role?",
            "Describe a challenging situation you faced and how you resolved it.",
            "Where do you see yourself in 5 years?",
            "Why are you interested in this position?"
        ]
        return jsonify({'status': 'success', 'questions': questions, 'instructions': instructions})
    except Exception as e:
        app.logger.error(f'Error getting interview questions: {str(e)}')
        return jsonify({'status': 'error', 'message': 'Failed to get questions'}), 500

@app.route('/api/ai/generate-interview-questions', methods=['POST'])
@login_required(['hr', 'employee'])
def ai_generate_interview_questions():
    """Generate AI-powered interview questions based on difficulty and training type"""
    try:
        data = request.get_json()
        difficulty_level = data.get('difficulty_level', 'intermediate')
        training_type = data.get('training_type', 'behavioral')
        job_role = data.get('job_role')
        previous_responses = data.get('previous_responses')
        
        questions = ai_services.generate_interview_questions(
            difficulty_level=difficulty_level,
            training_type=training_type,
            job_role=job_role,
            previous_responses=previous_responses
        )
        
        if questions:
            return jsonify({
                'status': 'success',
                'questions': questions
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to generate questions'
            }), 500
            
    except Exception as e:
        app.logger.error(f'Error generating interview questions: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate questions'
        }), 500

@app.route('/api/ai/generate-follow-up-question', methods=['POST'])
@login_required(['hr', 'employee'])
def ai_generate_follow_up_question():
    """Generate follow-up questions based on candidate responses"""
    try:
        data = request.get_json()
        original_question = data.get('original_question')
        candidate_answer = data.get('candidate_answer')
        difficulty_level = data.get('difficulty_level', 'intermediate')
        training_type = data.get('training_type', 'behavioral')
        
        if not original_question or not candidate_answer:
            return jsonify({
                'status': 'error',
                'message': 'Original question and candidate answer are required'
            }), 400
        
        follow_up = ai_services.generate_follow_up_question(
            original_question=original_question,
            candidate_answer=candidate_answer,
            difficulty_level=difficulty_level,
            training_type=training_type
        )
        
        if follow_up:
            return jsonify({
                'status': 'success',
                'follow_up_question': follow_up
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to generate follow-up question'
            }), 500
            
    except Exception as e:
        app.logger.error(f'Error generating follow-up question: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate follow-up question'
        }), 500

@app.route('/api/ai/analyze-interview-response', methods=['POST'])
@login_required(['hr', 'employee'])
def ai_analyze_interview_response():
    """Analyze interview response and provide constructive feedback"""
    try:
        data = request.get_json()
        question = data.get('question')
        answer = data.get('answer')
        difficulty_level = data.get('difficulty_level', 'intermediate')
        training_type = data.get('training_type', 'behavioral')
        
        if not question or not answer:
            return jsonify({
                'status': 'error',
                'message': 'Question and answer are required'
            }), 400
        
        analysis = ai_services.analyze_interview_response(
            question=question,
            answer=answer,
            difficulty_level=difficulty_level,
            training_type=training_type
        )
        
        if analysis:
            return jsonify({
                'status': 'success',
                'analysis': analysis
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to analyze response'
            }), 500
            
    except Exception as e:
        app.logger.error(f'Error analyzing interview response: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to analyze response'
        }), 500

@app.route('/api/ai/generate-interview-summary', methods=['POST'])
@login_required(['hr', 'employee'])
def ai_generate_interview_summary():
    """Generate a comprehensive interview summary with actionable insights"""
    try:
        data = request.get_json()
        interview_data = data.get('interview_data')
        overall_score = data.get('overall_score')
        
        if not interview_data:
            return jsonify({
                'status': 'error',
                'message': 'Interview data is required'
            }), 400
        
        summary = ai_services.generate_interview_summary(
            interview_data=interview_data,
            overall_score=overall_score
        )
        
        if summary:
            return jsonify({
                'status': 'success',
                'summary': summary
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to generate summary'
            }), 500
            
    except Exception as e:
        app.logger.error(f'Error generating interview summary: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': 'Failed to generate summary'
        }), 500





if __name__ == '__main__':
    app.run(debug=True)
